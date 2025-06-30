from django.http import HttpResponse, FileResponse
from rest_framework.decorators import api_view, permission_classes
from rest_framework.permissions import AllowAny
from rest_framework.response import Response
from rest_framework import status
from django.db import connection
from .models import Audit
import os
import io
import tempfile
import shutil
from django.conf import settings
import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from .notification_service import NotificationService
import json
from typing import Optional, Dict, Any
from django.views.decorators.csrf import csrf_exempt

@api_view(['GET'])
def generate_audit_report(request, audit_id):
    """
    Generate and download an audit report in DOCX format with tables for each finding
    """
    try:
        print(f"DEBUG: generate_audit_report called for audit_id: {audit_id}")
        
        # Check if a specific version is requested
        version = request.query_params.get('version')
        print(f"DEBUG: Version parameter: {version}")
        
        # If version is provided, check if it has an ApprovedRejected status
        if version:
            with connection.cursor() as cursor:
                cursor.execute("""
                    SELECT ApprovedRejected
                    FROM audit_version
                    WHERE AuditId = %s AND Version = %s
                """, [audit_id, version])
                
                version_row = cursor.fetchone()
                if not version_row:
                    return Response({"error": f"Version {version} not found for audit {audit_id}"}, 
                                   status=status.HTTP_404_NOT_FOUND)
                
                if version_row[0] is None:
                    return Response({"error": f"Version {version} does not have an approved or rejected status"}, 
                                   status=status.HTTP_400_BAD_REQUEST)
        
        # Create a temporary directory for the process
        temp_dir = tempfile.mkdtemp()
        output_file = os.path.join(temp_dir, f"audit_report_{audit_id}_{version if version else 'latest'}.docx")
        
        try:
            # Generate the report file
            report_file = generate_report_file(audit_id, output_file, version)
            
            if not report_file or not os.path.exists(output_file):
                print(f"ERROR: Failed to generate report for audit {audit_id}")
                return Response({"error": "Failed to generate report"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
            
            # Determine file name for download
            if version:
                download_filename = f"audit_report_{audit_id}_v{version}.docx"
            else:
                download_filename = f"audit_report_{audit_id}.docx"
            
            # Return the file for download
            response = FileResponse(
                open(output_file, 'rb'),
                as_attachment=True,
                filename=download_filename
            )
            
            # Notify relevant users about report generation
            try:
                notification_service = NotificationService()
                
                # Get user details
                with connection.cursor() as cursor:
                    cursor.execute("""
                        SELECT 
                            auditor.Email as auditor_email,
                            reviewer.Email as reviewer_email,
                            auditor.UserName as auditor_name
                        FROM 
                            audit a
                        JOIN users auditor ON a.auditor = auditor.UserId
                        LEFT JOIN users reviewer ON a.reviewer = reviewer.UserId
                        WHERE a.AuditId = %s
                    """, [audit_id])
                    
                    user_row = cursor.fetchone()
                    if user_row:
                        auditor_email, reviewer_email, auditor_name = user_row
                        
                        # Notify the user who generated the report (could be auditor or reviewer)
                        user_email = request.user.email if hasattr(request, 'user') and hasattr(request.user, 'email') else auditor_email
                        
                        notification_data = {
                            'notification_type': 'reportGenerated',
                            'email': user_email,
                            'email_type': 'gmail',
                            'template_data': [
                                "User",  # Generic name since we don't know exactly who generated
                                f"Audit #{audit_id}",
                                version if version else "latest",
                                datetime.datetime.now().strftime('%Y-%m-%d %H:%M')
                            ]
                        }
                        notification_service.send_multi_channel_notification(notification_data)
            except Exception as e:
                print(f"Failed to send notification: {str(e)}")
            
            return response
            
        finally:
            # Clean up temporary directory after sending response
            shutil.rmtree(temp_dir, ignore_errors=True)
    
    except Exception as e:
        print(f"ERROR in generate_audit_report: {str(e)}")
        import traceback
        traceback.print_exc()
        return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

def get_audit_data(audit_id: int) -> Optional[Dict[str, Any]]:
    """
    Get all necessary audit data for report generation
    """
    try:
        with connection.cursor() as cursor:
            # Get audit details
            cursor.execute("""
                SELECT 
                    a.Title, a.Scope, a.Objective, a.BusinessUnit,
                    a.Evidence as AuditEvidence, a.Comments as OverallAuditComments,
                    a.ReviewerComments as OverallReviewComments,
                    p.PolicyName as Policy, sp.SubPolicyName as SubPolicy, 
                    f.FrameworkName as Framework,
                    a.CompletionDate, a.ReviewDate
                FROM audit a
                LEFT JOIN policies p ON a.PolicyId = p.PolicyId
                LEFT JOIN subpolicies sp ON a.SubPolicyId = sp.SubPolicyId
                LEFT JOIN frameworks f ON a.FrameworkId = f.FrameworkId
                WHERE a.AuditId = %s
            """, [audit_id])
            audit_data = cursor.fetchone()
            
            if not audit_data:
                return None
                
            # Get audit findings with compliance details
            cursor.execute("""
                SELECT 
                    af.ComplianceId,
                    af.MajorMinor as TypeOfFinding,
                    af.CheckedDate as ReviewDate,
                    af.Evidence as ComplianceEvidence,
                    af.HowToVerify,
                    af.Impact,
                    af.DetailsOfFinding,
                    af.Comments,
                    af.Check as ReviewStatus,
                    af.ReviewComments,
                    COALESCE(af.SeverityRating, 0) as SeverityRating,
                    af.PredictiveRisks,
                    af.CorrectiveActions,
                    af.UnderlyingCause,
                    af.WhyToVerify,
                    af.WhatToVerify,
                    af.SuggestedActionPlan,
                    af.AssignedDate as MitigationDate,
                    af.ResponsibleForPlan,
                    COALESCE(af.ReAudit, 0) as ReAudit,
                    af.ReAuditDate,
                    af.Check as ComplianceStatus,
                    c.ComplianceTitle,
                    c.ComplianceItemDescription
                FROM audit_findings af
                LEFT JOIN compliance c ON af.ComplianceId = c.ComplianceId
                WHERE af.AuditId = %s
            """, [audit_id])
            findings = cursor.fetchall()
            
            return {
                'audit': audit_data,
                'findings': findings
            }
    except Exception as e:
        print(f"Error getting audit data: {str(e)}")
        return None

def generate_report_file(audit_id: int, output_path: str, version=None) -> Optional[str]:
    """
    Generate audit compliance report in Word format
    """
    try:
        # Get audit data
        data = get_audit_data(audit_id)
        if not data:
            return None
            
        audit_data = data['audit']
        findings = data['findings']
        
        # Create document
        doc = Document()
        
        # Add title
        title = doc.add_heading('Audit Compliance Report', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add a table for each finding
        for finding in findings:
            doc.add_heading(f'Compliance Item: {finding[22]}', level=2)  # ComplianceTitle
            
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            header_cells = table.rows[0].cells
            header_cells[0].text = 'ITEM'
            header_cells[1].text = 'DETAILS'
            
            # Map check status to readable format
            check_status_map = {
                '0': 'Not Started',
                '1': 'In Progress',
                '2': 'Completed',
                '3': 'Not Applicable'
            }
            
            # Map major/minor to readable format
            major_minor_map = {
                '0': 'Minor',
                '1': 'Major',
                '2': 'Not Applicable'
            }
            
            # Add finding details
            items = [
                ('Type of Finding', major_minor_map.get(finding[1], '') if finding[1] else ''),
                ('Finding ID', str(finding[0])),  # ComplianceId
                ('Date and Time', finding[2].strftime('%Y-%m-%d %H:%M:%S') if finding[2] else ''),
                ('Framework', audit_data[9]),  # Framework name
                ('Policy', audit_data[7]),  # Policy name
                ('Subpolicy', audit_data[8]),  # SubPolicy name
                ('Compliance Status', check_status_map.get(finding[21], '')),
                ('Type of Findings', major_minor_map.get(finding[1], '')),
                ('Severity Rating', str(finding[10])),
                ('What to Verify', finding[15] or ''),
                ('How to Verify', finding[4] or ''),
                ('Why to Verify', finding[14] or ''),
                ('Details of Findings', finding[6] or ''),
                ('Underlying Cause', finding[13] or ''),
                ('Impact', finding[5] or ''),
                ('Predictive Risks', finding[11] or '[]'),
                ('Corrective Actions', finding[12] or '[]'),
                ('Suggested Action Plan', finding[16] or ''),
                ('Responsible for Plan', finding[18] or ''),
                ('Mitigation Date', finding[17].strftime('%Y-%m-%d') if finding[17] else ''),
                ('Re-audit Required', 'Yes' if finding[19] else 'No'),
                ('Re-audit Date', finding[20].strftime('%Y-%m-%d') if finding[20] else ''),
                ('Comments', finding[7] or ''),
                ('Review Status', check_status_map.get(finding[8], '')),
                ('Review Comments', finding[9] or ''),
                ('Audit Evidence', audit_data[4] or ''),
                ('Compliance Evidence', finding[3] or ''),
                ('Overall Audit Comments', audit_data[5] or ''),
                ('Overall Review Comments', audit_data[6] or '')
            ]
            
            for item in items:
                row_cells = table.add_row().cells
                row_cells[0].text = item[0]
                row_cells[1].text = str(item[1] if item[1] is not None else '')
            
            # Add a page break after each finding except the last one
            if finding != findings[-1]:
                doc.add_page_break()
        
        # Save document
        doc.save(output_path)
        return output_path
        
    except Exception as e:
        print(f"Error generating report: {str(e)}")
        import traceback
        traceback.print_exc()
        return None

def create_incidents_for_findings(audit_id: int) -> None:
    """
    Create incidents for non-compliant and partially compliant audit findings
    """
    try:
        with connection.cursor() as cursor:
            # Get relevant audit findings and compliance details
            cursor.execute("""
                SELECT 
                    af.ComplianceId,
                    af.Check,
                    af.Comments,
                    af.user_id,
                    c.ComplianceTitle,
                    c.ComplianceItemDescription,
                    c.PossibleDamage,
                    c.Mitigation
                FROM audit_findings af
                JOIN compliance c ON af.ComplianceId = c.ComplianceId
                WHERE af.AuditId = %s 
                AND (
                    (af.Check = '0') OR  -- Not Compliant
                    (af.Check = '1')     -- Partially Compliant
                )
            """, [audit_id])
            
            findings = cursor.fetchall()
            
            current_datetime = datetime.datetime.now()
            current_date = current_datetime.date()
            current_time = current_datetime.time()
            
            # Create incidents for each finding
            for finding in findings:
                compliance_id = finding[0]
                check_status = finding[1]
                comments = finding[2]
                user_id = finding[3]
                compliance_title = finding[4]
                compliance_desc = finding[5]
                possible_damage = finding[6]
                mitigation = finding[7]
                
                # Check if incident already exists for this finding
                cursor.execute("""
                    SELECT COUNT(*) 
                    FROM incidents 
                    WHERE AuditId = %s AND ComplianceId = %s
                """, [audit_id, compliance_id])
                
                if cursor.fetchone()[0] > 0:
                    print(f"Incident already exists for AuditId {audit_id} and ComplianceId {compliance_id}")
                    continue
                
                # Create new incident
                cursor.execute("""
                    INSERT INTO incidents (
                        IncidentTitle,
                        Description,
                        PossibleDamage,
                        Mitigation,
                        AuditId,
                        ComplianceId,
                        Date,
                        Time,
                        UserId,
                        Origin,
                        Comments,
                        Status
                    ) VALUES (
                        %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s
                    )
                """, [
                    f"Non-Compliance of {compliance_title}",
                    compliance_desc,
                    possible_damage,
                    mitigation,
                    audit_id,
                    compliance_id,
                    current_date,
                    current_time,
                    user_id,
                    "Audit Finding",
                    comments,
                    "Open"
                ])
                
                print(f"Created incident for ComplianceId {compliance_id} in AuditId {audit_id}")
                
    except Exception as e:
        print(f"Error creating incidents: {str(e)}")
        import traceback
        traceback.print_exc()

@csrf_exempt
@api_view(['POST'])
@permission_classes([AllowAny])
def approve_audit_and_create_incidents(request, audit_id):
    """
    API endpoint to approve audit and create incidents for non-compliant findings
    """
    try:
        with connection.cursor() as cursor:
            # Update audit status
            cursor.execute("""
                UPDATE audit 
                SET Status = 'Approved',
                    CompletionDate = %s
                WHERE AuditId = %s
            """, [datetime.datetime.now(), audit_id])
            
            # Update LastChecklistItemVerified in audit_findings
            cursor.execute("""
                UPDATE audit_findings
                SET LastChecklistItemVerified = %s
                WHERE AuditId = %s
            """, [datetime.datetime.now(), audit_id])
            
            # Create incidents for non-compliant and partially compliant findings
            create_incidents_for_findings(audit_id)
            
            return Response({
                "message": "Audit approved and incidents created successfully",
                "audit_id": audit_id
            }, status=status.HTTP_200_OK)
            
    except Exception as e:
        print(f"Error in approve_audit_and_create_incidents: {str(e)}")
        import traceback
        traceback.print_exc()
        return Response({
            "error": str(e)
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR) 