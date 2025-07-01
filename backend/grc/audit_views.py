from rest_framework.decorators import api_view
from rest_framework.response import Response
from rest_framework import status
from django.conf import settings
from .models import Audit, Framework, Policy, Users, SubPolicy, Compliance, AuditFinding
from .serializers import FrameworkSerializer, PolicySerializer, UserSerializer, PolicyAllocationSerializer
from django.db import connection
from django.utils import timezone
import json
import random
from django.http import HttpResponse
import io
from docx import Document
from docx.shared import Inches, Pt
import os
from .notification_service import NotificationService

@api_view(['GET'])
def get_frameworks(request):
    """
    Get all frameworks
    """
    try:
        frameworks = Framework.objects.all()
        serializer = FrameworkSerializer(frameworks, many=True)
        return Response(serializer.data, status=status.HTTP_200_OK)
    except Exception as e:
        return Response({'error': str(e)}, status=status.HTTP_400_BAD_REQUEST)

@api_view(['GET'])
def get_policies_by_framework(request, framework_id):
    """
    Get all policies for a specific framework
    """
    try:
        policies = Policy.objects.filter(FrameworkId=framework_id)
        serializer = PolicySerializer(policies, many=True)
        return Response(serializer.data, status=status.HTTP_200_OK)
    except Exception as e:
        return Response({'error': str(e)}, status=status.HTTP_400_BAD_REQUEST)

 
def get_compliances_by_scope(framework_id, policy_id=None, subpolicy_id=None):
    """
    Helper function to get compliances based on assignment scope
    """
    try:
        if subpolicy_id:
            # Case 3: Specific SubPolicy
            return Compliance.objects.filter(SubPolicyId=subpolicy_id)
        elif policy_id:
            # Case 2: Specific Policy - get all compliances from its subpolicies
            subpolicies = SubPolicy.objects.filter(PolicyId=policy_id)
            return Compliance.objects.filter(SubPolicyId__in=subpolicies)
        else:
            # Case 1: Entire Framework - get all compliances from all policies and their subpolicies
            policies = Policy.objects.filter(FrameworkId=framework_id)
            subpolicies = SubPolicy.objects.filter(PolicyId__in=policies)
            return Compliance.objects.filter(SubPolicyId__in=subpolicies)
    except Exception as e:
        print(f"Error in get_compliances_by_scope: {str(e)}")
        return Compliance.objects.none()



@api_view(['GET'])
def get_assign_data(request):
    """
    Fetch frameworks, policies, subpolicies, and users for assignment data
    """
    try:
        frameworks = Framework.objects.all().values('FrameworkId', 'FrameworkName')
        policies = Policy.objects.all().values('PolicyId', 'PolicyName', 'FrameworkId_id')
        subpolicies = SubPolicy.objects.all().values('SubPolicyId', 'SubPolicyName', 'PolicyId_id')
        users = Users.objects.all().values('UserId', 'UserName')

        # Transform field names for frontend compatibility
        formatted_policies = []
        for policy in policies:
            formatted_policies.append({
                'PolicyId': policy['PolicyId'],
                'PolicyName': policy['PolicyName'],
                'FrameworkId': policy['FrameworkId_id']
            })

        formatted_subpolicies = []
        for subpolicy in subpolicies:
            formatted_subpolicies.append({
                'SubPolicyId': subpolicy['SubPolicyId'],
                'SubPolicyName': subpolicy['SubPolicyName'],
                'PolicyId': subpolicy['PolicyId_id']
            })

        return Response({
            'frameworks': list(frameworks),
            'policies': formatted_policies,
            'subpolicies': formatted_subpolicies,
            'users': list(users)
        }, status=status.HTTP_200_OK)
    except Exception as e:
        return Response({'error': str(e)}, status=status.HTTP_400_BAD_REQUEST)

 
@api_view(['GET'])
def get_all_audits(request):
    """
    Fetch all audits with related data for display in the audit table
    """
    try:
        print("DEBUG: get_all_audits was called")
        # Using raw SQL for better performance and to join multiple tables
        with connection.cursor() as cursor:
            print("DEBUG: Executing SQL query for get_all_audits")
            print("DEBUG: Dumping table schemas to verify column names")
            try:
                cursor.execute("DESCRIBE audit")
                audit_columns = cursor.fetchall()
                print("DEBUG: Audit table columns:", audit_columns)
            except Exception as e:
                print("DEBUG: Error describing audit table:", str(e))
                
            cursor.execute("""
                SELECT 
                    a.AuditId as audit_id,
                    f.FrameworkName as framework,
                    p.PolicyName as policy,
                    sp.SubPolicyName as subpolicy,
                    auditor_user.UserName as auditor,
                    a.DueDate as duedate,
                    a.Frequency as frequency,
                    reviewer_user.UserName as reviewer,
                    a.AuditType as audit_type,
                    CASE 
                        WHEN COUNT(af.AuditId) = 0 THEN 'Yet to Start'
                        WHEN COUNT(af.AuditId) > 0 AND SUM(CASE WHEN af.`Check` = '2' THEN 1 ELSE 0 END) = COUNT(af.AuditId) THEN 'Completed'
                        ELSE 'Work In progress'
                    END as status
                FROM 
                    audit a
                LEFT JOIN 
                    frameworks f ON a.FrameworkId = f.FrameworkId
                LEFT JOIN 
                    policies p ON a.PolicyId = p.PolicyId
                LEFT JOIN 
                    subpolicies sp ON a.SubPolicyId = sp.SubPolicyId
                LEFT JOIN 
                    users auditor_user ON a.auditor = auditor_user.UserId
                LEFT JOIN 
                    users reviewer_user ON a.reviewer = reviewer_user.UserId
                LEFT JOIN 
                    audit_findings af ON a.AuditId = af.AuditId
                GROUP BY 
                    a.AuditId, f.FrameworkName, p.PolicyName, sp.SubPolicyName, 
                    auditor_user.UserName, a.DueDate, a.Frequency, reviewer_user.UserName, a.AuditType
                ORDER BY 
                    a.AuditId DESC
            """)
            print("DEBUG: SQL query executed successfully")
            columns = [col[0] for col in cursor.description]
            print(f"DEBUG: Result columns: {columns}")
            audits = [dict(zip(columns, row)) for row in cursor.fetchall()]
            print(f"DEBUG: Fetched {len(audits)} audit records")

        # Process frequency to display text instead of number
        for audit in audits:
            # Format date
            if audit.get('duedate'):
                audit['duedate'] = audit['duedate'].strftime('%d/%m/%Y')
            
            # Convert frequency number to text
            freq = audit.get('frequency')
            if freq is not None:
                if freq == 0:
                    audit['frequency'] = 'Only Once'
                elif freq == 1:
                    audit['frequency'] = 'Daily'
                elif freq <= 60:
                    audit['frequency'] = 'Every 2 Months'
                elif freq <= 120:
                    audit['frequency'] = 'Every 4 Months'
                elif freq <= 182:
                    audit['frequency'] = 'Half Yearly'
                elif freq <= 365:
                    audit['frequency'] = 'Yearly'
                else:
                    audit['frequency'] = f'Every {freq} days'
            
            # Convert audit type from I/E to Internal/External
            if audit.get('audit_type') == 'I':
                audit['audit_type'] = 'Internal'
            elif audit.get('audit_type') == 'E':
                audit['audit_type'] = 'External'
            
            # Add report field
            audit['report'] = 'Download' if audit.get('status') == 'Completed' else 'Pending'

        return Response(audits, status=status.HTTP_200_OK)
    except Exception as e:
        print(f"ERROR in get_all_audits: {str(e)}")
        return Response({'error': str(e)}, status=status.HTTP_400_BAD_REQUEST)

@api_view(['GET'])
def get_my_audits(request):
    """
    Fetch audits assigned to the current user (as auditor)
    If user_id is not in session, default to user 1050
    """
    try:
        print("DEBUG: get_my_audits was called")
        
        # Get user_id from session, default to 1050 if not found
        user_id = request.session.get('user_id', 1050)
        print(f"DEBUG: Using user_id: {user_id}")
        
        # Using raw SQL to join multiple tables and get comprehensive data
        with connection.cursor() as cursor:
            print(f"DEBUG: Executing SQL query for my audits")
            cursor.execute("""
                SELECT 
                    a.AuditId as audit_id,
                    f.FrameworkName as framework,
                    p.PolicyName as policy,
                    sp.SubPolicyName as subpolicy,
                    a.DueDate as duedate,
                    a.Frequency as frequency,
                    reviewer_user.UserName as reviewer,
                    a.AuditType as audit_type,
                    assignee_user.UserName as assignee,
                    a.Status as status_from_db,
                    COALESCE(a.Reports, '[]') as reports,
                    CASE 
                        WHEN COUNT(af.AuditId) = 0 THEN 'Yet to Start'
                        WHEN COUNT(af.AuditId) > 0 AND SUM(CASE WHEN af.`Check` = '2' THEN 1 ELSE 0 END) = COUNT(af.AuditId) THEN 'Completed'
                        ELSE 'Work In progress'
                    END as calculated_status,
                    COUNT(af.AuditId) as total_compliances,
                    SUM(CASE WHEN af.`Check` = '2' THEN 1 ELSE 0 END) as completed_compliances,
                    a.CompletionDate as completion_date
                FROM 
                    audit a
                LEFT JOIN 
                    frameworks f ON a.FrameworkId = f.FrameworkId
                LEFT JOIN 
                    policies p ON a.PolicyId = p.PolicyId
                LEFT JOIN 
                    subpolicies sp ON a.SubPolicyId = sp.SubPolicyId
                LEFT JOIN 
                    users assignee_user ON a.assignee = assignee_user.UserId
                LEFT JOIN 
                    users reviewer_user ON a.reviewer = reviewer_user.UserId
                LEFT JOIN 
                    audit_findings af ON a.AuditId = af.AuditId
                WHERE 
                    a.auditor = %s
                GROUP BY 
                    a.AuditId, f.FrameworkName, p.PolicyName, sp.SubPolicyName, 
                    a.DueDate, a.Frequency, reviewer_user.UserName, a.AuditType, assignee_user.UserName, a.Status, a.CompletionDate, a.Reports
                ORDER BY 
                    a.DueDate ASC
            """, [user_id])
            
            columns = [col[0] for col in cursor.description]
            print(f"DEBUG: My audits query columns: {columns}")
            audits = [dict(zip(columns, row)) for row in cursor.fetchall()]
            print(f"DEBUG: Fetched {len(audits)} my audits")

        # Process and format audit data for display
        for audit in audits:
            # Format date
            if audit.get('duedate'):
                audit['duedate'] = audit['duedate'].strftime('%d/%m/%Y')
            
            # Format completion date if present
            if audit.get('completion_date'):
                audit['completion_date'] = audit['completion_date'].strftime('%d/%m/%Y %H:%M')
            
            # Calculate completion percentage
            total = audit.get('total_compliances') or 0
            completed = audit.get('completed_compliances') or 0
            audit['completion_percentage'] = round((completed / total) * 100) if total > 0 else 0
            
            # Use the database status instead of calculated status
            audit['status'] = audit.get('status_from_db') or audit.get('calculated_status') or 'Yet to Start'
            print(f"DEBUG: Status  for now: {audit['status']}---------------========================================")
            
            # Convert frequency number to text
            freq = audit.get('frequency')
            if freq is not None:
                if freq == 0:
                    audit['frequency_text'] = 'Only Once'
                elif freq == 1:
                    audit['frequency_text'] = 'Daily'
                elif freq <= 60:
                    audit['frequency_text'] = 'Every 2 Months'
                elif freq <= 120:
                    audit['frequency_text'] = 'Every 4 Months'
                elif freq <= 182:
                    audit['frequency_text'] = 'Half Yearly'
                elif freq <= 365:
                    audit['frequency_text'] = 'Yearly'
                else:
                    audit['frequency_text'] = f'Every {freq} days'
            
            # Convert audit type from I/E to Internal/External
            if audit.get('audit_type') == 'I':
                audit['audit_type_text'] = 'Internal'
            elif audit.get('audit_type') == 'E':
                audit['audit_type_text'] = 'External'
            
            # Add report status based on Reports field
            reports_data = audit.get('reports')
            print(f"Reports data for audit {audit.get('audit_id')}: {reports_data}")
            try:
                if isinstance(reports_data, str):
                    reports_data = json.loads(reports_data) if reports_data else []
                audit['reports'] = reports_data
                audit['report_available'] = bool(reports_data and reports_data != [] and reports_data != {})
            except (json.JSONDecodeError, TypeError):
                audit['reports'] = []
                audit['report_available'] = False
            print(f"Report available for audit {audit.get('audit_id')}: {audit['report_available']}")

        print(f"DEBUG: Successfully prepared my audits response")


        print(f"DEBUG: My audits response: {audits}---------------------------------------------{user_id}")

    
        return Response({
            'user_id': user_id,
            'audits': audits
        }, status=status.HTTP_200_OK)
    except Exception as e:
        print(f"ERROR in get_my_audits: {str(e)}")
        return Response({'error': str(e)}, status=status.HTTP_400_BAD_REQUEST)
 

@api_view(['GET'])
def get_audit_details(request, audit_id):
    """
    Fetch detailed information for a specific audit including compliance items
    """
    try:

        print("api--------------------------------------------------------",audit_id)

        request.session['current_audit_id'] = audit_id


        print(f"DEBUG: get_audit_details was called for audit_id {audit_id}", "this is clicked by the user")
        
        # Check if audit exists
        try:
            audit = Audit.objects.get(AuditId=audit_id)
            print(f"DEBUG: Found audit record with ID {audit_id}")
        except Audit.DoesNotExist:
            print(f"DEBUG: Audit with ID {audit_id} not found")
            return Response({'error': 'Audit not found'}, status=status.HTTP_404_NOT_FOUND)
        
        # Get basic audit information
        with connection.cursor() as cursor:
            print(f"DEBUG: Executing SQL query to get audit details")
            cursor.execute("""
                SELECT 
                    a.AuditId as audit_id,
                    f.FrameworkName as framework,
                    p.PolicyName as policy,
                    sp.SubPolicyName as subpolicy,
                    auditor_user.UserName as auditor,
                    auditor_user.UserId as auditor_id,
                    a.DueDate as duedate,
                    a.Frequency as frequency,
                    reviewer_user.UserName as reviewer,
                    reviewer_user.UserId as reviewer_id,
                    a.AuditType as audit_type,
                    assignee_user.UserName as assignee,
                    assignee_user.UserId as assignee_id,
                    a.Status as status_from_db,
                    a.CompletionDate as completion_date,
                    COALESCE(a.Reports, '[]') as reports,
                    CASE 
                        WHEN COUNT(af.AuditId) = 0 THEN 'Yet to Start'
                        WHEN COUNT(af.AuditId) > 0 AND SUM(CASE WHEN af.`Check` = '2' THEN 1 ELSE 0 END) = COUNT(af.AuditId) THEN 'Completed'
                        ELSE 'Work In Progress'
                    END as calculated_status,
                    COUNT(af.AuditId) as total_compliances,
                    SUM(CASE WHEN af.`Check` = '2' THEN 1 ELSE 0 END) as completed_compliances
                FROM 
                    audit a
                LEFT JOIN 
                    frameworks f ON a.FrameworkId = f.FrameworkId
                LEFT JOIN 
                    policies p ON a.PolicyId = p.PolicyId
                LEFT JOIN 
                    subpolicies sp ON a.SubPolicyId = sp.SubPolicyId
                LEFT JOIN 
                    users auditor_user ON a.auditor = auditor_user.UserId
                LEFT JOIN 
                    users reviewer_user ON a.reviewer = reviewer_user.UserId
                LEFT JOIN 
                    users assignee_user ON a.assignee = assignee_user.UserId
                LEFT JOIN 
                    audit_findings af ON a.AuditId = af.AuditId
                WHERE 
                    a.AuditId = %s
                GROUP BY 
                    a.AuditId, f.FrameworkName, p.PolicyName, sp.SubPolicyName, 
                    auditor_user.UserName, auditor_user.UserId, a.DueDate, a.Frequency, 
                    reviewer_user.UserName, reviewer_user.UserId, a.AuditType, 
                    assignee_user.UserName, assignee_user.UserId, a.Status, a.CompletionDate, a.Reports
            """, [audit_id])
            
            columns = [col[0] for col in cursor.description]
            print(f"DEBUG: Audit details query columns: {columns}")
            audit_data = dict(zip(columns, cursor.fetchone())) if cursor.rowcount > 0 else None
            print(f"DEBUG: Fetched audit details: {audit_data is not None}")
        
        if not audit_data:
            return Response({'error': 'Audit details not found'}, status=status.HTTP_404_NOT_FOUND)
        
        # Format date
        if audit_data.get('duedate'):
            audit_data['duedate'] = audit_data['duedate'].strftime('%Y-%m-%d')
            
        # Format completion date if present
        if audit_data.get('completion_date'):
            audit_data['completion_date'] = audit_data['completion_date'].strftime('%Y-%m-%d %H:%M:%S')
        
        # Use the database status instead of calculated status
        audit_data['status'] = audit_data.get('status_from_db') or audit_data.get('calculated_status') or 'Yet to Start'
        
        # Get compliance items for this audit
        with connection.cursor() as cursor:
            print(f"DEBUG: Executing SQL query for compliance items")
            cursor.execute("""
                SELECT 
                    af.AuditId as audit_id,
                    af.ComplianceId as ComplianceId,
                    c.ComplianceItemDescription,
                    c.IsRisk,
                    c.Criticality,
                    c.MandatoryOptional,
                    af.`Check` as status,
                    af.Comments as comment,
                    af.CheckedDate as checked_date
                FROM 
                    audit_findings af
                JOIN
                    compliance c ON af.ComplianceId = c.ComplianceId
                WHERE 
                    af.AuditId = %s
                ORDER BY
                    c.ComplianceId
            """, [audit_id])
            
            columns = [col[0] for col in cursor.description]
            print(f"DEBUG: Compliance items query columns: {columns}")
            compliance_items = [dict(zip(columns, row)) for row in cursor.fetchall()]
            print(f"DEBUG: Fetched {len(compliance_items)} compliance items")
        
        # Process compliance items
        for item in compliance_items:
            # Convert status from 0,1,2 to text
            if item['status'] == '0':
                item['status_text'] = 'Yet to Start'
            elif item['status'] == '1':
                item['status_text'] = 'In Progress'
            elif item['status'] == '2':
                item['status_text'] = 'Completed'
            
            # Format date if present
            if item.get('checked_date'):
                item['checked_date'] = item['checked_date'].strftime('%Y-%m-%d %H:%M:%S')
        
        # Combine audit data with compliance items
        audit_data['compliance_items'] = compliance_items
        
        print(f"DEBUG: Successfully prepared audit details response")
        return Response(audit_data, status=status.HTTP_200_OK)
    except Exception as e:
        print(f"ERROR in get_audit_details: {str(e)}")
        return Response({'error': str(e)}, status=status.HTTP_400_BAD_REQUEST)
 

def copy_review_data_from_r_to_a(audit_id, audit_data):
    """
    Copy review data (review_status, review_comments) from the latest R version
    to the audit_data (used for creating A versions)
    """
    try:
        with connection.cursor() as cursor:
            # Get the latest R version
            cursor.execute("""
                SELECT Version, ExtractedInfo FROM audit_version 
                WHERE AuditId = %s AND Version LIKE 'R%'
                ORDER BY Version DESC
                LIMIT 1
            """, [audit_id])
            
            r_version_row = cursor.fetchone()
            if not r_version_row:
                print(f"DEBUG: No R version found for audit {audit_id}")
                return audit_data
            
            latest_r_version = r_version_row[0]
            
            # Parse JSON data
            latest_review_data = None
            if isinstance(r_version_row[1], dict):
                latest_review_data = r_version_row[1]
            else:
                latest_review_data = json.loads(r_version_row[1])
            
            print(f"DEBUG: Found latest R version {latest_r_version} to copy review data from")
            
            # Copy review data (review_status, review_comments) from R version to audit_data
            if not latest_review_data:
                print(f"DEBUG: R version {latest_r_version} has no data")
                return audit_data
                
            for compliance_id, finding in latest_review_data.items():
                # Skip metadata and non-compliance entries
                if compliance_id == '__metadata__' or compliance_id == 'overall_comments' or not isinstance(finding, dict):
                    continue
                    
                # If this compliance exists in our audit data, copy review fields
                if compliance_id in audit_data:
                    # Copy review_status if available
                    if 'review_status' in finding:
                        audit_data[compliance_id]['review_status'] = finding['review_status']
                        print(f"DEBUG: Copied review_status '{finding['review_status']}' for compliance {compliance_id}")
                    
                    # Copy review_comments if available
                    if 'review_comments' in finding:
                        audit_data[compliance_id]['review_comments'] = finding['review_comments']
                        print(f"DEBUG: Copied review_comments for compliance {compliance_id}")
                        
                    # Copy accept_reject value if available
                    if 'accept_reject' in finding:
                        audit_data[compliance_id]['accept_reject'] = finding['accept_reject']
                        print(f"DEBUG: Copied accept_reject '{finding['accept_reject']}' for compliance {compliance_id}")
                        
                    # Copy reviewer_comments field if available (alternate name)
                    if 'reviewer_comments' in finding:
                        audit_data[compliance_id]['reviewer_comments'] = finding['reviewer_comments']
            
            # Copy overall comments if available
            if 'overall_comments' in latest_review_data:
                audit_data['overall_comments'] = latest_review_data['overall_comments']
                print(f"DEBUG: Copied overall comments from R version")
                
        return audit_data
    except Exception as e:
        print(f"DEBUG: Error copying review data from R version: {str(e)}")
        # Return original data if there's an error
        return audit_data

def create_audit_version(audit_id, user_id, custom_version=None):
    """
    Create a new version of an audit's findings in the audit_version table
    """
    try:
        print(f"DEBUG: Creating audit version for audit_id: {audit_id}, user_id: {user_id}")
        
        # Always get the next version number instead of using fixed version
        next_version = get_next_version_number(audit_id, "A")
        version = next_version if custom_version is None else custom_version
        
        print(f"DEBUG: Using version: {version}")
        
        # Get the audit findings
        audit_data = get_audit_findings_json(audit_id)
        
        # Copy review data from latest R version if available
        audit_data = copy_review_data_from_r_to_a(audit_id, audit_data)
        
        # Add metadata
        audit_data['__metadata__'] = {
            'version_date': timezone.now().strftime("%Y-%m-%d %H:%M:%S"),
            'auditor_id': user_id,
            'version_type': 'Auditor'
        }
        
        # Format the JSON
        json_data = json.dumps(audit_data, indent=2)
        print(f"DEBUG: Formatted JSON structure for audit version:\n{json_data}")
        
        # Insert into the audit_version table
        with connection.cursor() as cursor:
            # Check the actual column names in the audit_version table
            cursor.execute("DESCRIBE audit_version")
            columns = [column[0] for column in cursor.fetchall()]
            print(f"DEBUG: Available columns in audit_version table: {columns}")
            
            # Determine which column should store the JSON data
            # Based on the column names, it's likely 'ExtractedInfo' that should store the JSON
            data_column = 'ExtractedInfo'
            
            # Execute the query with the correct column names
            cursor.execute(
                f"""
                INSERT INTO audit_version (AuditId, Version, {data_column}, UserId, Date) 
                VALUES (%s, %s, %s, %s, NOW())
                """,
                [audit_id, version, json_data, user_id,timezone.now()]
            )
            
        print(f"DEBUG: Created new audit version {version} for audit {audit_id}")
        return version
    except Exception as e:
        if "Duplicate entry" in str(e):
            print(f"DEBUG: Version {version} already exists for audit {audit_id}, getting next version")
            return create_audit_version(audit_id, user_id)
        print(f"ERROR: Failed to create audit version: {str(e)}")
        return None
 
@api_view(['GET'])
def check_audit_reports(request):
    """
    Check for existing audit reports based on framework, policy, and subpolicy IDs
    """
    try:
        framework_id = request.GET.get('framework_id')
        policy_id = request.GET.get('policy_id')
        subpolicy_id = request.GET.get('subpolicy_id')

        if not framework_id:
            return Response({'error': 'Framework ID is required'}, status=status.HTTP_400_BAD_REQUEST)

        # Build query based on provided parameters
        query = """
            SELECT 
                ar.ReportId,
                a.CompletionDate,
                auditor.UserName as AuditorName,
                reviewer.UserName as ReviewerName
            FROM 
                audit_report ar
                JOIN audit a ON ar.AuditId = a.AuditId
                JOIN users auditor ON a.auditor = auditor.UserId
                LEFT JOIN users reviewer ON a.reviewer = reviewer.UserId
            WHERE 
                ar.FrameworkId = %s
        """
        params = [framework_id]

        if policy_id:
            query += " AND ar.PolicyId = %s"
            params.append(policy_id)
        else:
            query += " AND ar.PolicyId IS NULL"

        if subpolicy_id:
            query += " AND ar.SubPolicyId = %s"
            params.append(subpolicy_id)
        else:
            query += " AND ar.SubPolicyId IS NULL"

        with connection.cursor() as cursor:
            cursor.execute(query, params)
            columns = [col[0] for col in cursor.description]
            reports = [dict(zip(columns, row)) for row in cursor.fetchall()]

            # Format dates
            for report in reports:
                if report.get('CompletionDate'):
                    report['CompletionDate'] = report['CompletionDate'].strftime('%Y-%m-%d %H:%M:%S')

        return Response({'reports': reports}, status=status.HTTP_200_OK)

    except Exception as e:
        print(f"ERROR in check_audit_reports: {str(e)}")
        return Response({'error': str(e)}, status=status.HTTP_400_BAD_REQUEST)
 

@api_view(['GET'])
def get_users(request):
    """
    Get all users for allocation
    """
    try:
        users = Users.objects.all()
        serializer = UserSerializer(users, many=True)
        return Response(serializer.data, status=status.HTTP_200_OK)
    except Exception as e:
        return Response({'error': str(e)}, status=status.HTTP_400_BAD_REQUEST)

@api_view(['POST'])
def update_audit_status(request, audit_id):
    """Update the status of an audit"""
    try:
        print(f"DEBUG: Updating status for audit_id={audit_id}")
        print(f"DEBUG: Request data: {request.data}")
        
        # Get the audit record
        audit = Audit.objects.get(AuditId=audit_id)
        old_status = audit.Status
        new_status = request.data.get('status')
        
        print(f"DEBUG: Current status: {old_status}, New status: {new_status}")
        
        # Update the status
        audit.Status = new_status
        
        # Initialize notification service
      
        notification_service = NotificationService()
        
        # If changing to Ready for review
        if new_status == 'Ready for review' and old_status != 'Ready for review':
            try:
                # Need to get the reviewer information from the system - may need to adjust this based on your data model
                # For now, assuming a system administrator or specific reviewer is configured
                admin_users = Users.objects.filter(Role__contains='Admin')
                reviewer_emails = [admin.email for admin in admin_users if hasattr(admin, 'email')]
                
                # Send notification to each administrator/reviewer
                for reviewer_email in reviewer_emails:
                    notification_data = {
                        'notification_type': 'auditCompleted',
                        'email': reviewer_email,
                        'email_type': 'gmail',  # Or 'microsoft' based on your configuration
                        'template_data': [
                            'Administrator',  # Generic placeholder for the reviewer name
                            audit.Title,
                            audit.AssignedTo.get_full_name() if hasattr(audit.AssignedTo, 'get_full_name') else audit.AssignedTo.username,
                            (timezone.now() + timezone.timedelta(days=5)).strftime('%Y-%m-%d')  # Set review due date 5 days from now
                        ]
                    }
                    notification_service.send_multi_channel_notification(notification_data)
                    print(f"DEBUG: Sent 'audit ready for review' notification to {reviewer_email}")
            except Exception as e:
                print(f"ERROR: Failed to send notification: {str(e)}")
                # Don't fail the status update if notification fails
        
        # If changing to Under review, set the ReviewStartDate
        if new_status == 'Under review' and old_status != 'Under review':
            audit.ReviewStartDate = timezone.now()
            print(f"DEBUG: Setting ReviewStartDate to {audit.ReviewStartDate}")
            
            # Send notification to auditor that their audit is under review
            try:
                # Get auditor details from the audit
                auditor = audit.AssignedTo
                if auditor and hasattr(auditor, 'email'):
                    notification_data = {
                        'notification_type': 'auditCompleted',
                        'email': auditor.email,
                        'email_type': 'gmail',  # Or 'microsoft' based on your configuration
                        'template_data': [
                            auditor.get_full_name() if hasattr(auditor, 'get_full_name') else auditor.username,
                            audit.Title,
                            request.user.get_full_name() if hasattr(request.user, 'get_full_name') else request.user.username,
                            timezone.now().strftime('%Y-%m-%d')
                        ]
                    }
                    notification_service.send_multi_channel_notification(notification_data)
                    print(f"DEBUG: Sent 'audit under review' notification to {auditor.email}")
            except Exception as e:
                print(f"ERROR: Failed to send notification: {str(e)}")
                # Don't fail the status update if notification fails
        
        # If changing to Completed, set the CompletionDate
        if new_status == 'Completed' and old_status != 'Completed':
            audit.CompletionDate = timezone.now()
            print(f"DEBUG: Setting CompletionDate to {audit.CompletionDate}")
            
            # Send notification to relevant stakeholders that audit is completed
            try:
                # Get auditor details from the audit
                auditor = audit.AssignedTo
                if auditor and hasattr(auditor, 'email'):
                    notification_data = {
                        'notification_type': 'auditReviewed',
                        'email': auditor.email,
                        'email_type': 'gmail',  # Or 'microsoft' based on your configuration
                        'template_data': [
                            auditor.get_full_name() if hasattr(auditor, 'get_full_name') else auditor.username,
                            audit.Title,
                            'Approved',  # Status is "Completed" which we map to "Approved" in the notification
                            request.user.get_full_name() if hasattr(request.user, 'get_full_name') else request.user.username,
                            request.data.get('comment', 'Audit completed successfully')
                        ]
                    }
                    notification_service.send_multi_channel_notification(notification_data)
                    print(f"DEBUG: Sent 'audit completed' notification to {auditor.email}")
            except Exception as e:
                print(f"ERROR: Failed to send notification: {str(e)}")
                # Don't fail the status update if notification fails
        
        # If changing to Rejected, notify the auditor
        elif new_status == 'Rejected' and old_status != 'Rejected':
            try:
                # Get auditor details from the audit
                auditor = audit.AssignedTo
                if auditor and hasattr(auditor, 'email'):
                    notification_data = {
                        'notification_type': 'auditReviewed',
                        'email': auditor.email,
                        'email_type': 'gmail',  # Or 'microsoft' based on your configuration
                        'template_data': [
                            auditor.get_full_name() if hasattr(auditor, 'get_full_name') else auditor.username,
                            audit.Title,
                            'Rejected',  # Status is explicitly "Rejected"
                            request.user.get_full_name() if hasattr(request.user, 'get_full_name') else request.user.username,
                            request.data.get('comment', 'Please review and address the issues noted.')
                        ]
                    }
                    notification_service.send_multi_channel_notification(notification_data)
                    print(f"DEBUG: Sent 'audit rejected' notification to {auditor.email}")
            except Exception as e:
                print(f"ERROR: Failed to send notification: {str(e)}")
                # Don't fail the status update if notification fails
        
        # For any other status change
        elif old_status != new_status and new_status not in ['Under review', 'Completed', 'Ready for review', 'Rejected']:
            try:
                # Get auditor details from the audit
                auditor = audit.AssignedTo
                if auditor and hasattr(auditor, 'email'):
                    notification_data = {
                        'notification_type': 'policyStatusChange',
                        'email': auditor.email,
                        'email_type': 'gmail',  # Or 'microsoft' based on your configuration
                        'template_data': [
                            auditor.get_full_name() if hasattr(auditor, 'get_full_name') else auditor.username,
                            audit.Title,
                            new_status,
                            request.user.get_full_name() if hasattr(request.user, 'get_full_name') else request.user.username,
                            timezone.now().strftime('%Y-%m-%d %H:%M:%S')
                        ]
                    }
                    notification_service.send_multi_channel_notification(notification_data)
                    print(f"DEBUG: Sent general status change notification to {auditor.email}")
            except Exception as e:
                print(f"ERROR: Failed to send notification: {str(e)}")
                # Don't fail the status update if notification fails
        
        # Save the changes
        audit.save()
        
        # Just update the status without creating a version
        # Version creation is now handled separately when explicitly saving
        
        # If audit is marked as completed, generate and upload report
        if new_status == 'Completed' and old_status != 'Completed':
            try:
                # Import the report utility function
                from .report_utils import generate_and_upload_report
                
                # Get the user ID from the request
                user_id = request.data.get('user_id', 'system')
                
                # Generate and upload the report asynchronously (in a separate thread)
                import threading
                thread = threading.Thread(
                    target=generate_and_upload_report,
                    args=(audit_id, user_id)
                )
                thread.daemon = True  # Make sure the thread doesn't block server shutdown
                thread.start()
                
                print(f"DEBUG: Started report generation thread for audit {audit_id}")
            except Exception as e:
                print(f"ERROR: Failed to start report generation: {str(e)}")
                # Don't fail the status update if report generation fails
        
        return Response({
            "status": "success",
            "message": f"Audit status updated to {new_status}",
            "audit_id": audit_id
        })
    
    except Audit.DoesNotExist:
        print(f"ERROR: Audit with ID {audit_id} not found")
        return Response({"status": "error", "message": "Audit not found"}, status=404)
    except Exception as e:
        print(f"ERROR in update_audit_status: {str(e)}")
        import traceback
        traceback.print_exc()
        return Response({"status": "error", "message": str(e)}, status=500)
 

@api_view(['GET'])
def get_audit_status(request, audit_id):
    """
    Get just the status of a specific audit
    """
    try:
        print(f"DEBUG: get_audit_status called for audit_id: {audit_id}")
        
        try:
            audit = Audit.objects.get(AuditId=audit_id)
        except Audit.DoesNotExist:
            return Response({'error': 'Audit not found'}, status=status.HTTP_404_NOT_FOUND)
        
        # Get all audit findings
        findings = AuditFinding.objects.filter(AuditId=audit_id)
        total_findings = findings.count()
        
        # Count findings with different statuses
        completed_findings = findings.filter(Check='2').count()
        not_applicable_findings = findings.filter(Check='3').count()
        
        # Count items that are considered "done" (completed or not applicable)
        done_findings = completed_findings + not_applicable_findings
        
        # Calculate progress percentage based on total minus Not Applicable
        effective_total = total_findings - not_applicable_findings
        if effective_total > 0:
            completion_percentage = round((completed_findings / effective_total) * 100)
        else:
            # If all findings are Not Applicable, then we're done
            completion_percentage = 100
        
        # Determine the calculated status
        if total_findings == 0:
            calculated_status = 'Yet to Start'
        elif done_findings == total_findings:
            calculated_status = 'Completed'
        else:
            calculated_status = 'Work In Progress'
        
        # Use the database status if available, else fallback to calculated
        final_status = audit.Status or calculated_status
        
        # Get completion date if available
        completion_date = None
        if audit.CompletionDate:
            completion_date = audit.CompletionDate.strftime('%Y-%m-%d %H:%M:%S')
        
        return Response({
            'audit_id': audit_id,
            'status': final_status,
            'saved_status': audit.Status,
            'calculated_status': calculated_status,
            'completion_percentage': completion_percentage,
            'completion_date': completion_date,
            'total_findings': total_findings,
            'completed_findings': completed_findings
        }, status=status.HTTP_200_OK)
    except Exception as e:
        print(f"ERROR in get_audit_status: {str(e)}")
        return Response({'error': str(e)}, status=status.HTTP_400_BAD_REQUEST)



@api_view(['GET'])
def get_all_compliance(request):
    """
    Get all compliance items
    """
    try:
        compliance_items = Compliance.objects.all()
        return Response({
            'count': len(compliance_items),
            'message': f'Found {len(compliance_items)} compliance items'
        }, status=status.HTTP_200_OK)
    except Exception as e:
        return Response({'error': str(e)}, status=status.HTTP_400_BAD_REQUEST)


@api_view(['POST'])
def create_compliance(request):
    """
    Create a new compliance item
    """
    try:
        from datetime import date
        
        # Get subpolicy_id from request
        subpolicy_id = request.data.get('subpolicy_id')
        if not subpolicy_id:
            return Response({'error': 'subpolicy_id is required'}, status=status.HTTP_400_BAD_REQUEST)
        
        # Get description
        description = request.data.get('description')
        if not description:
            return Response({'error': 'description is required'}, status=status.HTTP_400_BAD_REQUEST)
        
        # Get or validate SubPolicy
        try:
            subpolicy = SubPolicy.objects.get(SubPolicyId=subpolicy_id)
        except SubPolicy.DoesNotExist:
            return Response({'error': f'Subpolicy with ID {subpolicy_id} does not exist'}, 
                          status=status.HTTP_404_NOT_FOUND)
        
        # Create compliance item with default values for required fields
        compliance = Compliance.objects.create(
            SubPolicyId=subpolicy,
            ComplianceItemDescription=description,
            IsRisk=request.data.get('is_risk', False),
            PossibleDamage=request.data.get('possible_damage', ''),
            Criticality=request.data.get('criticality', 'Medium'),
            MandatoryOptional=request.data.get('mandatory_optional', 'Mandatory'),
            ManualAutomatic=request.data.get('manual_automatic', 'Manual'),
            Impact=request.data.get('impact', 'Medium'),
            Probability=request.data.get('probability', 'Medium'),
            ActiveInactive=request.data.get('active_inactive', 'Active'),
            PermanentTemporary=request.data.get('permanent_temporary', 'Permanent'),
            CreatedByName=request.data.get('created_by_name', 'System'),
            CreatedByDate=date.today(),
            AuthorizedByName=request.data.get('authorized_by_name', 'System'),
            AuthorizedByDate=date.today(),
            ComplianceVersion='1.0'
        )
        
        return Response({
            'message': 'Compliance item created successfully',
            'compliance_id': compliance.ComplianceId
        }, status=status.HTTP_201_CREATED)
    except Exception as e:
        return Response({'error': str(e)}, status=status.HTTP_400_BAD_REQUEST)


@api_view(['GET'])
def get_compliance_by_subpolicy(request, subpolicy_id):
    """
    Get compliance items for a specific subpolicy
    """
    try:
        compliance_items = Compliance.objects.filter(SubPolicyId=subpolicy_id)
        from .serializers import ComplianceSerializer
        serializer = ComplianceSerializer(compliance_items, many=True)
        return Response({
            'count': len(compliance_items),
            'items': serializer.data
        }, status=status.HTTP_200_OK)
    except Exception as e:
        return Response({'error': str(e)}, status=status.HTTP_400_BAD_REQUEST)
 
@api_view(['GET'])
def get_audit_compliances(request, audit_id):
    """
    Get all compliances organized by policy and subpolicy hierarchy for a specific audit
    """
    try:
        print(f"DEBUG: get_audit_compliances called for audit_id: {audit_id}")
        
        # Check if audit exists
        try:
            audit = Audit.objects.get(AuditId=audit_id)
        except Audit.DoesNotExist:
            return Response({'error': 'Audit not found'}, status=status.HTTP_404_NOT_FOUND)

        # Get all audit findings with compliance details, including policy and subpolicy info
        with connection.cursor() as cursor:
            cursor.execute("""
                SELECT 
                    af.AuditId,
                    p.PolicyId,
                    p.PolicyName,
                    sp.SubPolicyId,
                    sp.SubPolicyName,
                    c.ComplianceId,
                    c.ComplianceItemDescription,
                    c.IsRisk,
                    c.Criticality,
                    c.MandatoryOptional,
                    af.`Check` as status,
                    af.Comments,
                    af.Evidence,
                    af.HowToVerify,
                    af.Impact,
                    af.Recommendation,
                    af.DetailsOfFinding,
                    af.CheckedDate,
                    af.MajorMinor
                FROM 
                    audit_findings af
                JOIN 
                    compliance c ON af.ComplianceId = c.ComplianceId
                JOIN 
                    subpolicies sp ON c.SubPolicyId = sp.SubPolicyId
                JOIN 
                    policies p ON sp.PolicyId = p.PolicyId
                WHERE 
                    af.AuditId = %s
                ORDER BY 
                    p.PolicyId, sp.SubPolicyId, c.ComplianceId
            """, [audit_id])
            
            columns = [col[0] for col in cursor.description]
            findings = [dict(zip(columns, row)) for row in cursor.fetchall()]

        # Organize findings by policy and subpolicy
        organized_data = {}
        for finding in findings:
            policy_id = finding['PolicyId']
            subpolicy_id = finding['SubPolicyId']
            
            # Convert status codes to text and check for special Not Applicable flag
            status_text = 'Not Started'
            compliance_status = 'Not Compliant'  # Default value
            
            # First check for status '3' which is explicitly Not Applicable
            if finding['status'] == '3':
                status_text = 'Not Applicable'
                compliance_status = 'Not Applicable'
                print(f"DEBUG: Compliance ID {finding['ComplianceId']} status code '3', setting compliance_status to 'Not Applicable'")
            # Then check for Not Applicable marker in comments as fallback
            elif finding['Comments'] and '[Not Applicable]' in finding['Comments']:
                status_text = 'Not Applicable'
                compliance_status = 'Not Applicable'
                print(f"DEBUG: Compliance ID {finding['ComplianceId']} has Not Applicable marker in comments, setting compliance_status to 'Not Applicable'")
            elif finding['status'] == '2':
                status_text = 'Completed'
                compliance_status = 'Fully Compliant'
                print(f"DEBUG: Compliance ID {finding['ComplianceId']} status code '2', setting compliance_status to 'Fully Compliant'")
            elif finding['status'] == '1':
                status_text = 'In Progress'
                compliance_status = 'Partially Compliant'
                print(f"DEBUG: Compliance ID {finding['ComplianceId']} status code '1', setting compliance_status to 'Partially Compliant'")
            else:
                status_text = 'Not Started'
                compliance_status = 'Not Compliant'
                print(f"DEBUG: Compliance ID {finding['ComplianceId']} status code '{finding['status']}', setting compliance_status to 'Not Compliant'")
            
            # Format dates if present
            checked_date = None
            if finding.get('CheckedDate'):
                checked_date = finding['CheckedDate'].strftime('%Y-%m-%d %H:%M:%S')

            # Format finding data
            formatted_finding = {
                'compliance_id': finding['ComplianceId'],
                'description': finding['ComplianceItemDescription'],
                'is_risk': finding['IsRisk'],
                'mandatory_optional': finding['MandatoryOptional'],
                'status': finding['status'],
                'status_text': status_text,
                'compliance_status': compliance_status,
                'comments': finding['Comments'] or '',
                'evidence': finding['Evidence'] or '',
                'how_to_verify': finding['HowToVerify'] or '',
                'impact': finding['Impact'] or '',
                'recommendation': finding['Recommendation'] or '',
                'details_of_finding': finding['DetailsOfFinding'] or '',
                'checked_date': checked_date,
                'major_minor': finding.get('MajorMinor')  # Include MajorMinor from audit_findings
            }
            
            # Convert numeric criticality values to text - prioritize MajorMinor from audit_findings
            criticality_text = 'Not Applicable'  # Default
            
            # First check if we have a MajorMinor value from audit_findings
            if finding.get('MajorMinor') is not None:
                major_minor_value = finding['MajorMinor']
                print(f"DEBUG: Using MajorMinor value {major_minor_value} from audit_findings for compliance {finding['ComplianceId']}")
                
                if major_minor_value == '0':
                    criticality_text = 'Minor'
                elif major_minor_value == '1':
                    criticality_text = 'Major'
                elif major_minor_value == '2':
                    criticality_text = 'Not Applicable'
            else:
                # Fall back to criticality from the compliance table
                criticality_value = finding['Criticality']
                
                # Map numeric criticality to text representation
                try:
                    # Convert to int if it's a string representation of a number
                    if isinstance(criticality_value, str) and criticality_value.isdigit():
                        criticality_value = int(criticality_value)
                    
                    if criticality_value == 0:
                        criticality_text = 'Minor'
                    elif criticality_value == 1:
                        criticality_text = 'Major'
                    elif criticality_value == 2:
                        criticality_text = 'Not Applicable'
                    else:
                        # For any other values, use as-is if it's a string, otherwise default
                        criticality_text = criticality_value if isinstance(criticality_value, str) else 'Not Applicable'
                        
                    print(f"DEBUG: Mapped compliance table criticality {criticality_value} to '{criticality_text}' for compliance {finding['ComplianceId']}")
                except Exception as e:
                    print(f"ERROR mapping criticality: {str(e)}")
                    # Use the original value as fallback
                    criticality_text = criticality_value if isinstance(criticality_value, str) else 'Not Applicable'
            
            # Set the mapped criticality in the formatted data
            formatted_finding['criticality'] = criticality_text
            
            # Debug log for criticality values
            print(f"DEBUG: Compliance ID {finding['ComplianceId']} final criticality value: {criticality_text}")
            
            # Create policy entry if it doesn't exist
            if policy_id not in organized_data:
                organized_data[policy_id] = {
                    'policy_id': policy_id,
                    'policy_name': finding['PolicyName'],
                    'subpolicies': {}
                }

            # Create subpolicy entry if it doesn't exist
            if subpolicy_id not in organized_data[policy_id]['subpolicies']:
                organized_data[policy_id]['subpolicies'][subpolicy_id] = {
                    'subpolicy_id': subpolicy_id,
                    'subpolicy_name': finding['SubPolicyName'],
                    'compliances': []
                }

            # Add compliance to subpolicy
            organized_data[policy_id]['subpolicies'][subpolicy_id]['compliances'].append(formatted_finding)

        # Convert the organized data to a list format
        formatted_response = []
        for policy_data in organized_data.values():
            policy = {
                'policy_id': policy_data['policy_id'],
                'policy_name': policy_data['policy_name'],
                'subpolicies': []
            }
            
            for subpolicy_data in policy_data['subpolicies'].values():
                subpolicy = {
                    'subpolicy_id': subpolicy_data['subpolicy_id'],
                    'subpolicy_name': subpolicy_data['subpolicy_name'],
                    'compliances': subpolicy_data['compliances']
                }
                policy['subpolicies'].append(subpolicy)
            
            formatted_response.append(policy)

        return Response({
            'audit_id': audit_id,
            'total_policies': len(formatted_response),
            'policies': formatted_response
        }, status=status.HTTP_200_OK)

    except Exception as e:
        print(f"ERROR in get_audit_compliances: {str(e)}")
        return Response({'error': str(e)}, status=status.HTTP_400_BAD_REQUEST)

 
def create_new_version(audit_id, user_id, data, prefix):
    """Create a new version for audit data with proper error handling for duplicates"""
    try:
        # Ensure audit_id is valid
        audit_id = int(audit_id)

        # Check if audit exists
        with connection.cursor() as cursor:
            cursor.execute("SELECT COUNT(*) FROM audit WHERE AuditId = %s", [audit_id])
            if cursor.fetchone()[0] == 0:
                print(f"ERROR: Audit ID {audit_id} does not exist")
                return None
        
        # Get the next version number
        next_version = get_next_version_number(audit_id, prefix)
        
        # Get table structure to determine the correct columns
        with connection.cursor() as cursor:
            # Try to select one row to check structure (safer than DESCRIBE)
            try:
                cursor.execute("SELECT * FROM audit_version LIMIT 0")
                columns = [col[0] for col in cursor.description]
                print(f"DEBUG: audit_version columns: {columns}")
            except Exception as e:
                print(f"ERROR getting table structure: {str(e)}")
                columns = ['AuditId', 'Version', 'ExtractedInfo', 'UserId', 'Date']
            
            # Now try to insert with transaction to ensure atomicity
            with connection.cursor() as transaction_cursor:
                # Start transaction
                transaction_cursor.execute("START TRANSACTION")
                
                try:
                    # Check for duplicates one more time
                    transaction_cursor.execute(
                        "SELECT COUNT(*) FROM audit_version WHERE AuditId = %s AND Version = %s FOR UPDATE",
                        [audit_id, next_version]
                    )
                    if transaction_cursor.fetchone()[0] > 0:
                        # Version exists, increment version number until we find an unused one
                        version_num = int(next_version[1:])
                        found_unique = False
                        
                        for _ in range(10):  # Try up to 10 times
                            version_num += 1
                            candidate_version = f"{prefix}{version_num}"
                            transaction_cursor.execute(
                                "SELECT COUNT(*) FROM audit_version WHERE AuditId = %s AND Version = %s",
                                [audit_id, candidate_version]
                            )
                            if transaction_cursor.fetchone()[0] == 0:
                                next_version = candidate_version
                                found_unique = True
                                break
                        
                        if not found_unique:
                            # Generate a truly unique version name with timestamp
                            timestamp = timezone.now().strftime('%H%M%S')
                            next_version = f"{prefix}{version_num}_{timestamp}"
                    
                    # Prepare column list based on available columns
                    columns_str = "AuditId, Version, ExtractedInfo, UserId, Date"
                    values_str = "%s, %s, %s, %s, NOW()"
                    params = [audit_id, next_version, json.dumps(data), user_id]
                    
                    # Check for ApprovedRejected value in metadata
                    if isinstance(data, dict) and "__metadata__" in data and "ApprovedRejected" in data["__metadata__"]:
                        approvedRejected = data["__metadata__"]["ApprovedRejected"]
                        columns_str += ", ApprovedRejected"
                        values_str += ", %s"
                        params.append(approvedRejected)
                        print(f"DEBUG: Setting ApprovedRejected to {approvedRejected} in new version")
                    
                    # Now insert the new version
                    query = f"INSERT INTO audit_version ({columns_str}) VALUES ({values_str})"
                    print(f"DEBUG: Inserting version {next_version} for audit {audit_id}")
                    transaction_cursor.execute(query, params)
                    
                    # Commit transaction
                    transaction_cursor.execute("COMMIT")
                    return next_version
                    
                except Exception as e:
                    # Rollback on error
                    transaction_cursor.execute("ROLLBACK")
                    print(f"ERROR in create_new_version transaction: {str(e)}")
                    raise
    
    except Exception as e:
        print(f"ERROR in create_new_version: {str(e)}")
        import traceback
        traceback.print_exc()
        return None
 

def get_initial_audit_data(audit_id, compliance_id=None):
    """
    Create initial audit data structure when no version exists
    """
    try:
        findings_data = {}
        
        # Get all compliance items for this audit
        with connection.cursor() as cursor:
            query = """
                SELECT 
                    af.ComplianceId,
                    c.ComplianceItemDescription,
                    af.`Check`,
                    af.Comments,
                    af.Evidence,
                    af.HowToVerify,
                    af.Impact,
                    af.Recommendation,
                    af.DetailsOfFinding,
                    af.MajorMinor
                FROM 
                    audit_findings af
                JOIN
                    compliance c ON af.ComplianceId = c.ComplianceId
                WHERE 
                    af.AuditId = %s
            """
            params = [audit_id]
            
            if compliance_id:
                query += " AND af.ComplianceId = %s"
                params.append(compliance_id)
                
            cursor.execute(query, params)
            
            for row in cursor.fetchall():
                compliance_id = str(row[0])
                findings_data[compliance_id] = {
                    'description': row[1],
                    'status': row[2] or '0',
                    'comments': row[3] or '',
                    'evidence': row[4] or '',
                    'how_to_verify': row[5] or '',
                    'impact': row[6] or '',
                    'recommendation': row[7] or '',
                    'details_of_finding': row[8] or '',
                    'major_minor': row[9] or '0',
                    'compliance_status': 'Not Compliant',
                    # Add consistent review fields
                    'review_status': 'In Review',
                    'review_comments': '',
                    'reviewer_comments': '',
                    'accept_reject': '0'  # 0=In Review, 1=Accept, 2=Reject
                }
        
        # Add metadata
        findings_data['__metadata__'] = {
            'created_date': timezone.now().strftime('%Y-%m-%d %H:%M:%S'),
            'version_type': 'Auditor'
        }
        
        # Add overall comments
        findings_data['overall_comments'] = ''
        
        return findings_data
    except Exception as e:
        print(f"ERROR in get_initial_audit_data: {str(e)}")
        return None

 

@api_view(['POST'])
def update_audit_finding(request, compliance_id):
    """Update an audit finding for a specific compliance item"""
    print(f"DEBUG: update_audit_finding called for compliance_id: {compliance_id}")
    print(f"DEBUG: Full request data: {request.data}")
    
    try:
        # Get the audit_id from the session - we trust this is correct
        audit_id = request.session['current_audit_id']
        print(f"DEBUG: Using audit_id {audit_id} from session")
        
        if not audit_id:
            # Only if we don't have an audit_id in session, get it from request
            audit_id = request.data.get('audit_id') or request.query_params.get('audit_id')
            print(f"DEBUG: No audit_id in session, using {audit_id} from request")
            
            # If still no audit_id, then we have a problem
            if not audit_id:
                print(f"ERROR: No audit_id found in session or request")
                return Response({"error": "No audit ID found. Please specify an audit ID."}, status=400)
            
            # Store this audit_id in session for future use
            # request.session['current_audit_id'] = audit_id
            print(f"DEBUG: Stored audit_id {audit_id} in session for future use")
        
        # Get the AuditFinding record using the audit_id from session
        try:
            finding = AuditFinding.objects.get(ComplianceId=compliance_id, AuditId=audit_id)
            print(f"DEBUG: Found AuditFinding with AuditId {audit_id}, ComplianceId {compliance_id}")
        except AuditFinding.DoesNotExist:
            print(f"ERROR: No finding found for compliance_id {compliance_id} with audit_id {audit_id}")
            return Response({"error": f"Audit finding not found for compliance {compliance_id} in audit {audit_id}"}, status=404)
        
        # Update the fields from the request data
        if 'how_to_verify' in request.data:
            finding.HowToVerify = request.data['how_to_verify']
        
        if 'impact' in request.data:
            finding.Impact = request.data['impact']
        
        if 'details_of_finding' in request.data:
            finding.DetailsOfFinding = request.data['details_of_finding']
        
        if 'comments' in request.data:
            finding.Comments = request.data['comments']
        
        if 'recommendation' in request.data:
            finding.Recommendation = request.data['recommendation']
        
        if 'criticality' in request.data:
            criticality = request.data['criticality'].strip().lower() if isinstance(request.data['criticality'], str) else ''
            if criticality == 'major':
                finding.MajorMinor = '1'  # Major: 1
            elif criticality == 'minor':
                finding.MajorMinor = '0'  # Minor: 0
            else:
                finding.MajorMinor = '0'  # Default to Minor
        
        # Update compliance status
        if 'compliance_status' in request.data:
            compliance_status = request.data['compliance_status']
            if compliance_status == 'Fully Compliant':
                finding.Check = '2'  # Completed
            elif compliance_status == 'Partially Compliant':
                finding.Check = '1'  # In Progress
            elif compliance_status == 'Not Compliant':
                finding.Check = '0'  # Not Started
            elif compliance_status == 'Not Applicable':
                finding.Check = '3'  # Not Applicable
        
        # Set checked date if not auto_save
        auto_save = request.data.get('auto_save', False)
        if not auto_save:
            finding.CheckedDate = timezone.now()
        
        # Save the changes
        finding.save()
        print(f"DEBUG: Successfully updated finding for compliance_id {compliance_id} with audit_id {audit_id}")
        
        # Return success response
        return Response({
            "status": "success", 
            "message": "Audit finding updated successfully",
            "audit_id": audit_id,
            "compliance_id": compliance_id
        })
    
    except Exception as e:
        print(f"ERROR in update_audit_finding: {str(e)}")
        import traceback
        traceback.print_exc()
        return Response({"error": str(e)}, status=500)



# @api_view(['POST'])
# def upload_evidence(request, compliance_id):
#     """
#     Upload evidence for a specific audit finding
#     - Supports auto-save functionality with tracking
#     """
#     try:
#         is_auto_save = request.POST.get('auto_save', 'true').lower() == 'true'
#         print(f"DEBUG: upload_evidence called for compliance_id: {compliance_id} (auto-save: {is_auto_save})")
        
#         # Check if there's a file in the request - check both 'evidence' and 'file' names
#         file = None
#         if 'evidence' in request.FILES:
#             file = request.FILES['evidence']
#             print("DEBUG: File found in 'evidence' field")
#         elif 'file' in request.FILES:
#             file = request.FILES['file'] 
#             print("DEBUG: File found in 'file' field")
#         else:
#             return Response({'error': 'No evidence file provided'}, status=status.HTTP_400_BAD_REQUEST)
            
#         # Get the file information
#         file_name = file.name
#         file_size = file.size
        
#         print(f"DEBUG: Processing file upload: {file_name} ({file_size} bytes)")
        
#         # Use raw SQL to update the audit finding
#         with connection.cursor() as cursor:
#             # First check if the audit finding exists
#             cursor.execute("""
#                 SELECT 
#                     af.`Check`
#                 FROM 
#                     audit_findings af
#                 JOIN 
#                     compliance c ON af.ComplianceId = c.ComplianceId 
#                 WHERE 
#                     c.ComplianceId = %s
#                 LIMIT 1
#             """, [compliance_id])
            
#             result = cursor.fetchone()
#             if not result:
#                 return Response({'error': 'Audit finding not found'}, status=status.HTTP_404_NOT_FOUND)
            
#             current_check = result[0]
            
#             # Prepare the update
#             update_fields = ["Evidence = %s"]
#             update_values = [file_name]
            
#             # If status was not set yet, mark it as "In Progress"
#             if current_check == '0':
#                 update_fields.append("`Check` = %s")
#                 update_values.append('1')  # Mark as In Progress
#                 print(f"DEBUG: Will update status to 'In Progress' as evidence was uploaded")
            
#             # Execute the update
#             update_sql = f"""
#                 UPDATE audit_findings af
#                 JOIN compliance c ON af.ComplianceId = c.ComplianceId
#                 SET {', '.join(update_fields)}
#                 WHERE c.ComplianceId = %s
#             """
#             update_values.append(compliance_id)
            
#             cursor.execute(update_sql, update_values)
#             print(f"DEBUG: Updated {cursor.rowcount} audit finding record(s) with evidence")
        
#         print(f"DEBUG: Evidence '{file_name}' uploaded for compliance {compliance_id} via {'auto-save' if is_auto_save else 'manual save'}")
        
#         # Generate an S3 URL for the uploaded file
#         s3_url = f"https://grc-files-vardaan.s3.amazonaws.com/evidence/{compliance_id}_{file_name}"
        
#         return Response({
#             'message': f"Evidence {'auto-saved' if is_auto_save else 'uploaded'} successfully",
#             'compliance_id': compliance_id,
#             'filename': file_name,
#             'file_size': file_size,
#             's3_url': s3_url,  # Include S3 URL in the response
#             'url': s3_url,     # Also include as url for backwards compatibility
#             'timestamp': timezone.now().strftime('%Y-%m-%d %H:%M:%S')
#         }, status=status.HTTP_200_OK)
        
#     except Exception as e:
#         print(f"ERROR in upload_evidence: {str(e)}")
#         return Response({'error': str(e)}, status=status.HTTP_400_BAD_REQUEST)

 
 
 
def get_audit_findings_json(audit_id, overall_comments=None):
    """
    Get a JSON representation of all findings for an audit
    """
    try:
        print(f"DEBUG: get_audit_findings_json called for audit_id: {audit_id}")
        findings_data = {}
        
        # Get all compliance items for this audit
        with connection.cursor() as cursor:
            query = """
                SELECT 
                    af.ComplianceId,
                    c.ComplianceItemDescription,
                    af.`Check`,
                    af.Comments,
                    af.Evidence,
                    af.HowToVerify,
                    af.Impact,
                    af.Recommendation,
                    af.DetailsOfFinding,
                    af.MajorMinor
                FROM 
                    audit_findings af
                JOIN
                    compliance c ON af.ComplianceId = c.ComplianceId
                WHERE 
                    af.AuditId = %s
            """
            params = [audit_id]
            
            cursor.execute(query, params)
            rows = cursor.fetchall()
            print(f"DEBUG: Found {len(rows)} findings for audit {audit_id}")
            
            for row in rows:  # CRITICAL FIX: changed from cursor.fetchall() which would skip rows
                compliance_id = str(row[0])
                check_value = row[2] or '0'
                major_minor = row[9] or '0'
                
                # Map check value to compliance status
                if check_value == '0':
                    compliance_status = 'Not Compliant'
                elif check_value == '1':
                    compliance_status = 'Partially Compliant'
                elif check_value == '2':
                    compliance_status = 'Fully Compliant'
                elif check_value == '3':
                    compliance_status = 'Not Applicable'
                else:
                    compliance_status = 'Not Compliant'
                
                # Map major/minor value to criticality
                if major_minor == '0':
                    criticality = 'Minor'
                elif major_minor == '1':
                    criticality = 'Major'
                elif major_minor == '2':
                    criticality = 'Not Applicable'
                else:
                    criticality = 'Minor'
                
                findings_data[compliance_id] = {
                    'description': row[1],
                    'status': check_value,
                    'compliance_status': compliance_status,
                    'comments': row[3] or '',
                    'evidence': row[4] or '',
                    'how_to_verify': row[5] or '',
                    'impact': row[6] or '',
                    'recommendation': row[7] or '',
                    'details_of_finding': row[8] or '',
                    'major_minor': major_minor,
                    'criticality': criticality,
                    # Add consistent review fields to A versions
                    'review_status': 'In Review',
                    'review_comments': '',
                    'reviewer_comments': '',
                    'accept_reject': '0'  # 0=In Review, 1=Accept, 2=Reject
                }
                print(f"DEBUG: Processed compliance ID {compliance_id}: {compliance_status}, {criticality}")
        
        # Add metadata
        findings_data['__metadata__'] = {
            'created_date': timezone.now().strftime('%Y-%m-%d %H:%M:%S'),
            'version_type': 'Auditor'
        }
        
        # Add overall comments
        findings_data['overall_comments'] = overall_comments or ''
        
        print(f"DEBUG: get_audit_findings_json created JSON with {len(findings_data)-2} compliance items")
        return findings_data
    except Exception as e:
        print(f"ERROR in get_audit_findings_json: {str(e)}")
        import traceback
        traceback.print_exc()
        return None


@api_view(['GET'])
def add_majorminor_column(request):
    """
    Add MajorMinor column to audit_findings table
    """
    try:
        with connection.cursor() as cursor:
            # Check if the column already exists
            cursor.execute("""
                SELECT COUNT(*) 
                FROM information_schema.COLUMNS 
                WHERE TABLE_NAME = 'audit_findings' 
                AND COLUMN_NAME = 'MajorMinor'
            """)
            
            column_exists = cursor.fetchone()[0] > 0
            
            if column_exists:
                return Response({
                    'message': 'MajorMinor column already exists in audit_findings table.'
                }, status=status.HTTP_200_OK)
            else:
                # Add the column
                cursor.execute("""
                    ALTER TABLE audit_findings
                    ADD COLUMN MajorMinor CHAR(1) NULL;
                """)
                return Response({
                    'message': 'MajorMinor column added successfully to audit_findings table.'
                }, status=status.HTTP_200_OK)
    except Exception as e:
        print(f"ERROR adding MajorMinor column: {str(e)}")
        return Response({'error': str(e)}, status=status.HTTP_400_BAD_REQUEST)

# @api_view(['POST'])
# def add_compliance_to_audit(request, audit_id):
#     """
#     Add a new compliance item in the context of an audit
#     - Creates a new compliance item
#     - Creates a corresponding audit_finding
#     - Uses audit context to determine framework/policy/subpolicy constraints
#     """
#     try:
#         print(f"DEBUG: add_compliance_to_audit called for audit_id: {audit_id}")
#         print(f"DEBUG: Request data: {request.data}")
        
#         # Extract data from the request
#         compliance_data = request.data.copy()
        
#         # First, verify the audit exists and get necessary data using raw SQL 
#         # to avoid ORM issues with model fields not matching database schema
#         with connection.cursor() as cursor:
#             cursor.execute("""
#                 SELECT 
#                     a.AuditId, a.SubPolicyId, a.auditor, a.FrameworkId, a.AssignedDate
#                 FROM 
#                     audit a
#                 WHERE 
#                     a.AuditId = %s
#             """, [audit_id])
            
#             audit_row = cursor.fetchone()
#             if not audit_row:
#                 return Response({'error': f'Audit with ID {audit_id} not found'}, status=status.HTTP_404_NOT_FOUND)
            
#             # Extract data from the result
#             audit_id, audit_subpolicy_id, auditor_id, framework_id, audit_assigned_date = audit_row
#             print(f"DEBUG: Found audit with ID {audit_id}, AssignedDate: {audit_assigned_date}")
            
#             # If audit has no AssignedDate, get it from existing audit findings
#             if not audit_assigned_date:
#                 cursor.execute("""
#                     SELECT AssignedDate 
#                     FROM audit_findings 
#                     WHERE AuditId = %s 
#                     ORDER BY AssignedDate 
#                     LIMIT 1
#                 """, [audit_id])
                
#                 existing_date_row = cursor.fetchone()
#                 if existing_date_row:
#                     audit_assigned_date = existing_date_row[0]
#                     print(f"DEBUG: Using AssignedDate from existing audit finding: {audit_assigned_date}")
#                 else:
#                     # If still no date, use current time
#                     audit_assigned_date = timezone.now()
#                     print(f"DEBUG: No existing AssignedDate found, using current time: {audit_assigned_date}")
        
#         # Handle case where subpolicy is provided directly
#         if 'subpolicy_id' in compliance_data and compliance_data['subpolicy_id']:
#             subpolicy_id = compliance_data['subpolicy_id']
#         # Otherwise, use the subpolicy from the audit
#         elif audit_subpolicy_id:
#             subpolicy_id = audit_subpolicy_id
#             print(f"DEBUG: Using subpolicy {subpolicy_id} from audit")
#         else:
#             return Response({'error': 'No subpolicy specified and audit does not have a subpolicy'}, 
#                            status=status.HTTP_400_BAD_REQUEST)
        
#         # Verify subpolicy exists using raw SQL
#         with connection.cursor() as cursor:
#             cursor.execute("""
#                 SELECT 
#                     sp.SubPolicyId, sp.SubPolicyName 
#                 FROM 
#                     subpolicies sp
#                 WHERE 
#                     sp.SubPolicyId = %s
#             """, [subpolicy_id])
            
#             subpolicy_row = cursor.fetchone()
#             if not subpolicy_row:
#                 return Response({'error': f'Subpolicy with ID {subpolicy_id} does not exist'}, 
#                               status=status.HTTP_404_NOT_FOUND)
            
#             subpolicy_id, subpolicy_name = subpolicy_row
#             print(f"DEBUG: Found subpolicy: ID={subpolicy_id}, Name={subpolicy_name}")
        
#         # Get current user info
#         user_id = request.session.get('user_id', auditor_id)
        
#         # Get user data
#         with connection.cursor() as cursor:
#             cursor.execute("""
#                 SELECT 
#                     u.UserName
#                 FROM 
#                     users u
#                 WHERE 
#                     u.UserId = %s
#             """, [user_id])
            
#             user_row = cursor.fetchone()
#             if not user_row:
#                 user_name = "System"
#             else:
#                 user_name = user_row[0]
        
#         # Get current date
#         current_date = timezone.now().date()
        
#         # Insert compliance directly using raw SQL
#         with connection.cursor() as cursor:
#             cursor.execute("""
#                 INSERT INTO compliance (
#                     SubPolicyId, ComplianceItemDescription, IsRisk, PossibleDamage,
#                     Criticality, MandatoryOptional, ManualAutomatic, Impact,
#                     Probability, ActiveInactive, PermanentTemporary,
#                     CreatedByName, CreatedByDate, AuthorizedByName, AuthorizedByDate,
#                     ComplianceVersion, Status
#                 ) VALUES (
#                     %s, %s, %s, %s,
#                     %s, %s, %s, %s,
#                     %s, %s, %s,
#                     %s, %s, %s, %s,
#                     %s, %s
#                 )
#             """, [
#                 subpolicy_id, 
#                 compliance_data.get('description', ''),
#                 1 if compliance_data.get('is_risk', False) else 0,
#                 compliance_data.get('possible_damage', ''),
#                 compliance_data.get('criticality', 'Medium'),
#                 compliance_data.get('mandatory_optional', 'Mandatory'),
#                 compliance_data.get('manual_automatic', 'Manual'),
#                 compliance_data.get('impact', 'Medium'),
#                 compliance_data.get('probability', 'Medium'),
#                 compliance_data.get('active_inactive', 'Active'),
#                 compliance_data.get('permanent_temporary', 'Temporary'),
#                 user_name,
#                 current_date,
#                 user_name,
#                 current_date,
#                 compliance_data.get('ComplianceVersion', '0'),
#                 'Active'
#             ])
            
#             # Get the newly inserted ID
#             cursor.execute("SELECT LAST_INSERT_ID()")
#             compliance_id = cursor.fetchone()[0]
#             print(f"DEBUG: Created new compliance item with ID {compliance_id}")
        
#         # Create audit finding with raw SQL - use the exact same AssignedDate from the audit
#         with connection.cursor() as cursor:
#             cursor.execute("""
#                 INSERT INTO audit_findings (
#                     AuditId, ComplianceId, UserId, Evidence, 
#                     `Check`, Comments, AssignedDate
#                 ) VALUES (
#                     %s, %s, %s, %s,
#                     %s, %s, %s
#                 )
#             """, [
#                 audit_id,
#                 compliance_id,
#                 user_id,
#                 '',  # Evidence
#                 '0',  # Check = Yet to Start
#                 '',  # Comments
#                 audit_assigned_date  # Use the AssignedDate from the audit
#             ])
            
#             print(f"DEBUG: Created audit finding for compliance ID: {compliance_id} with AssignedDate: {audit_assigned_date}")
        
#         # Return success response
#         return Response({
#             'message': 'Compliance item added successfully',
#             'compliance_id': compliance_id,
#             'audit_finding_created': True
#         }, status=status.HTTP_201_CREATED)
        
#     except Exception as e:
#         print(f"ERROR in add_compliance_to_audit: {str(e)}")
#         return Response({'error': str(e)}, status=status.HTTP_400_BAD_REQUEST)
 
@api_view(['GET'])
def fix_subpolicy_version_field(request):
    """
    Check and fix the Version field in the subpolicies table
    """
    try:
        with connection.cursor() as cursor:
            # Check if the Version column exists
            cursor.execute("""
                SELECT COUNT(*) 
                FROM information_schema.COLUMNS 
                WHERE TABLE_NAME = 'subpolicies' 
                AND COLUMN_NAME = 'Version'
            """)
            
            version_exists = cursor.fetchone()[0] > 0
            
            if version_exists:
                # Check if there are records with NULL Version values
                cursor.execute("""
                    SELECT COUNT(*) 
                    FROM subpolicies 
                    WHERE Version IS NULL OR Version = ''
                """)
                
                null_versions = cursor.fetchone()[0]
                
                if null_versions > 0:
                    # Update the NULL versions to a default value
                    cursor.execute("""
                        UPDATE subpolicies
                        SET Version = '1.0'
                        WHERE Version IS NULL OR Version = ''
                    """)
                    
                    return Response({
                        'message': f'Updated {null_versions} subpolicies with default Version value.'
                    }, status=status.HTTP_200_OK)
                else:
                    return Response({
                        'message': 'All subpolicies have Version values set.'
                    }, status=status.HTTP_200_OK)
            else:
                # Add the Version column
                cursor.execute("""
                    ALTER TABLE subpolicies
                    ADD COLUMN Version VARCHAR(50) NOT NULL DEFAULT '1.0'
                """)
                
                return Response({
                    'message': 'Version column added to subpolicies table with default value 1.0.'
                }, status=status.HTTP_200_OK)
    except Exception as e:
        print(f"ERROR fixing subpolicy version field: {str(e)}")
        return Response({'error': str(e)}, status=status.HTTP_400_BAD_REQUEST)

  

@api_view(['GET'])
def get_my_reviews(request):
    """
    Fetch audits assigned to the current user (as reviewer)
    If user_id is not in session, default to user 1020
    """
    try:
        print("DEBUG: get_my_reviews was called")

        # Get user_id from session, default to 1020 if not found
        user_id = request.session.get('user_id', 1020)
        print(f"DEBUG: Using user_id: {user_id}")
        
        # Check if ReviewDate and ReviewComments columns exist in the audit table
        review_date_exists = False
        review_comments_exists = False
        try:
            with connection.cursor() as cursor:
                # Check for ReviewDate column
                cursor.execute("""
                    SELECT COUNT(*) 
                    FROM information_schema.COLUMNS 
                    WHERE TABLE_NAME = 'audit' 
                    AND COLUMN_NAME = 'ReviewDate'
                """)
                review_date_exists = cursor.fetchone()[0] > 0
                print(f"DEBUG: ReviewDate column exists: {review_date_exists}")
                
                # Check for ReviewComments column
                cursor.execute("""
                    SELECT COUNT(*) 
                    FROM information_schema.COLUMNS 
                    WHERE TABLE_NAME = 'audit' 
                    AND COLUMN_NAME = 'ReviewComments'
                """)
                review_comments_exists = cursor.fetchone()[0] > 0
                print(f"DEBUG: ReviewComments column exists: {review_comments_exists}")
        except Exception as e:
            print(f"DEBUG: Error checking for columns: {str(e)}")
            review_date_exists = False
            review_comments_exists = False
        
        # Using raw SQL to join multiple tables and get comprehensive data
        with connection.cursor() as cursor:
            print(f"DEBUG: Executing SQL query for my reviews")
            
            # Build SQL query based on available columns
            select_fields = [
                "a.AuditId as audit_id",
                "f.FrameworkName as framework",
                "p.PolicyName as policy",
                "sp.SubPolicyName as subpolicy",
                "a.DueDate as duedate",
                "a.Frequency as frequency",
                "auditor_user.UserName as auditor",
                "auditor_user.UserId as auditor_id",
                "a.AuditType as audit_type",
                "assignee_user.UserName as assignee",
                "a.Status as status",
                "COUNT(af.AuditId) as total_compliances",
                "SUM(CASE WHEN af.`Check` = '2' THEN 1 ELSE 0 END) as completed_compliances",
                "a.CompletionDate as completion_date",
                "(SELECT ApprovedRejected FROM audit_version WHERE AuditId = a.AuditId ORDER BY Version DESC LIMIT 1) as approved_rejected"
            ]
            
            # Only include ReviewStatus if it exists
            try:
                cursor.execute("""
                    SELECT COUNT(*) 
                    FROM information_schema.COLUMNS 
                    WHERE TABLE_NAME = 'audit' 
                    AND COLUMN_NAME = 'ReviewStatus'
                """)
                review_status_exists = cursor.fetchone()[0] > 0
                if review_status_exists:
                    select_fields.append("a.ReviewStatus as review_status")
                    print(f"DEBUG: ReviewStatus column exists, adding to query")
                else:
                    print(f"DEBUG: ReviewStatus column does not exist, skipping")
            except Exception as e:
                print(f"DEBUG: Error checking for ReviewStatus column: {str(e)}")
                review_status_exists = False
            
            # Add ReviewComments field only if it exists
            if review_comments_exists:
                select_fields.append("a.ReviewComments as review_comments")
                print(f"DEBUG: Adding ReviewComments to query")
            
            # Add ReviewDate field only if it exists
            if review_date_exists:
                select_fields.append("a.ReviewDate as review_date")
                print(f"DEBUG: Adding ReviewDate to query")
                
            # Join fields and build the complete query
            select_clause = ", ".join(select_fields)
            
            # Group by fields without ReviewDate and ReviewComments
            group_by_fields = [
                "a.AuditId", "f.FrameworkName", "p.PolicyName", "sp.SubPolicyName", 
                "a.DueDate", "a.Frequency", "auditor_user.UserName", "auditor_user.UserId", "a.AuditType",
                "assignee_user.UserName", "a.Status", "a.CompletionDate"
            ]
            
            # Add ReviewStatus to GROUP BY if it exists
            if review_status_exists:
                group_by_fields.append("a.ReviewStatus")
            
            # Add ReviewComments to GROUP BY if it exists
            if review_comments_exists:
                group_by_fields.append("a.ReviewComments")
                
            # Add ReviewDate to GROUP BY if it exists
            if review_date_exists:
                group_by_fields.append("a.ReviewDate")
                
            group_by_clause = ", ".join(group_by_fields)
            
            query = f"""
                SELECT 
                    {select_clause}
                FROM 
                    audit a
                LEFT JOIN 
                    frameworks f ON a.FrameworkId = f.FrameworkId
                LEFT JOIN 
                    policies p ON a.PolicyId = p.PolicyId
                LEFT JOIN 
                    subpolicies sp ON a.SubPolicyId = sp.SubPolicyId
                LEFT JOIN 
                    users assignee_user ON a.assignee = assignee_user.UserId
                LEFT JOIN 
                    users auditor_user ON a.auditor = auditor_user.UserId
                LEFT JOIN 
                    audit_findings af ON a.AuditId = af.AuditId
                WHERE 
                    a.reviewer = %s
                GROUP BY 
                    {group_by_clause}
                ORDER BY 
                    a.DueDate ASC
            """
            
            print(f"DEBUG: Executing SQL query: {query}")
            cursor.execute(query, [user_id])
            
            columns = [col[0] for col in cursor.description]
            print(f"DEBUG: My reviews query columns: {columns}")
            audits = [dict(zip(columns, row)) for row in cursor.fetchall()]
            print(f"DEBUG: Fetched {len(audits)} audits for review")

        # Process and format audit data for display
        for audit in audits:
            # Format date
            if audit.get('duedate'):
                audit['duedate'] = audit['duedate'].strftime('%d/%m/%Y')
            
            # Format completion date if present
            if audit.get('completion_date'):
                audit['completion_date'] = audit['completion_date'].strftime('%d/%m/%Y %H:%M')
                
            # Format review date if present
            if audit.get('review_date'):
                audit['review_date'] = audit['review_date'].strftime('%d/%m/%Y %H:%M')
            else:
                # If ReviewDate doesn't exist in DB, provide a default value
                audit['review_date'] = None
            
            # Set default review comments if not present
            if 'review_comments' not in audit:
                audit['review_comments'] = None
                
            # Set default review status if not present or null
            if 'review_status' not in audit or not audit.get('review_status'):
                audit['review_status'] = "Yet to Start"
            else:
                # Convert integer review status to string representation
                review_status_int = audit.get('review_status')
                if isinstance(review_status_int, int) or (isinstance(review_status_int, str) and review_status_int.isdigit()):
                    # Map review status integer to string
                    status_map = {
                        0: 'Yet to Start',
                        1: 'In Review',
                        2: 'Accept',
                        3: 'Reject'
                    }
                    review_status_int = int(review_status_int)
                    audit['review_status'] = status_map.get(review_status_int, 'Unknown')
                    print(f"DEBUG: Mapped ReviewStatus {review_status_int} to '{audit['review_status']}' for audit {audit['audit_id']}")
                
            # Map audit status to review status for display purposes
            audit_status = audit.get('status', '')
            
            # Determine if review can be updated and the display status
            can_update_review = False
            display_review_status = audit['review_status']
            
            # Check for ApprovedRejected value from the latest version
            approved_rejected = audit.get('approved_rejected')
            if approved_rejected:
                print(f"DEBUG: Audit {audit['audit_id']} has ApprovedRejected value: {approved_rejected}")
                
                # Automatically update the review status based on ApprovedRejected
                if approved_rejected == 'Approved':
                    # Set review status to Accept if ApprovedRejected is Approved
                    audit['review_status'] = 'Accept'
                    display_review_status = 'Accept'
                    print(f"DEBUG: Setting review status to 'Accept' based on ApprovedRejected")
                elif approved_rejected == 'Rejected':
                    # Set review status to Reject if ApprovedRejected is Rejected
                    audit['review_status'] = 'Reject'
                    display_review_status = 'Reject'
                    print(f"DEBUG: Setting review status to 'Reject' based on ApprovedRejected")
            
            if audit_status in ['Yet to Start', 'Work In Progress']:
                # For audit in progress, show "Under Audit" for the reviewer
                display_review_status = "Under Audit"
                can_update_review = False
                print(f"DEBUG: Audit {audit['audit_id']} status '{audit_status}' - display review status 'Under Audit'")
            elif audit_status == 'Under review':
                # When audit is submitted for review, reviewer can perform review
                # Default to "Yet to Start" for the reviewer's initial state
                if display_review_status == "Yet to Start":
                    print(f"DEBUG: Audit {audit['audit_id']} is ready for review - review status will start with 'Yet to Start'")
                can_update_review = True
                print(f"DEBUG: Audit {audit['audit_id']} status 'Under review' - reviewer can update status")
            elif audit_status == 'Completed':
                # If audit is already completed, don't allow further review updates
                can_update_review = False
                print(f"DEBUG: Audit {audit['audit_id']} is already completed - review status locked")
                
            # Add the mapped values to the audit record
            audit['display_review_status'] = display_review_status
            audit['can_update_review'] = can_update_review
            audit['approved_rejected'] = approved_rejected  # Add ApprovedRejected to the response
            
            # Calculate completion percentage
            total = audit.get('total_compliances') or 0
            completed = audit.get('completed_compliances') or 0
            audit['completion_percentage'] = round((completed / total) * 100) if total > 0 else 0
            
            # Convert frequency number to text
            freq = audit.get('frequency')
            if freq is not None:
                if freq == 0:
                    audit['frequency_text'] = 'Only Once'
                elif freq == 1:
                    audit['frequency_text'] = 'Daily'
                elif freq <= 60:
                    audit['frequency_text'] = 'Every 2 Months'
                elif freq <= 120:
                    audit['frequency_text'] = 'Every 4 Months'
                elif freq <= 182:
                    audit['frequency_text'] = 'Half Yearly'
                elif freq <= 365:
                    audit['frequency_text'] = 'Yearly'
                else:
                    audit['frequency_text'] = f'Every {freq} days'
            
            # Convert audit type from I/E to Internal/External
            if audit.get('audit_type') == 'I':
                audit['audit_type_text'] = 'Internal'
            elif audit.get('audit_type') == 'E':
                audit['audit_type_text'] = 'External'

        print(f"DEBUG: Successfully prepared my reviews response")
        return Response({
            'user_id': user_id,
            'audits': audits
        }, status=status.HTTP_200_OK)
    except Exception as e:
        print(f"ERROR in get_my_reviews: {str(e)}")
        return Response({'error': str(e)}, status=status.HTTP_400_BAD_REQUEST)
    

@api_view(['POST'])
def upload_evidence(request, compliance_id):
    """
    Upload evidence for a specific audit finding
    - Supports auto-save functionality with tracking
    - Integrates with S3 storage for file uploads
    """
    try:
        is_auto_save = request.POST.get('auto_save', 'false').lower() == 'true'
        print(f"DEBUG: upload_evidence called for compliance_id: {compliance_id} (auto-save: {is_auto_save})")
        
        # Check if there's a file in the request - check both 'evidence' and 'file' names
        file = None
        if 'evidence' in request.FILES:
            file = request.FILES['evidence']
            print("DEBUG: File found in 'evidence' field")
        elif 'file' in request.FILES:
            file = request.FILES['file'] 
            print("DEBUG: File found in 'file' field")
        else:
            return Response({'error': 'No evidence file provided'}, status=status.HTTP_400_BAD_REQUEST)
            
        # Get file information
        file_name = file.name
        file_size = file.size
        
        print(f"DEBUG: Processing file upload: {file_name} ({file_size} bytes)")
        
        # Get user information
        user_id = request.user.id if request.user.is_authenticated else 'anonymous'
        
        # Get audit_id from request if available
        audit_id = request.POST.get('audit_id', None)
        
        # Get table name and storage column from request or use defaults
        table_name = request.POST.get('table_name', 'audit_findings')
        storage_column = request.POST.get('storage_column', 'Evidence')
        
        print(f"DEBUG: Upload parameters - audit_id: {audit_id}, table_name: {table_name}, storage_column: {storage_column}")
        
        # Create a temporary file to save the uploaded file
        import tempfile
        import os
        from django.utils import timezone
        
        temp_file_path = os.path.join(tempfile.gettempdir(), file_name)
        with open(temp_file_path, 'wb+') as destination:
            for chunk in file.chunks():
                destination.write(chunk)
        
        # Upload to S3 if the S3 service is configured
        s3_location = ""
        try:
            from .s3_functions import S3Client
            s3_client = S3Client()
            
            # Upload to S3 with additional parameters
            upload_params = {
                'compliance_id': str(compliance_id),
                'documentType': 'evidence',
                'table_name': table_name,
                'storage_column': storage_column
            }
            
            # Add audit_id if available
            if audit_id:
                upload_params['audit_id'] = str(audit_id)
            
            # Upload to S3
            upload_result = s3_client.upload_file(
                file_path=temp_file_path,
                user_id=str(user_id),
                file_name=file_name,
                **upload_params
            )
            
            s3_location = upload_result.get('file_id', '')
            print(f"DEBUG: File uploaded to S3, location: {s3_location}")
        except Exception as s3_error:
            print(f"ERROR in S3 upload: {str(s3_error)}")
            # If S3 upload fails, create a fallback URL with timestamp and filename
            timestamp = timezone.now().strftime('%Y%m%d%H%M%S')
            s3_location = f"local-storage://{timestamp}_{file_name}"
            print(f"DEBUG: Using fallback location for evidence: {s3_location}")
        
        # Clean up the temporary file
        try:
            os.unlink(temp_file_path)
        except Exception as e:
            print(f"WARNING: Failed to delete temporary file: {str(e)}")
        
        # Use raw SQL to update the audit finding
        with connection.cursor() as cursor:
            # First check if the audit finding exists
            cursor.execute("""
                SELECT 
                    af.`Check`
                FROM 
                    audit_findings af
                JOIN 
                    compliance c ON af.ComplianceId = c.ComplianceId 
                WHERE 
                    c.ComplianceId = %s
                LIMIT 1
            """, [compliance_id])
            
            result = cursor.fetchone()
            if not result:
                return Response({'error': 'Audit finding not found'}, status=status.HTTP_404_NOT_FOUND)
            
            current_check = result[0]
            
            # First get the existing evidence URLs
            cursor.execute("""
                SELECT Evidence
                FROM audit_findings
                WHERE ComplianceId = %s
            """, [compliance_id])
            
            evidence_row = cursor.fetchone()
            existing_evidence = evidence_row[0] if evidence_row and evidence_row[0] else ''
            
            # Append the new URL to existing ones, separated by commas
            if existing_evidence and s3_location:
                # Check if the URL is already in the list to avoid duplicates
                existing_urls = existing_evidence.split(',')
                if s3_location not in existing_urls:
                    combined_evidence = f"{existing_evidence},{s3_location}"
                else:
                    combined_evidence = existing_evidence
            else:
                combined_evidence = s3_location
            
            # Prepare the update for audit finding
            update_fields = ["Evidence = %s"]
            update_values = [combined_evidence]  # Store the combined URLs
            
            # If status was not set yet, mark it as "In Progress"
            if current_check == '0':
                update_fields.append("`Check` = %s")
                update_values.append('1')  # Mark as In Progress
                print(f"DEBUG: Will update status to 'In Progress' as evidence was uploaded")
            
            # Execute the update
            update_sql = f"""
                UPDATE audit_findings af
                JOIN compliance c ON af.ComplianceId = c.ComplianceId
                SET {', '.join(update_fields)}
                WHERE c.ComplianceId = %s
            """
            update_values.append(compliance_id)
            
            cursor.execute(update_sql, update_values)
            print(f"DEBUG: Updated {cursor.rowcount} audit finding record(s) with S3 URL in Evidence column")
        
        print(f"DEBUG: Evidence '{file_name}' uploaded for compliance {compliance_id} via {'auto-save' if is_auto_save else 'manual save'}")
        
        return Response({
            'message': f"Evidence {'auto-saved' if is_auto_save else 'uploaded'} successfully. S3 URL stored in Evidence column.",
            'compliance_id': compliance_id,
            'audit_id': audit_id,
            'filename': file_name,
            'file_size': file_size,
            's3_location': s3_location,
            'all_evidence': combined_evidence,
            'timestamp': timezone.now().strftime('%Y-%m-%d %H:%M:%S')
        }, status=status.HTTP_200_OK)
        
    except Exception as e:
        print(f"ERROR in upload_evidence: {str(e)}")
        import traceback
        traceback.print_exc()
        return Response({'error': str(e)}, status=status.HTTP_400_BAD_REQUEST)

 
@api_view(['POST'])
def submit_audit_findings(request, audit_id):
    """
    Mark an audit as ready for review and submit all findings.
    This is explicitly called when the auditor clicks "Submit for Review" button,
    not automatically when changing status.
    """
    try:
        print(f"DEBUG: submit_audit_findings called for audit_id: {audit_id}")
        print(f"DEBUG: Request data: {request.data}")
        
        # Find the audit
        try:
            audit = Audit.objects.get(AuditId=audit_id)
        except Audit.DoesNotExist:
            return Response({'error': 'Audit not found'}, status=status.HTTP_404_NOT_FOUND)
        
        # Get user ID from session or request
        user_id = request.session.get('user_id')
        if not user_id:
            # If no user_id in session, attempt to get auditor ID from audit
            try:
                # Try getting it from auditor field
                user_id = audit.auditor_id
                print(f"DEBUG: Using auditor_id from audit: {user_id}")
            except Exception as e:
                # If that fails, try to get it from the audit_findings
                print(f"DEBUG: Error getting auditor_id from audit object: {str(e)}")
                try:
                    # Get the first audit finding for this audit to get the UserId
                    with connection.cursor() as cursor:
                        cursor.execute("""
                            SELECT UserId FROM audit_findings
                            WHERE AuditId = %s
                            LIMIT 1
                        """, [audit_id])
                        row = cursor.fetchone()
                        if row:
                            user_id = row[0]
                            print(f"DEBUG: Got user_id {user_id} from audit_findings")
                except Exception as e:
                    print(f"DEBUG: Error getting user_id from audit_findings: {str(e)}")
                
                # If all else fails, use a default value
                if not user_id:
                    user_id = 1050  # Default auditor ID
                    print(f"DEBUG: Using default user_id: {user_id}")
        
        # Update the audit status to "Under review"
        audit.Status = 'Under review'
        audit.ReviewStatus = 0  # Integer value for "Yet to Start" (0)
        audit.CompletionDate = timezone.now()
        audit.save()
        
        print(f"DEBUG: Audit {audit_id} status set to 'Under review' and ReviewStatus set to 0 (Yet to Start)")
        
        # Extract overall comments if provided in request
        overall_comments = request.data.get('overall_comments', 'Overall comments about the audit process')
        
        # Get all findings for this audit and mark them as checked if not already
        findings = AuditFinding.objects.filter(AuditId=audit_id)
        updated_findings = 0
        
        for finding in findings:
            if finding.Check != '2':  # If not already marked as completed
                finding.Check = '2'  # Mark as completed
                finding.CheckedDate = timezone.now()
                finding.save()
                updated_findings += 1
        
        print(f"DEBUG: Updated {updated_findings} of {len(findings)} audit findings to 'Completed'")
        
        # Get the next version number
        version = "A1"  # Default
        
        with connection.cursor() as cursor:
            cursor.execute("""
                SELECT Version FROM audit_version 
                WHERE AuditId = %s AND Version LIKE 'A%'
                ORDER BY Version DESC
                LIMIT 1
            """, [audit_id])
            
            existing_version = cursor.fetchone()
            if existing_version:
                current_version = existing_version[0]
                print(f"DEBUG: Found existing audit version: {current_version}")
                
                # Extract number part and increment it
                try:
                    version_num = int(current_version[1:])
                    version = f"A{version_num + 1}"
                    print(f"DEBUG: Incrementing from existing version {current_version} to {version}")
                except ValueError:
                    print(f"DEBUG: Could not increment version '{current_version}', using default A1")
                    version = "A1"
        
        # Get all findings in structured JSON format
        structured_json = get_audit_findings_json(audit_id, overall_comments)
        print("DEBUG: Formatted JSON structure for audit findings:")
        print(json.dumps(structured_json, indent=2))
        
        # Create an audit version with the new version number
        version_result = None
        if user_id:
            print(f"DEBUG: Creating audit version {version} with user_id: {user_id}")
            version_result = create_audit_version(audit_id, user_id, version)
        
        response_data = {
            'message': 'Audit submitted for review successfully',
            'audit_id': audit_id,
            'status': 'Under review',
            'review_status': 'Yet to Start',
            'review_status_int': 0,
            'completion_date': audit.CompletionDate.strftime('%Y-%m-%d %H:%M:%S'),
            'findings_updated': updated_findings
        }
        
        # Include version creation result if applicable
        if version_result:
            response_data['version_created'] = version_result['success']
            if version_result['success']:
                response_data['version'] = version_result['version']
                response_data['findings_saved'] = version_result['findings_count']
            else:
                response_data['version_error'] = version_result['error']
        
        return Response(response_data, status=status.HTTP_200_OK)
        
    except Exception as e:
        print(f"ERROR in submit_audit_findings: {str(e)}")
        return Response({'error': str(e)}, status=status.HTTP_400_BAD_REQUEST)
    
 
@api_view(['POST'])
def allocate_policy(request):
    """
    Allocate a policy to users and create audit findings based on the selected scope
    """
    try:
        data = request.data.copy()
        print("="*50)
        print("DEBUG: allocate_policy was called")
        print("Received data:", data)
        print("Request headers:", request.headers)
        print("Request method:", request.method)
        
        # If assignee is not provided in the request data, get it from the session
        if 'assignee' not in data or not data['assignee']:
            print("DEBUG: No assignee provided, using default from session")
            data['assignee'] = request.session.get('user_id', 1020)
        
        # Debug: convert types as needed
        for key, value in data.items():
            print(f"DEBUG: {key} = {value} (type: {type(value)})")
            
        serializer = PolicyAllocationSerializer(data=data)
        
        if serializer.is_valid():
            print("DEBUG: Serializer validation passed")
            # Get framework_id from validated data
            framework_id = serializer.validated_data['framework']
            policy_id = serializer.validated_data.get('policy')
            subpolicy_id = serializer.validated_data.get('subpolicy')
            
            # Get user IDs
            assignee_id = serializer.validated_data['assignee']
            auditor_id = serializer.validated_data['auditor']
            reviewer_id = serializer.validated_data.get('reviewer')

            print(f"Creating audit with: assignee={assignee_id}, auditor={auditor_id}, reviewer={reviewer_id}")
            print(f"Other fields: framework={framework_id}, policy={policy_id}, subpolicy={subpolicy_id}")
            print(f"Date/Type: duedate={serializer.validated_data['duedate']}, frequency={serializer.validated_data['frequency']}, audit_type={serializer.validated_data['audit_type']}")

            # Create a timezone-aware datetime for AssignedDate
            assigned_date = timezone.now()

            # Get selected reports from validated data
            selected_reports = serializer.validated_data.get('selected_reports', [])
            print(f"Selected reports: {selected_reports}")

            # Convert selected reports to JSON string
            if selected_reports:
                reports_array = []
                for i, report_id in enumerate(selected_reports, 1):
                    # Get report details from database
                    try:
                        with connection.cursor() as cursor:
                            cursor.execute("""
                                SELECT 
                                    ReportId,
                                    AuditId,
                                    Report,
                                    PolicyId,
                                    SubPolicyId,
                                    FrameworkId
                                FROM audit_report
                                WHERE ReportId = %s
                            """, [report_id])
                            report_data = cursor.fetchone()
                            
                            if report_data:
                                report_dict = {
                                    f"Report_{i}": {
                                        "ReportId": report_data[0],
                                        "AuditId": report_data[1],
                                        "Report": report_data[2],
                                        "PolicyId": report_data[3],
                                        "SubPolicyId": report_data[4],
                                        "FrameworkId": report_data[5]
                                    }
                                }
                                reports_array.append(report_dict)
                            
                    except Exception as e:
                            print(f"Error fetching report {report_id}: {str(e)}")
                            continue        
                reports_dict = {"reports": reports_array}
                reports_json = json.dumps(reports_dict)
                print(f"DEBUG: Generated reports JSON: {reports_json}")
            else:
                reports_json = None



            # Create the Audit record with AssignedDate and Report
            audit = Audit.objects.create(
                Assignee_id=assignee_id,
                Auditor_id=auditor_id,
                Reviewer_id=reviewer_id,
                FrameworkId_id=framework_id,
                PolicyId_id=policy_id,
                SubPolicyId_id=subpolicy_id,
                DueDate=serializer.validated_data['duedate'],
                Frequency=serializer.validated_data['frequency'],
                AuditType=serializer.validated_data['audit_type'],
                Status='Yet to Start',
                AssignedDate=assigned_date,  # Set the AssignedDate in Audit
                Reports=reports_json  # Save selected reports in JSON format
            )

            # Get compliances based on the selected scope using helper function
            compliances = get_compliances_by_scope(framework_id, policy_id, subpolicy_id)
            print(f"DEBUG: Found {len(compliances)} compliances for the selected scope")

            # Create audit findings for each compliance
            audit_findings = []
            
            for compliance in compliances:
                try:
                    print(f"DEBUG: Creating audit finding for compliance ID: {compliance.ComplianceId}")
                    
                    # Map Criticality value from compliance to MajorMinor value
                    major_minor = None
                    if hasattr(compliance, 'Criticality'):
                        criticality = compliance.Criticality
                        # Convert to string if it's a number
                        if isinstance(criticality, int) or (isinstance(criticality, str) and criticality.isdigit()):
                            criticality_value = int(criticality) if isinstance(criticality, str) else criticality
                            if criticality_value == 0:
                                major_minor = '0'  # Minor
                            elif criticality_value == 1:
                                major_minor = '1'  # Major
                            elif criticality_value == 2:
                                major_minor = '2'  # Not Applicable
                        # Handle string values of Criticality
                        elif isinstance(criticality, str):
                            if criticality.lower() == 'minor':
                                major_minor = '0'
                            elif criticality.lower() == 'major':
                                major_minor = '1'
                            elif criticality.lower() == 'not applicable':
                                major_minor = '2'
                    
                    print(f"DEBUG: Setting initial MajorMinor value to '{major_minor}' for compliance ID: {compliance.ComplianceId}")
                    
                    audit_finding = AuditFinding.objects.create(
                        AuditId=audit,
                        ComplianceId=compliance,
                        UserId_id=auditor_id,
                        Evidence='',
                        Check='0',  # 0 = Yet to Start
                        Comments='',
                        MajorMinor=major_minor,  # Set initial MajorMinor value
                        AssignedDate=assigned_date  # Use the same AssignedDate as the audit
                    )
                    audit_findings.append(f"Added finding for compliance {compliance.ComplianceId}")
                    print(f"DEBUG: Successfully created audit finding for compliance {compliance.ComplianceId}")
                except Exception as e:
                    print(f"ERROR creating audit finding: {e}")
                    print(f"Detailed error info - AuditId: {audit.AuditId}, ComplianceId: {compliance.ComplianceId}, UserId: {auditor_id}")
                    # If there's an error with audit findings, delete the audit to maintain consistency
                    audit.delete()
                    return Response({
                        'error': f'Failed to create audit finding: {str(e)}'
                    }, status=status.HTTP_400_BAD_REQUEST)

            print(f"DEBUG: Created {len(audit_findings)} audit findings")
            
            # Send notifications to assigned users
            try:
                notification_service = NotificationService()
                
                # Get user information
                try:
                    # Get user information
                    auditor = Users.objects.get(UserId=auditor_id)
                    print("Auditor ",auditor,"============================================================")
                    assignee = Users.objects.get(UserId=assignee_id)
                    print("Assignee ",assignee,"============================================================")
                    
                    # Get framework details - use only the columns that exist in the table
                    with connection.cursor() as cursor:
                        cursor.execute("""
                            SELECT FrameworkId, FrameworkName
                            FROM frameworks
                            WHERE FrameworkId = %s
                        """, [framework_id])
                        framework_data = cursor.fetchone()
                    
                    if not framework_data:
                        raise Framework.DoesNotExist("Framework not found")
                    
                    framework_name = framework_data[1]
                    print("Framework name:", framework_name, "============================================================")
                    
                    # Create scope description
                    scope_description = f"Framework: {framework_name}"
                    print("Scope Description ",scope_description,"============================================================")
                    
                    # Get policy name if needed
                    if policy_id:
                        with connection.cursor() as cursor:
                            cursor.execute("""
                                SELECT PolicyId, PolicyName
                                FROM policies
                                WHERE PolicyId = %s
                            """, [policy_id])
                            policy_data = cursor.fetchone()
                            
                        if policy_data:
                            policy_name = policy_data[1]
                            scope_description += f", Policy: {policy_name}"
                    
                    # Get subpolicy name if needed
                    if subpolicy_id:
                        with connection.cursor() as cursor:
                            cursor.execute("""
                                SELECT SubPolicyId, SubPolicyName
                                FROM subpolicies
                                WHERE SubPolicyId = %s
                            """, [subpolicy_id])
                            subpolicy_data = cursor.fetchone()
                            
                        if subpolicy_data:
                            subpolicy_name = subpolicy_data[1]
                            scope_description += f", SubPolicy: {subpolicy_name}"
                    
                    # Format the due date
                    due_date = serializer.validated_data['duedate'].strftime('%Y-%m-%d')
                    
                    # Set a title for the audit based on framework/policy
                    audit_title = f"Audit of {framework_name}"
                    if policy_id and 'policy_name' in locals():
                        audit_title += f" - {policy_name}"
                    
                    # Update the audit title
                    audit.Title = audit_title
                    audit.save()
                    
                    # Get username fields
                    auditor_name = auditor.UserName if hasattr(auditor, 'UserName') else f"User {auditor_id}"
                    assignee_name = assignee.UserName if hasattr(assignee, 'UserName') else f"User {assignee_id}"
                    
                    # Get email fields using direct DB query
                    with connection.cursor() as cursor:
                        cursor.execute("""
                            SELECT Email FROM users WHERE UserId = %s
                        """, [auditor_id])
                        auditor_email_data = cursor.fetchone()
                        
                        cursor.execute("""
                            SELECT Email FROM users WHERE UserId = %s
                        """, [assignee_id])
                        assignee_email_data = cursor.fetchone()
                        
                        if reviewer_id:
                            cursor.execute("""
                                SELECT Email, UserName FROM users WHERE UserId = %s
                            """, [reviewer_id])
                            reviewer_data = cursor.fetchone()
                            if reviewer_data:
                                reviewer_email = reviewer_data[0]
                                reviewer_name = reviewer_data[1] if reviewer_data[1] else f"User {reviewer_id}"
                            else:
                                reviewer_email = None
                                reviewer_name = f"User {reviewer_id}"
                    
                    auditor_email = auditor_email_data[0] if auditor_email_data else None
                    assignee_email = assignee_email_data[0] if assignee_email_data else None
                    
                    # 1. Notify the auditor
                    if auditor_email:
                        auditor_notification = {
                            'notification_type': 'auditAssigned',
                            'email': auditor_email,
                            'email_type': 'gmail',  # Or 'microsoft' based on your configuration
                            'template_data': [
                                auditor_name,
                                audit_title,
                                scope_description,
                                due_date
                            ]
                        }
                        notification_result = notification_service.send_multi_channel_notification(auditor_notification)
                        print(f"DEBUG: Sent audit assignment notification to auditor {auditor_email}, result: {notification_result['success']}")
                    
                    # 2. Notify the assignee if different from auditor
                    if assignee_id != auditor_id and assignee_email:
                        assignee_notification = {
                            'notification_type': 'auditAssigned',
                            'email': assignee_email,
                            'email_type': 'gmail',
                            'template_data': [
                                assignee_name,
                                audit_title,
                                scope_description,
                                due_date
                            ]
                        }
                        notification_result = notification_service.send_multi_channel_notification(assignee_notification)
                        print(f"DEBUG: Sent audit assignment notification to assignee {assignee_email}, result: {notification_result['success']}")
                    
                    # 3. Notify the reviewer if specified
                    if reviewer_id and reviewer_email:
                        reviewer_notification = {
                            'notification_type': 'auditAssigned',
                            'email': reviewer_email,
                            'email_type': 'gmail',
                            'template_data': [
                                reviewer_name,
                                audit_title,
                                scope_description,
                                due_date
                            ]
                        }
                        notification_result = notification_service.send_multi_channel_notification(reviewer_notification)
                        print(f"DEBUG: Sent audit assignment notification to reviewer {reviewer_email}, result: {notification_result['success']}")
                
                except Exception as e:
                    print(f"ERROR: Failed to get user details for notifications: {str(e)}")
                    # Continue even if getting user details fails
            
            except Exception as e:
                print(f"ERROR: Failed to send notifications: {str(e)}")
                # Continue even if notifications fail
            
            return Response({
                'message': 'Policy allocated successfully',
                'audit_id': audit.AuditId,
                'findings_created': len(audit_findings),
                'scope': {
                    'framework_id': framework_id,
                    'policy_id': policy_id,
                    'subpolicy_id': subpolicy_id
                }
            }, status=status.HTTP_201_CREATED)
        else:
            print("="*50)
            print("DEBUG: Serializer validation failed")
            print("Validation errors:", serializer.errors)
            
            error_messages = {}
            for field, errors in serializer.errors.items():
                print(f"Field '{field}' errors:", errors)
                if field in data:
                    print(f"  Value provided: '{data[field]}' (type: {type(data[field])})")
                else:
                    print(f"  No value provided for field")
                error_messages[field] = [str(error) for error in errors]
            
            return Response({
                'error': 'Validation failed',
                'details': error_messages
            }, status=status.HTTP_400_BAD_REQUEST)
    except Exception as e:
        print(f"ERROR in allocate_policy: {str(e)}")
        import traceback
        traceback.print_exc()
        return Response({
            'error': str(e)
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
    
def create_review_version(audit_id, user_id, compliance_reviews=None, overall_comments=None, custom_version=None):
    """
    Create a new entry in the audit_version table with "R" prefix for reviewer versions
    - Extract all audit findings for the given audit_id
    - Add review status and comments for each compliance item
    - Add overall review comments
    - Use "R1" as version prefix for reviewer data
    - Save to the audit_version table
    """
    try:
        print(f"DEBUG: Creating review version for audit_id: {audit_id}, user_id: {user_id}")
        
        # Check if table exists
        with connection.cursor() as cursor:
            cursor.execute("""
                SELECT COUNT(*) FROM information_schema.tables 
                WHERE table_schema = DATABASE() 
                AND table_name = 'audit_version'
            """)
            if cursor.fetchone()[0] == 0:
                print("DEBUG: audit_version table doesn't exist, creating it")
                cursor.execute("""
                    CREATE TABLE IF NOT EXISTS audit_version (
                        AuditId INT,
                        Version VARCHAR(45),
                        ExtractedInfo JSON,
                        UserId INT,
                        ApproverId INT NULL,
                        ApprovedRejected VARCHAR(45) NULL,
                        Date DATETIME,
                        PRIMARY KEY (AuditId, Version)
                    )
                """)
                print("DEBUG: audit_version table created")
        
        # Check if a review version already exists for this audit
        version = custom_version or "R1"  # Use provided version or default to R1
        print(f"DEBUG: Using version: {version}")
        
        existing_version_data = None
        
        with connection.cursor() as cursor:
            cursor.execute("""
                SELECT Version, ExtractedInfo 
                FROM audit_version 
                WHERE AuditId = %s AND Version LIKE 'R%'
                ORDER BY Version DESC
                LIMIT 1
            """, [audit_id])
            
            existing_version_row = cursor.fetchone()
            if existing_version_row:
                existing_version = existing_version_row[0]
                print(f"DEBUG: Found existing review version: {existing_version}")
                # Parse existing data for comparison
                try:
                    if isinstance(existing_version_row[1], dict):
                        existing_version_data = existing_version_row[1]
                    else:
                        existing_version_data = json.loads(existing_version_row[1])
                    print(f"DEBUG: Loaded existing review version data with {len(existing_version_data) if existing_version_data else 0} findings")
                except Exception as e:
                    print(f"DEBUG: Failed to parse existing version data: {str(e)}")
                    existing_version_data = {}
        
        # First get the most recent audit version record (A prefix)
        audit_version_data = None
        with connection.cursor() as cursor:
            cursor.execute("""
                SELECT ExtractedInfo FROM audit_version 
                WHERE AuditId = %s AND Version LIKE 'A%'
                ORDER BY Version DESC, Date DESC
                LIMIT 1
            """, [audit_id])
            
            version_row = cursor.fetchone()
            if version_row:
                try:
                    # Parse the JSON from the audit version
                    if isinstance(version_row[0], dict):
                        audit_version_data = version_row[0]
                    else:
                        audit_version_data = json.loads(version_row[0])
                    print(f"DEBUG: Found base audit version data with {len(audit_version_data) if audit_version_data else 0} findings")
                except Exception as e:
                    print(f"DEBUG: Failed to parse audit version data: {str(e)}")
                    audit_version_data = {}
            else:
                print("DEBUG: No audit version found, will create findings data from scratch")
        
        # If no audit version was found or it's empty, fetch the current findings
        if not audit_version_data or len(audit_version_data) == 0:
            audit_version_data = {}
            with connection.cursor() as cursor:
                cursor.execute("""
                    SELECT 
                        af.ComplianceId,
                        af.Evidence,
                        af.`Check`,
                        af.HowToVerify,
                        af.Impact,
                        af.Recommendation,
                        af.DetailsOfFinding,
                        af.Comments,
                        af.MajorMinor,
                        c.ComplianceItemDescription
                    FROM 
                        audit_findings af
                    JOIN
                        compliance c ON af.ComplianceId = c.ComplianceId
                    WHERE 
                        af.AuditId = %s
                """, [audit_id])
                
                findings = cursor.fetchall()
                columns = ['ComplianceId', 'Evidence', 'Check', 'HowToVerify', 
                           'Impact', 'Recommendation', 'DetailsOfFinding', 
                           'Comments', 'MajorMinor', 'ComplianceItemDescription']
                
                # Construct the JSON data
                for finding in findings:
                    compliance_id = finding[0]  # First column is ComplianceId
                    finding_data = {}
                    
                    # Add all columns except ComplianceId and ComplianceItemDescription
                    for i in range(1, len(columns) - 1):  # Skip the last column
                        finding_data[columns[i]] = finding[i]
                    
                    # Add description as a special field for UI display purposes
                    finding_data['description'] = finding[9]  # ComplianceItemDescription
                    
                    # Add empty fields for review data
                    finding_data['accept_reject'] = "0"  # 0 = Not reviewed
                    finding_data['review_comments'] = ""
                    finding_data['review_status'] = "In Review"  # Default review status
                    
                    # Use ComplianceId as the key in the JSON
                    audit_version_data[str(compliance_id)] = finding_data
        
        # Now update with the current review data
        review_updated = False
        findings_data = audit_version_data.copy() if audit_version_data else {}
        
        # Add or update metadata
        metadata = findings_data.get('__metadata__', {})
        if not isinstance(metadata, dict):
            metadata = {}
        
        old_comments = metadata.get('overall_comments', '')
        
        # Check if comments changed
        comments_changed = overall_comments is not None and old_comments != overall_comments
        if comments_changed:
            print(f"DEBUG: Overall comments changed from '{old_comments}' to '{overall_comments}'")
            metadata['overall_comments'] = overall_comments
            review_updated = True
        else:
            metadata['overall_comments'] = old_comments
            
        # Calculate overall review status from compliance reviews
        if compliance_reviews:
            has_rejected = False
            all_accepted = True
            for review in compliance_reviews:
                if not isinstance(review, dict):
                    continue
                    
                if review.get('review_status') == 'Reject':
                    has_rejected = True
                    all_accepted = False
                elif review.get('review_status') != 'Accept':
                    all_accepted = False
            
            if has_rejected:
                overall_status = 'Reject'
            elif all_accepted and len(compliance_reviews) > 0:
                overall_status = 'Accept'
            else:
                overall_status = 'In Review'
            
            metadata['overall_status'] = overall_status
            print(f"DEBUG: Set overall review status to '{overall_status}'")
            
        metadata['review_date'] = timezone.now().strftime('%Y-%m-%d %H:%M:%S')
        metadata['reviewer_id'] = user_id
        findings_data['__metadata__'] = metadata
        
        # Update with the compliance review data if provided
        if compliance_reviews:
            print(f"DEBUG: Processing {len(compliance_reviews)} compliance reviews")
            for review in compliance_reviews:
                if not isinstance(review, dict):
                    print(f"DEBUG: Skipping non-dict review: {type(review)}")
                    continue
                    
                compliance_id = str(review.get('compliance_id'))
                
                if not compliance_id:
                    print(f"DEBUG: Skipping review with no compliance_id")
                    continue
                    
                # If we don't have this compliance in our data, try to fetch it
                if compliance_id not in findings_data:
                    print(f"DEBUG: Compliance ID {compliance_id} not found in version data, attempting to fetch from database")
                    try:
                        with connection.cursor() as cursor:
                            cursor.execute("""
                                SELECT 
                                    af.Evidence,
                                    af.`Check`,
                                    af.HowToVerify,
                                    af.Impact,
                                    af.Recommendation,
                                    af.DetailsOfFinding,
                                    af.Comments,
                                    af.MajorMinor,
                                    c.ComplianceItemDescription
                                FROM 
                                    audit_findings af
                                JOIN
                                    compliance c ON af.ComplianceId = c.ComplianceId
                                WHERE 
                                    af.AuditId = %s AND af.ComplianceId = %s
                            """, [audit_id, compliance_id])
                            
                            finding_row = cursor.fetchone()
                            if finding_row:
                                # Add this finding to our data
                                fields = ['Evidence', 'Check', 'HowToVerify', 'Impact', 'Recommendation', 
                                         'DetailsOfFinding', 'Comments', 'MajorMinor']
                                
                                finding_data = {}
                                for i in range(len(fields)):
                                    finding_data[fields[i]] = finding_row[i]
                                
                                # Add description
                                finding_data['description'] = finding_row[8]  # ComplianceItemDescription
                                
                                # Add empty review fields
                                finding_data['accept_reject'] = "0"  # 0 = Not reviewed
                                finding_data['review_comments'] = ""
                                finding_data['review_status'] = "In Review"  # Default review status
                                
                                findings_data[compliance_id] = finding_data
                                print(f"DEBUG: Added missing compliance {compliance_id} from database")
                            else:
                                print(f"DEBUG: Compliance {compliance_id} not found in database, skipping")
                                continue
                    except Exception as e:
                        print(f"ERROR fetching compliance data: {str(e)}")
                        continue
                
                # Get old values for comparison
                old_finding = findings_data[compliance_id]
                old_accept_reject = old_finding.get('accept_reject', "0")
                old_review_comments = old_finding.get('review_comments', '')
                
                # Map review_status to accept_reject value
                new_accept_reject = "0"  # Default (not accepted/rejected)
                if review.get('review_status') == 'Accept':
                    new_accept_reject = "1"  # 1 = Accepted
                elif review.get('review_status') == 'Reject':
                    new_accept_reject = "2"  # 2 = Rejected

                # Get review comments from the review data
                new_review_comments = review.get('review_comments', '')
                new_review_status = review.get('review_status', 'In Review')  # Get the actual string status
                
                # Check if anything changed
                comments_changed = new_review_comments != old_review_comments
                status_changed = new_accept_reject != old_accept_reject
                
                if comments_changed or status_changed:
                    # Update the finding data with review information
                    # For structured JSON compatibility
                    findings_data[compliance_id]['accept_reject'] = new_accept_reject
                    findings_data[compliance_id]['comments'] = new_review_comments  # Using 'comments' for the structured JSON
                    findings_data[compliance_id]['review_comments'] = new_review_comments  # Keep old field for backward compatibility
                    findings_data[compliance_id]['review_status'] = new_review_status  # Store the string status too
                    review_updated = True
                    print(f"DEBUG: Updated finding {compliance_id}: status changed={status_changed}, comments changed={comments_changed}")
        
        # If using an explicit custom version, don't compare with existing data
        if not custom_version and existing_version_data and not review_updated:
            # Additional check - compare the updated data with the existing version
            print("DEBUG: Comparing new data with existing version data")
            
            # If we already have version data with the same review statuses and comments,
            # and if overall comments are the same, don't create a new version
            is_same = True
            
            # Check metadata
            existing_metadata = existing_version_data.get('__metadata__', {})
            if not isinstance(existing_metadata, dict):
                existing_metadata = {}
                
            existing_comments = existing_metadata.get('overall_comments', '')
            if overall_comments is not None and overall_comments != existing_comments:
                is_same = False
                print(f"DEBUG: Overall comments different: '{existing_comments}' vs '{overall_comments}'")
            
            # Check individual findings
            if is_same:
                for compliance_id, finding in findings_data.items():
                    if compliance_id == '__metadata__':
                        continue
                        
                    if compliance_id not in existing_version_data:
                        is_same = False
                        print(f"DEBUG: Compliance {compliance_id} not in existing data")
                        break
                        
                    existing_finding = existing_version_data[compliance_id]
                    if finding.get('accept_reject') != existing_finding.get('accept_reject'):
                        is_same = False
                        print(f"DEBUG: Different accept_reject for {compliance_id}: {finding.get('accept_reject')} vs {existing_finding.get('accept_reject')}")
                        break
                        
                    # Check both comments field options
                    new_comments = finding.get('comments', finding.get('review_comments', ''))
                    existing_comments = existing_finding.get('comments', existing_finding.get('review_comments', ''))
                    if new_comments != existing_comments:
                        is_same = False
                        print(f"DEBUG: Different comments for {compliance_id}")
                        break
                
            if is_same:
                print("DEBUG: No changes detected compared to existing version, using existing version")
                return {
                    'success': True,
                    'audit_id': audit_id,
                    'version': existing_version,  # Use the existing version ID since we're not creating a new one
                    'findings_count': len(findings_data) - 1 if findings_data else 0,  # Subtract 1 for metadata
                    'message': 'No changes detected, using existing version'
                }
        
        # Force creation if using custom version, otherwise check for changes
        if not custom_version and not review_updated and not comments_changed:
            print("DEBUG: No review data was updated, not creating a new version")
            return {
                'success': False,
                'message': 'No review data was updated'
            }
        
        # Structure the data in the expected JSON format
        # We're already using compliance_ids as keys, so the structure is already close
        # Just need to ensure the right field names are used consistently
        for compliance_id, finding in list(findings_data.items()):
            if compliance_id == '__metadata__':
                continue
                
            # Make sure both the "comments" and "review_comments" fields are present
            if 'review_comments' in finding and 'comments' not in finding:
                finding['comments'] = finding['review_comments']
            elif 'comments' in finding and 'review_comments' not in finding:
                finding['review_comments'] = finding['comments']
        
            # Ensure review_status is present for all items
            if 'review_status' not in finding:
                # Derive from accept_reject
                accept_reject = finding.get('accept_reject', "0")
                if accept_reject == "1" or accept_reject == 1:
                    finding['review_status'] = 'Accept'
                elif accept_reject == "2" or accept_reject == 2:
                    finding['review_status'] = 'Reject'
                else:
                    finding['review_status'] = 'In Review'
                    
        # Add overall_comments directly to the root
        if 'overall_comments' not in findings_data and overall_comments:
            findings_data['overall_comments'] = overall_comments
        
        # Convert the findings data to JSON
        extracted_info_json = json.dumps(findings_data)
        
        # Get current timestamp
        current_time = timezone.now()
        
        # Insert the new version
        with connection.cursor() as cursor:
            # First check if a version with this ID already exists
            cursor.execute("""
                SELECT COUNT(*) FROM audit_version
                WHERE AuditId = %s AND Version = %s
            """, [audit_id, version])
            
            version_exists = cursor.fetchone()[0] > 0
            
            if version_exists:
                # Update the existing record
                cursor.execute("""
                    UPDATE audit_version
                    SET ExtractedInfo = %s,
                        UserId = %s,
                        Date = %s
                    WHERE AuditId = %s AND Version = %s
                """, [extracted_info_json, user_id, current_time, audit_id, version])
                print(f"DEBUG: Updated existing review version {version} for audit {audit_id}")
            else:
                # Insert a new record
                cursor.execute("""
                    INSERT INTO audit_version (
                        AuditId, Version, ExtractedInfo, UserId, 
                        ApproverId, ApprovedRejected, Date
                    )
                    VALUES (%s, %s, %s, %s, NULL, NULL, %s)
                """, [audit_id, version, extracted_info_json, user_id, timezone.localtime(current_time_utc)])
                print(f"DEBUG: Created new review version {version} for audit {audit_id}")
        
        # Also update the audit_findings table with the review data so it persists
        if compliance_reviews:
            try:
                with connection.cursor() as cursor:
                    # First check if the required columns exist
                    cursor.execute("""
                        SELECT COUNT(*) 
                        FROM information_schema.COLUMNS 
                        WHERE TABLE_NAME = 'audit_findings' 
                        AND COLUMN_NAME = 'ReviewRejected'
                    """)
                    column_exists = cursor.fetchone()[0] > 0
                    
                    if not column_exists:
                        print("DEBUG: Adding review columns to audit_findings")
                        cursor.execute("""
                            ALTER TABLE audit_findings
                            ADD COLUMN ReviewRejected TINYINT DEFAULT 0,
                            ADD COLUMN ReviewComments TEXT NULL,
                            ADD COLUMN ReviewStatus VARCHAR(50) NULL,
                            ADD COLUMN ReviewDate DATETIME NULL
                        """)
                        print("DEBUG: Columns added successfully")
                
                # Update each finding with review data
                for review in compliance_reviews:
                    if not isinstance(review, dict):
                        continue
                        
                    compliance_id = review.get('compliance_id')
                    if not compliance_id:
                        continue
                        
                    review_status = review.get('review_status', 'In Review')
                    review_comments = review.get('review_comments', '')
                    
                    # Map status to ReviewRejected value
                    review_rejected = 0  # Default to not rejected
                    if review_status == 'Reject':
                        review_rejected = 1
                    
                    cursor.execute("""
                        UPDATE audit_findings
                        SET ReviewRejected = %s,
                            ReviewComments = %s,
                            ReviewStatus = %s,
                            ReviewDate = %s
                        WHERE AuditId = %s AND ComplianceId = %s
                    """, [
                        review_rejected,
                        review_comments,
                        review_status,
                        current_time,
                        audit_id,
                        compliance_id
                    ])
                    
                print(f"DEBUG: Updated audit_findings table with review data")
                
                # If we have overall comments but no specific compliance reviews,
                # make sure we still update the audit table with the comments
                if overall_comments and not compliance_reviews:
                    cursor.execute("""
                        UPDATE audit
                        SET ReviewComments = %s,
                            ReviewDate = %s
                        WHERE AuditId = %s
                    """, [
                        overall_comments,
                        current_time,
                        audit_id
                    ])
                    print(f"DEBUG: Updated audit table with overall comments: {overall_comments}")
            except Exception as e:
                print(f"WARNING: Failed to update audit_findings table: {str(e)}")
                # Continue even if this fails - the version data is the primary storage
        
        return {
            'success': True,
            'audit_id': audit_id,
            'version': version,
            'findings_count': len(findings_data) - 1 if findings_data else 0  # Subtract 1 for metadata
        }
    except Exception as e:
        print(f"ERROR in create_review_version: {str(e)}")
        return {
            'success': False,
            'error': str(e)
        }



@api_view(['GET'])
def fix_audit_table(request):
    """
    Fix the audit table schema by adding any missing columns defined in the models.py
    This is helpful to upgrade older database instances without requiring a full migration
    """
    try:
        print("DEBUG: fix_audit_table called")
        
        # These are the columns we expect to have based on our models.py
        # Order is important - columns with dependencies should be added first
        expected_columns = [
            {"name": "ReviewStatus", "type": "VARCHAR(45)", "nullable": "NULL"},
            {"name": "ReviewComments", "type": "VARCHAR(255)", "nullable": "NULL"},
            {"name": "ReviewStartDate", "type": "DATETIME", "nullable": "NULL"},
            {"name": "ReviewDate", "type": "DATETIME", "nullable": "NULL"},
            {"name": "Evidence", "type": "TEXT", "nullable": "NULL"},
            {"name": "Comments", "type": "TEXT", "nullable": "NULL"}
        ]
        
        columns_added = []
        columns_exist = []
        errors = []
        
        # Check each column and add if it doesn't exist
        with connection.cursor() as cursor:
            for column in expected_columns:
                col_name = column["name"]
                col_type = column["type"]
                nullable = column["nullable"]
                
                try:
                    # Check if column exists
                    cursor.execute(f"""
                        SELECT COUNT(*) 
                        FROM information_schema.COLUMNS 
                        WHERE TABLE_NAME = 'audit' 
                        AND COLUMN_NAME = '{col_name}'
                    """)
                    column_exists = cursor.fetchone()[0] > 0
                    
                    if not column_exists:
                        # Add the column
                        print(f"DEBUG: Adding column {col_name} to audit table")
                        alter_sql = f"""
                            ALTER TABLE audit
                            ADD COLUMN {col_name} {col_type} {nullable}
                        """
                        print(f"DEBUG: Executing SQL: {alter_sql}")
                        cursor.execute(alter_sql)
                        columns_added.append(col_name)
                    else:
                        print(f"DEBUG: Column {col_name} already exists in audit table")
                        columns_exist.append(col_name)
                except Exception as e:
                    error_msg = f"ERROR adding column {col_name}: {str(e)}"
                    print(error_msg)
                    errors.append(error_msg)
        
        return Response({
            'message': 'Audit table schema check completed',
            'columns_added': columns_added,
            'columns_already_exist': columns_exist,
            'errors': errors
        }, status=status.HTTP_200_OK)
    except Exception as e:
        print(f"ERROR in fix_audit_table: {str(e)}")
        return Response({'error': str(e)}, status=status.HTTP_400_BAD_REQUEST)

@api_view(['GET'])
def get_audit_versions(request, audit_id):
    """
    Get all saved versions for a specific audit
    """
    try:
        print(f"DEBUG: get_audit_versions called for audit_id: {audit_id}")
        
        # Check if the audit exists
        try:
            audit = Audit.objects.get(AuditId=audit_id)
        except Audit.DoesNotExist:
            return Response({'error': 'Audit not found'}, status=status.HTTP_404_NOT_FOUND)
        
        # Check if audit_version table exists
        with connection.cursor() as cursor:
            cursor.execute("""
                SELECT COUNT(*) FROM information_schema.tables 
                WHERE table_schema = DATABASE() 
                AND table_name = 'audit_version'
            """)
            if cursor.fetchone()[0] == 0:
                return Response({
                    'message': 'No versions found',
                    'audit_id': audit_id,
                    'versions': []
                }, status=status.HTTP_200_OK)
        
        # Get all versions for this audit
        with connection.cursor() as cursor:
            cursor.execute("""
                SELECT 
                    av.AuditId,
                    av.Version,
                    u.UserName as Auditor,
                    av.UserId,
                    CASE 
                        WHEN av.ApproverId IS NOT NULL THEN (
                            SELECT UserName FROM users WHERE UserId = av.ApproverId
                        )
                        ELSE NULL
                    END as Approver,
                    av.ApproverId,
                    av.ApprovedRejected,
                    av.Date,
                    JSON_LENGTH(av.ExtractedInfo) as FindingsCount,
                    CASE 
                        WHEN av.Version LIKE 'A%' THEN 'Auditor'
                        WHEN av.Version LIKE 'R%' THEN 'Reviewer'
                        ELSE 'Unknown'
                    END as VersionType
                FROM 
                    audit_version av
                LEFT JOIN
                    users u ON av.UserId = u.UserId
                WHERE 
                    av.AuditId = %s
                ORDER BY 
                    av.Version DESC
            """, [audit_id])
            
            columns = [col[0] for col in cursor.description]
            versions = [dict(zip(columns, row)) for row in cursor.fetchall()]
        
        # Format date for each version
        for version in versions:
            if version.get('Date'):
                version['Date'] = version['Date'].strftime('%Y-%m-%d %H:%M:%S')
        
        return Response({
            'audit_id': audit_id,
            'total_versions': len(versions),
            'versions': versions
        }, status=status.HTTP_200_OK)
    
    except Exception as e:
        print(f"ERROR in get_audit_versions: {str(e)}")
        return Response({'error': str(e)}, status=status.HTTP_400_BAD_REQUEST)

@api_view(['GET'])
def get_audit_version_details(request, audit_id, version):
    """
    Get detailed information for a specific audit version
    """
    try:
        print(f"DEBUG: get_audit_version_details called for audit_id: {audit_id}, version: {version}")
        
        # Check if the audit exists
        try:
            audit = Audit.objects.get(AuditId=audit_id)
        except Audit.DoesNotExist:
            return Response({'error': 'Audit not found'}, status=status.HTTP_404_NOT_FOUND)
        
        # Get version details
        with connection.cursor() as cursor:
            # Use parameter placeholders for all values to avoid format string issues
            sql_query = """
                SELECT 
                    av.AuditId,
                    av.Version,
                    u.UserName as Auditor,
                    av.UserId,
                    CASE 
                        WHEN av.ApproverId IS NOT NULL THEN (
                            SELECT UserName FROM users WHERE UserId = av.ApproverId
                        )
                        ELSE NULL
                    END as Approver,
                    av.ApproverId,
                    av.ApprovedRejected,
                    av.Date,
                    av.ExtractedInfo,
                    CASE 
                        WHEN av.Version LIKE 'A%%' THEN 'Auditor'
                        WHEN av.Version LIKE 'R%%' THEN 'Reviewer'
                        ELSE 'Unknown'
                    END as VersionType
                FROM 
                    audit_version av
                LEFT JOIN
                    users u ON av.UserId = u.UserId
                WHERE 
                    av.AuditId = %s AND av.Version = %s
            """
            cursor.execute(sql_query, [audit_id, version])
            
            row = cursor.fetchone()
            if not row:
                return Response({'error': f'Version {version} not found for audit {audit_id}'}, 
                              status=status.HTTP_404_NOT_FOUND)
            
            columns = [col[0] for col in cursor.description]
            version_data = dict(zip(columns, row))
        
        # Format date
        if version_data.get('Date'):
            version_data['Date'] = version_data['Date'].strftime('%Y-%m-%d %H:%M:%S')
        
        # Parse the JSON data
        findings = {}
        metadata = {}
        if version_data.get('ExtractedInfo'):
            try:
                # If it's already a dictionary, use it directly
                if isinstance(version_data['ExtractedInfo'], dict):
                    all_data = version_data['ExtractedInfo']
                # Otherwise, parse it as JSON
                else:
                    all_data = json.loads(version_data['ExtractedInfo'])
                
                # Extract metadata if it exists
                if '__metadata__' in all_data:
                    metadata = all_data.pop('__metadata__')
                    version_data['metadata'] = metadata
                
                # Filter out metadata keys from the findings
                findings = {k: v for k, v in all_data.items() if k != '__metadata__' and k != 'overall_comments'}
                
            except json.JSONDecodeError:
                print(f"ERROR: Invalid JSON in ExtractedInfo for audit {audit_id}, version {version}")
                findings = {}
        
        # Get compliance details to enhance the finding data
        compliance_ids = [k for k in findings.keys() if k.isdigit()]
        compliances = {}
        
        if compliance_ids:
            try:
                # Construct query with placeholders
                placeholders = ', '.join(['%s' for _ in compliance_ids])
                with connection.cursor() as cursor:
                    query = """
                    SELECT 
                        c.ComplianceId,
                        c.ComplianceItemDescription,
                        c.MandatoryOptional,
                        sp.SubPolicyName,
                        p.PolicyName
                    FROM 
                        compliance c
                    JOIN 
                        subpolicies sp ON c.SubPolicyId = sp.SubPolicyId
                    JOIN 
                        policies p ON sp.PolicyId = p.PolicyId
                    WHERE 
                            c.ComplianceId IN (""" + placeholders + ")"
                    
                    # Convert compliance_ids to integers for the database query
                    cursor.execute(query, [int(cid) for cid in compliance_ids])
                
                for row in cursor.fetchall():
                    compliance_id = str(row[0])
                    compliances[compliance_id] = {
                        'description': row[1],
                        'mandatory_optional': row[2],
                        'subpolicy_name': row[3],
                        'policy_name': row[4]
                    }
            except Exception as inner_e:
                print(f"ERROR in get_audit_version_details while fetching compliances: {str(inner_e)}")
                # Continue with empty compliances rather than failing
                compliances = {}
        
        # Combine compliance details with findings
        enhanced_findings = []
        for compliance_id, finding in findings.items():
            if not compliance_id.isdigit():
                # Skip non-numeric keys like metadata
                continue
                
            compliance_info = compliances.get(compliance_id, {
                'description': 'Unknown',
                'mandatory_optional': 'Unknown',
                'subpolicy_name': 'Unknown',
                'policy_name': 'Unknown'
            })
            
            # Combine compliance info with finding data
            enhanced_finding = {
                'compliance_id': compliance_id,
                'description': compliance_info['description'],
                'mandatory_optional': compliance_info['mandatory_optional'],
                'subpolicy_name': compliance_info['subpolicy_name'],
                'policy_name': compliance_info['policy_name']
            }
            # Add all finding data
            for key, value in finding.items():
                enhanced_finding[key] = value
            # For reviewer versions, map accept_reject codes to review_status
            if version.startswith('R'):
                accept_reject_map = {
                    "0": "In Review",
                    "1": "Accept",
                    "2": "Reject"
                }
                if 'accept_reject' in enhanced_finding:
                    accept_reject_code = str(enhanced_finding['accept_reject'])
                    enhanced_finding['review_status'] = accept_reject_map.get(accept_reject_code, "In Review")
            
            enhanced_findings.append(enhanced_finding)
        
        # Sort findings by policy, subpolicy, compliance ID
        try:
            enhanced_findings.sort(key=lambda x: (
                    x.get('policy_name', ''), 
                    x.get('subpolicy_name', ''), 
                    int(x.get('compliance_id', 0))
                ))
        except Exception as sort_error:
            print(f"WARNING: Could not sort enhanced_findings: {str(sort_error)}")
            # Continue without sorting
            
        # Remove the raw JSON from the response
        if 'ExtractedInfo' in version_data:
            del version_data['ExtractedInfo']
        
        return Response({
            'version_info': version_data,
            'findings_count': len(enhanced_findings),
            'findings': enhanced_findings
        }, status=status.HTTP_200_OK)
    
    except Exception as e:
        error_msg = str(e)
        print(f"ERROR in get_audit_version_details: {error_msg}")
        
        # For debugging only - in production, use the generic message
        if settings.DEBUG:
            return Response({'error': f'Error retrieving audit version details: {error_msg}'}, 
                          status=status.HTTP_400_BAD_REQUEST)
        else:
            return Response({'error': 'Error retrieving audit version details'}, 
                          status=status.HTTP_400_BAD_REQUEST)

def get_latest_version_data(audit_id):
    """
    Helper to get the latest version data
    """
    try:
        with connection.cursor() as cursor:
            cursor.execute("""
                SELECT ExtractedInfo 
                FROM audit_version 
                WHERE AuditId = %s 
                ORDER BY Date DESC 
                LIMIT 1
            """, [audit_id])
            
            row = cursor.fetchone()
            if row:
                return json.loads(row[0]) if isinstance(row[0], str) else row[0]
            return None
    except Exception as e:
        print(f"ERROR in get_latest_version_data: {str(e)}")
        return None
  
@api_view(['POST'])
def save_review_progress(request, audit_id):
    """
    Save reviewer progress - always creates new version with reviewer_status and reviewer_comments in JSON format
    """
    try:
        # Validate audit state
        audit = Audit.objects.get(AuditId=audit_id)
        if audit.Status != 'Under review':
            return Response({
                'error': 'Cannot save review when audit is not under review'
            }, status=status.HTTP_400_BAD_REQUEST)
 
        # Get latest version data
        latest_data = get_latest_version_data(audit_id)
        if not latest_data:
            return Response({'error': 'No version data found'}, status=status.HTTP_404_NOT_FOUND)
 
        # Extract compliance reviews from request data
        compliance_reviews = request.data.get('compliance_reviews', [])
        overall_comments = request.data.get('review_comments', '')
        save_only = request.data.get('save_only', False)  # New parameter to control status update
        cancel_action = request.data.get('cancel_action', False)  # Flag to indicate this was a cancel action
       
        print(f"DEBUG: Saving review progress for audit {audit_id} with {len(compliance_reviews)} compliance reviews")
        print(f"DEBUG: Overall comments: {overall_comments}")
        print(f"DEBUG: Save only (no status update): {save_only}")
        
        if cancel_action:
            print(f"DEBUG: This is a CANCEL action - will only save version without any status changes or reports")
       
        # Create the structured JSON data for the version
        structured_data = latest_data.copy() if latest_data else {}
       
        # Add or update metadata
        metadata = structured_data.get('__metadata__', {})
        if not isinstance(metadata, dict):
            metadata = {}
           
        metadata['overall_comments'] = overall_comments
        metadata['review_date'] = timezone.now().strftime('%Y-%m-%d %H:%M:%S')
        metadata['reviewer_id'] = request.session.get('user_id', 1020)  # Default to reviewer ID
       
        # Calculate overall review status from compliance reviews
        has_rejected = False
        all_accepted = True
       
        for review in compliance_reviews:
            if not isinstance(review, dict):
                continue
               
            compliance_id = str(review.get('compliance_id'))
            if not compliance_id:
                continue
               
            # Get current compliance data or initialize it
            if compliance_id not in structured_data:
                structured_data[compliance_id] = {}
               
            review_status = review.get('review_status', 'In Review')
            review_comments = review.get('review_comments', '')
           
            # Update the status tracking variables
            if review_status == 'Reject':
                has_rejected = True
                all_accepted = False
            elif review_status != 'Accept':
                all_accepted = False
               
            # Update fields in the structured data
            structured_data[compliance_id]['review_status'] = review_status
            structured_data[compliance_id]['review_comments'] = review_comments
           
            # Also add accept_reject field for compatibility
            accept_reject = "0"  # Default (In Review)
            if review_status == 'Accept':
                accept_reject = "1"
            elif review_status == 'Reject':
                accept_reject = "2"
               
            structured_data[compliance_id]['accept_reject'] = accept_reject
           
            print(f"DEBUG: Updated compliance {compliance_id} with status={review_status}, comments={review_comments}")
       
        # Set overall status in metadata
        if has_rejected:
            overall_status = 'Reject'
            approvedRejected = 'Rejected'
           
            # Update audit status back to "Work In Progress" only if not save_only mode
            if not save_only:
                try:
                    audit = Audit.objects.get(AuditId=audit_id)
                    if audit.Status == 'Under review':
                        audit.Status = 'Work In Progress'
                        audit.save()
                        print(f"DEBUG: Audit {audit_id} has rejections, setting status to 'Work In Progress'")
                except Exception as e:
                    print(f"ERROR updating audit status: {str(e)}")
            else:
                print(f"DEBUG: Save only mode - not updating audit status despite rejections")
                # In save_only mode, set approvedRejected to None to prevent updating audit_findings
                approvedRejected = None
               
        elif all_accepted and len(compliance_reviews) > 0:
            overall_status = 'Accept'
            
            # Only set approvedRejected if not in save_only mode
            if not save_only:
                approvedRejected = 'Approved'
                
                # Update audit status to "Completed" - ONLY if not in save_only mode
                try:
                    audit = Audit.objects.get(AuditId=audit_id)
                    audit.Status = 'Completed'
                    audit.save()
                    print(f"DEBUG: Audit {audit_id} approved, setting status to 'Completed'")
                   
                    # Save approved data to audit_findings table
                    try:
                        with connection.cursor() as cursor:
                            for review in compliance_reviews:
                                compliance_id = review.get('compliance_id')
                                review_status = review.get('review_status', 'Accept')
                                review_comments = review.get('review_comments', '')
                               
                                # Find the compliance in structured_data to get compliance_status
                                compliance_data = structured_data.get(str(compliance_id), {})
                                compliance_status = compliance_data.get('compliance_status', '')
                               
                                # Map compliance_status to Check value
                                check_value = '0'  # Default: Not Started (Not Compliant)
                                if compliance_status == 'Fully Compliant':
                                    check_value = '2'  # Completed
                                elif compliance_status == 'Partially Compliant':
                                    check_value = '1'  # In Progress
                                elif compliance_status == 'Not Applicable':
                                    check_value = '3'  # Not Applicable
                               
                                # Update audit_findings with approved status
                                cursor.execute("""
                                    UPDATE audit_findings
                                    SET ReviewStatus = %s,
                                        ReviewComments = %s,
                                        ReviewRejected = %s,
                                        ReviewDate = NOW(),
                                        `Check` = %s
                                    WHERE AuditId = %s AND ComplianceId = %s
                                """, [
                                    review_status,
                                    review_comments,
                                    0,  # Not rejected as this is approved
                                    check_value,
                                    audit_id,
                                    compliance_id
                                ])
                        print(f"DEBUG: Successfully updated audit_findings table with approved data from save_review_progress")
                       
                        # Generate and upload report since all findings are accepted
                        try:
                            from .report_utils import generate_and_upload_report
                            from .checklist_utils import update_lastchecklistitem_verified
                           
                            # Update lastchecklistitemverified table
                            update_result = update_lastchecklistitem_verified(audit_id)
                            if update_result:
                                print(f"DEBUG: Successfully updated lastchecklistitemverified table for audit {audit_id}")
                            else:
                                print(f"WARNING: Failed to update lastchecklistitemverified table for audit {audit_id}")
                           
                            # Get the user ID from the session
                            user_id = request.session.get('user_id', 'system')
                           
                            # Generate and upload the report asynchronously
                            import threading
                            thread = threading.Thread(
                                target=generate_and_upload_report,
                                args=(audit_id, user_id)
                            )
                            thread.daemon = True
                            thread.start()
                           
                            print(f"DEBUG: Started report generation thread for audit {audit_id} after all findings accepted")
                        except Exception as e:
                            print(f"ERROR: Failed to start report generation or update checklist: {str(e)}")
                            # Don't fail the save operation if report generation fails
                    except Exception as e:
                        print(f"ERROR updating audit_findings in save_review_progress: {str(e)}")
                except Exception as e:
                    print(f"ERROR updating audit status: {str(e)}")
            else:
                # In save_only mode (Cancel button was clicked)
                approvedRejected = None
                print(f"DEBUG: Save only mode (Cancel button) - not updating audit status or generating report despite all accepted findings")
        else:
            overall_status = 'In Review'
            approvedRejected = None  # Don't set ApprovedRejected until final decision
           
        metadata['overall_status'] = overall_status
        # Only set ApprovedRejected in metadata if we're not in save_only mode
        if approvedRejected and not save_only:
            metadata['ApprovedRejected'] = approvedRejected
        structured_data['__metadata__'] = metadata
       
        # Add overall_comments directly to the root for backward compatibility
        structured_data['overall_comments'] = overall_comments
       
        # Create a new version with reviewer data
        new_version = create_new_version(
            audit_id,
            request.session.get('user_id', 1020),  # Default to reviewer ID
            structured_data,
            "R"  # Always R for reviewer changes
        )
       
        print(f"DEBUG: Created new reviewer version {new_version} with {len(structured_data)} entries")
 
        return Response({
            'message': 'Review saved in new version',
            'review_version': new_version
        }, status=status.HTTP_200_OK)
 
    except Exception as e:
        print(f"ERROR in save_review_progress: {str(e)}")
        return Response({'error': str(e)}, status=status.HTTP_400_BAD_REQUEST)
 

@api_view(['GET'])
def check_audit_version(request, audit_id):
    """
    Debug endpoint to check if an audit version exists for a given audit ID
    """
    try:
        print(f"DEBUG: check_audit_version called for audit_id: {audit_id}")
        
        # Check if the audit exists
        try:
            audit = Audit.objects.get(AuditId=audit_id)
        except Audit.DoesNotExist:
            return Response({'error': 'Audit not found'}, status=status.HTTP_404_NOT_FOUND)
        
        # Check if audit_version table exists
        with connection.cursor() as cursor:
            cursor.execute(
                "SELECT COUNT(*) FROM information_schema.tables WHERE table_schema = DATABASE() AND table_name = 'audit_version'",
                []
            )
            if cursor.fetchone()[0] == 0:
                return Response({
                    'audit_id': audit_id,
                    'table_exists': False,
                    'message': 'audit_version table does not exist'
                }, status=status.HTTP_200_OK)
            
            # Check for auditor versions for this audit
            cursor.execute(
                "SELECT Version, JSON_LENGTH(ExtractedInfo) as FindingsCount, Date, UserId, ApproverId, ApprovedRejected FROM audit_version WHERE AuditId = %s AND Version LIKE %s ORDER BY Version DESC",
                [audit_id, "A%"]
            )
            
            columns = [col[0] for col in cursor.description]
            auditor_versions = []
            for row in cursor.fetchall():
                version_data = dict(zip(columns, row))
                if version_data.get('Date'):
                    version_data['CreatedDate'] = version_data['Date'].strftime('%Y-%m-%d %H:%M:%S')
                auditor_versions.append(version_data)
            
            # Check for reviewer versions for this audit
            cursor.execute(
                "SELECT Version, JSON_LENGTH(ExtractedInfo) as FindingsCount, Date, UserId, ApproverId, ApprovedRejected FROM audit_version WHERE AuditId = %s AND Version LIKE %s ORDER BY Version DESC",
                [audit_id, "R%"]
            )
            
            columns = [col[0] for col in cursor.description]
            reviewer_versions = []
            for row in cursor.fetchall():
                version_data = dict(zip(columns, row))
                if version_data.get('Date'):
                    version_data['CreatedDate'] = version_data['Date'].strftime('%Y-%m-%d %H:%M:%S')
                reviewer_versions.append(version_data)
            
            # Get current audit status
            audit_status = audit.Status
            
            # Check if audit has proper user IDs
            user_info = {}
            try:
                user_info['auditor_id'] = audit.auditor_id if hasattr(audit, 'auditor_id') else None
                user_info['reviewer_id'] = audit.reviewer_id if hasattr(audit, 'reviewer_id') else None
            except Exception as e:
                user_info['error'] = str(e)
            
            # Try manual SQL to get user IDs
            try:
                cursor.execute(
                    "SELECT auditor, reviewer FROM audit WHERE AuditId = %s",
                    [audit_id]
                )
                row = cursor.fetchone()
                if row:
                    user_info['sql_auditor_id'] = row[0]
                    user_info['sql_reviewer_id'] = row[1]
            except Exception as e:
                user_info['sql_error'] = str(e)
        
        # Get current session user_id
        session_user_id = request.session.get('user_id')
        
        # Determine if we should use reviewer version
        use_reviewer_version = False
        latest_version = None
        
        # If this is a reviewer (match session_user_id with reviewer_id)
        if (session_user_id and 
            (session_user_id == user_info.get('reviewer_id') or 
             session_user_id == user_info.get('sql_reviewer_id'))):
            use_reviewer_version = True
            if reviewer_versions:
                latest_version = reviewer_versions[0]
        
        # If no reviewer version or session doesn't match reviewer, use auditor version
        if not latest_version and auditor_versions:
            latest_version = auditor_versions[0]
        
        return Response({
            'audit_id': audit_id,
            'audit_status': audit_status,
            'auditor_versions_found': len(auditor_versions),
            'auditor_versions': auditor_versions,
            'reviewer_versions_found': len(reviewer_versions),
            'reviewer_versions': reviewer_versions,
            'session_user_id': session_user_id,
            'use_reviewer_version': use_reviewer_version,
            'recommended_version': latest_version['Version'] if latest_version else None,
            'user_info': user_info
        }, status=status.HTTP_200_OK)
    
    except Exception as e:
        print(f"ERROR in check_audit_version: {str(e)}")
        import traceback
        traceback.print_exc()
        return Response({'error': str(e)}, status=status.HTTP_400_BAD_REQUEST)
 
# Add this helper function to determine next version number
def get_next_version_number(audit_id, prefix):
    """
    Get the next version number for an audit
    """
    try:
        with connection.cursor() as cursor:
            # Get the latest version for this audit with the given prefix
            cursor.execute(
                """
                SELECT Version FROM audit_version 
                WHERE AuditId = %s AND Version LIKE %s
                ORDER BY Date DESC
                LIMIT 1
                """,
                [audit_id, f"{prefix}%"]
            )
            latest = cursor.fetchone()
            
            if latest:
                # Extract the number part and increment
                version = latest[0]
                print(f"DEBUG: Latest version found: {version}")
                
                if len(version) > 1 and version[0] == prefix:
                    try:
                        number = int(version[1:])
                        next_number = number + 1
                        return f"{prefix}{next_number}"
                    except ValueError:
                        print(f"DEBUG: Could not parse number from version {version}, using {prefix}1")
                        return f"{prefix}1"
            
            # No version found, start with 1
            print(f"DEBUG: No previous version found, starting with {prefix}1")
            return f"{prefix}1"
    except Exception as e:
        print(f"ERROR: Exception in get_next_version_number: {str(e)}")
        # Fallback to a safe default
        return f"{prefix}1"

@api_view(['GET'])
def load_review_data(request, audit_id):
    """
    Load the latest audit version data when a reviewer clicks 'Continue Review' button
    - Fetches the latest audit version (A-prefix) that is submitted for review
    - Returns the JSON data in a format for the reviewer to start reviewing
    - If any review data already exists (R-prefix), it loads that instead
    """
    try:
        print(f"DEBUG: load_review_data called for audit_id: {audit_id}")
        
        # Find the audit
        try:
            audit = Audit.objects.get(AuditId=audit_id)
        except Audit.DoesNotExist:
            return Response({'error': 'Audit not found'}, status=status.HTTP_404_NOT_FOUND)
        
        # Check if audit is in the correct state for review
        if audit.Status != 'Under review':
            return Response({
                'error': 'Cannot load review data when audit is not under review',
                'current_audit_status': audit.Status
            }, status=status.HTTP_400_BAD_REQUEST)
        
        # Get user ID from session
        user_id = request.session.get('user_id')
        if not user_id and hasattr(audit, 'reviewer_id'):
            user_id = audit.reviewer_id
            
        if not user_id:
            # Default to 1020 if no user_id can be determined
            user_id = 1020
            print(f"DEBUG: No user_id found in session or audit, using default reviewer ID: {user_id}")
        
        # Check if there is already a review version (R-prefix)
        latest_review_version = None
        with connection.cursor() as cursor:
            # Fix SQL query to avoid formatting errors
            cursor.execute(
                "SELECT Version, ExtractedInfo FROM audit_version WHERE AuditId = %s AND Version LIKE %s ORDER BY Version DESC LIMIT 1",
                [audit_id, "R%"]
            )
            
            review_row = cursor.fetchone()
            if review_row:
                latest_review_version = review_row[0]
                # Handle both dict and string JSON representations
                if isinstance(review_row[1], dict):
                    version_data = review_row[1]
                else:
                    try:
                        version_data = json.loads(review_row[1]) if review_row[1] else {}
                    except Exception as e:
                        print(f"ERROR parsing JSON from review version: {str(e)}")
                        version_data = {}
                
                print(f"DEBUG: Found existing review version: {latest_review_version}")
                
                # Return the data from the existing review version
                return Response({
                    'audit_id': audit_id,
                    'version': latest_review_version,
                    'data': version_data,
                    'message': f'Loaded existing review version {latest_review_version}'
                }, status=status.HTTP_200_OK)
        
        # If no review version exists, get the latest audit version (A-prefix)
        latest_audit_version = None
        with connection.cursor() as cursor:
            # Fix SQL query to avoid formatting errors
            cursor.execute(
                "SELECT Version, ExtractedInfo FROM audit_version WHERE AuditId = %s AND Version LIKE %s ORDER BY Version DESC LIMIT 1",
                [audit_id, "A%"]
            )
            
            audit_row = cursor.fetchone()
            if not audit_row:
                # If no audit version exists, create empty data structure
                print(f"DEBUG: No audit version found for audit_id: {audit_id}")
                try:
                    version_data = get_audit_findings_json(audit_id, "")
                    latest_audit_version = "A1"
                except Exception as e:
                    print(f"ERROR creating audit findings JSON: {str(e)}")
                    # Create minimal structure if we can't get findings
                    version_data = {
                        "__metadata__": {
                            "reviewer_id": user_id,
                            "review_date": timezone.now().strftime('%Y-%m-%d %H:%M:%S'),
                            "version_type": "Reviewer"
                        },
                        "overall_comments": "Overall comments about the review process"
                    }
                    latest_audit_version = "New"
            else:
                latest_audit_version = audit_row[0]
                # Handle both dict and string JSON representations
                if isinstance(audit_row[1], dict):
                    version_data = audit_row[1]
                else:
                    try:
                        version_data = json.loads(audit_row[1]) if audit_row[1] else {}
                    except Exception as e:
                        print(f"ERROR parsing JSON from audit version: {str(e)}")
                        version_data = {}
                print(f"DEBUG: Found audit version: {latest_audit_version}")
                # Make sure the data has the expected fields for review
                # Process each compliance item to ensure it has accept_reject and comments fields
                for key, value in list(version_data.items()):
                    # Skip metadata and overall_comments
                    if key == '__metadata__' or key == 'overall_comments':
                        continue
                        
                    if isinstance(value, dict):
                        # Ensure the accept_reject field exists (default: 0)
                        if 'accept_reject' not in value:
                            value['accept_reject'] = "0"
                            
                        # Ensure the comments field exists for reviewer comments
                        if 'comments' not in value:
                            value['comments'] = ""
                            
                # Ensure overall_comments field exists
                if 'overall_comments' not in version_data:
                    version_data['overall_comments'] = "Overall comments about the review process"
                
                # Update metadata for the reviewer
                if '__metadata__' in version_data:
                    metadata = version_data['__metadata__']
                    metadata['reviewer_id'] = user_id
                    metadata['review_date'] = timezone.now().strftime('%Y-%m-%d %H:%M:%S')
                    if 'version_type' not in metadata:
                        metadata['version_type'] = 'Reviewer'
                else:
                    version_data['__metadata__'] = {
                        'reviewer_id': user_id,
                        'review_date': timezone.now().strftime('%Y-%m-%d %H:%M:%S'),
                        'version_type': 'Reviewer'
                    }
        
        # Print the structure we're returning for debugging
        print("DEBUG: JSON structure being returned:")
        try:
            print(json.dumps(version_data, indent=2, default=str))
        except Exception as e:
            print(f"DEBUG: Error printing JSON: {str(e)}")
        
        # Return the data from the audit version
        return Response({
            'audit_id': audit_id,
            'version': latest_audit_version,
            'data': version_data,
            'message': f'Loaded audit version {latest_audit_version} for review'
        }, status=status.HTTP_200_OK)
        
    except Exception as e:
        print(f"ERROR in load_review_data: {str(e)}")
        import traceback
        traceback.print_exc()
        return Response({'error': str(e)}, status=status.HTTP_400_BAD_REQUEST) 


@api_view(['POST'])
def update_audit_version_review_data(request, audit_id, version_id):
    """
    Update an existing audit version with new review data
    - Only updates accept_reject and comments fields
    - Preserves all other data
    - Prints the updated JSON before saving
    """
    try:
        print(f"DEBUG: update_audit_version_review_data called for audit_id: {audit_id}, version_id: {version_id}")
        print(f"DEBUG: Request data: {request.data}")
        
        # Check if the audit and version exist
        with connection.cursor() as cursor:
            cursor.execute("""
                SELECT ExtractedInfo 
                FROM audit_version 
                WHERE AuditId = %s AND Version = %s
            """, [audit_id, version_id])
            
            version_row = cursor.fetchone()
            if not version_row:
                return Response({
                    'error': f'Version {version_id} not found for audit {audit_id}'
                }, status=status.HTTP_404_NOT_FOUND)
            
            # Parse the existing JSON data
            if isinstance(version_row[0], dict):
                existing_data = version_row[0]
            else:
                existing_data = json.loads(version_row[0])
            
            print(f"DEBUG: Loaded existing data for version {version_id}")
        
        # Get user ID from session or request
        user_id = request.session.get('user_id', 1020)  # Default to reviewer ID if not found
        
        # Get the updated review data from the request
        updated_reviews = request.data
        
        # Create a deep copy of the existing data to avoid modifying it directly
        updated_data = json.loads(json.dumps(existing_data))
        
        # Track changes for logging
        changes = []
        
        # Update specific fields for each compliance ID in the request
        for compliance_id, review_data in updated_reviews.items():
            # Skip metadata and special fields
            if compliance_id in ['__metadata__', 'overall_comments'] or not isinstance(review_data, dict):
                continue    
            # Skip invalid compliance IDs
            if compliance_id not in updated_data:
                print(f"DEBUG: Compliance ID {compliance_id} not found in existing data, skipping")
                continue
            # Update specific fields only
            if 'accept_reject' in review_data:
                # Store old value for logging
                old_value = updated_data[compliance_id].get('accept_reject', '0')
                new_value = review_data['accept_reject']
                
                if old_value != new_value:
                    updated_data[compliance_id]['accept_reject'] = new_value
                    changes.append(f"Compliance {compliance_id}: accept_reject {old_value} -> {new_value}")
            
            if 'comments' in review_data:
                # Store old value for logging
                old_value = updated_data[compliance_id].get('comments', '')
                new_value = review_data['comments']
                
                if old_value != new_value:
                    updated_data[compliance_id]['comments'] = new_value
                    changes.append(f"Compliance {compliance_id}: comments updated")
        
        # Update overall_comments if provided
        if 'overall_comments' in updated_reviews:
            old_comments = updated_data.get('overall_comments', '')
            new_comments = updated_reviews['overall_comments']
            
            if old_comments != new_comments:
                updated_data['overall_comments'] = new_comments
                changes.append("Overall comments updated")
        
        # Update metadata if it exists
        if '__metadata__' in updated_data:
            updated_data['__metadata__']['review_date'] = timezone.now().strftime('%Y-%m-%d %H:%M:%S')
            updated_data['__metadata__']['reviewer_id'] = user_id
        else:
            # Create metadata if it doesn't exist
            updated_data['__metadata__'] = {
                'review_date': timezone.now().strftime('%Y-%m-%d %H:%M:%S'),
                'reviewer_id': user_id,
                'version_type': 'Reviewer'
            }
        
        # Print the updated JSON before saving
        print("DEBUG: Updated JSON structure before saving:")
        formatted_json = json.dumps(updated_data, indent=2)
        print(formatted_json)
        
        # If no changes were made, return without updating
        if not changes:
            return Response({
                'message': 'No changes detected, skipping update',
                'audit_id': audit_id,
                'version_id': version_id
            }, status=status.HTTP_200_OK)
        
        # Save the updated JSON back to the database
        with connection.cursor() as cursor:
            cursor.execute("""
                UPDATE audit_version
                SET ExtractedInfo = %s,
                    Date = %s
                WHERE AuditId = %s AND Version = %s
            """, [json.dumps(updated_data), timezone.now(), audit_id, version_id])
            
            print(f"DEBUG: Updated audit_version with {len(changes)} changes")
        
        return Response({
            'message': 'Audit version updated successfully',
            'audit_id': audit_id,
            'version_id': version_id,
            'changes_count': len(changes),
            'changes': changes
        }, status=status.HTTP_200_OK)
        
    except Exception as e:
        print(f"ERROR in update_audit_version_review_data: {str(e)}")
        return Response({'error': str(e)}, status=status.HTTP_400_BAD_REQUEST)

 
# @api_view(['POST'])
# def update_review_status(request, audit_id):
#     """
#     Update the review status of a specific audit
#     """
#     print("--------------views.py update_review_status---------------------------------------")
#     try:
#         print(f"DEBUG: update_review_status called for audit_id: {audit_id}")
#         print(f"DEBUG: Request data: {request.data}")
        
#         # Validate input
#         if 'review_status' not in request.data:
#             return Response({'error': 'review_status field is required'}, status=status.HTTP_400_BAD_REQUEST)
        
#         # Get the requested status string
#         new_status_str = request.data['review_status']
#         valid_statuses = ['Yet to Start', 'In Review', 'Accept', 'Reject']
#         if new_status_str not in valid_statuses:
#             return Response({'error': f'Invalid status. Must be one of: {", ".join(valid_statuses)}'}, 
#                           status=status.HTTP_400_BAD_REQUEST)
                          
#         # Check if we received compliance review data in the request
#         has_rejected = False
#         all_accepted = True
#         compliance_reviews = request.data.get('compliance_reviews', [])
        
#         # Check if any compliance has a 'Reject' status
#         for review in compliance_reviews:
#             if review.get('review_status') == 'Reject':
#                 has_rejected = True
#                 all_accepted = False
#                 break
#             elif review.get('review_status') != 'Accept':
#                 all_accepted = False
        
#         print(f"DEBUG: Review status check - has_rejected: {has_rejected}, all_accepted: {all_accepted}")
        
#         # Map string status to integer
#         status_map = {
#             'Yet to Start': 0,
#             'In Review': 1,
#             'Accept': 2,
#             'Reject': 3
#         }
#         new_status_int = status_map.get(new_status_str)
        
#         # Find and update the audit
#         try:
#             audit = Audit.objects.get(AuditId=audit_id)
#         except Audit.DoesNotExist:
#             return Response({'error': 'Audit not found'}, status=status.HTTP_404_NOT_FOUND)
        
#         # Check if audit is in the correct state for review
#         if audit.Status != 'Under review':
#             return Response({
#                 'error': 'Cannot update review status when audit is not under review',
#                 'current_audit_status': audit.Status
#             }, status=status.HTTP_400_BAD_REQUEST)

#         # Store the old status before updating
#         old_status_int = audit.ReviewStatus
#         # Map integer to string for logging
#         status_reverse_map = {0: 'Yet to Start', 1: 'In Review', 2: 'Accept', 3: 'Reject'}
#         old_status_str = status_reverse_map.get(old_status_int, 'Unknown') if old_status_int is not None else 'None'
#         print(f"DEBUG: Changing review status from '{old_status_str}' ({old_status_int}) to '{new_status_str}' ({new_status_int})")
        
#         # Update the review status with integer value
#         audit.ReviewStatus = new_status_int
        
#         # Add review comments if provided
#         review_comments = None
#         if 'review_comments' in request.data:
#             review_comments = request.data['review_comments']
#             audit.ReviewComments = review_comments
#             print(f"DEBUG: Setting review comments: {review_comments}")
        
#         # Set review date if the field exists
#         current_time = timezone.now()
#         review_date_exists = hasattr(audit, 'ReviewDate')
        
#         if review_date_exists:
#             audit.ReviewDate = current_time
#             print(f"DEBUG: Setting ReviewDate to {current_time}")
#         else:
#             print("DEBUG: ReviewDate field doesn't exist, skipping")
            
#             # Try to add the column if it doesn't exist
#             try:
#                 with connection.cursor() as cursor:
#                     cursor.execute("""
#                         SELECT COUNT(*) 
#                         FROM information_schema.COLUMNS 
#                         WHERE TABLE_NAME = 'audit' 
#                         AND COLUMN_NAME = 'ReviewDate'
#                     """)
#                     column_exists = cursor.fetchone()[0] > 0
                    
#                     if not column_exists:
#                         print("DEBUG: Attempting to add ReviewDate column to audit table")
#                         cursor.execute("""
#                             ALTER TABLE audit
#                             ADD COLUMN ReviewDate DATETIME NULL
#                         """)
#                         print("DEBUG: ReviewDate column added successfully")
                        
#                         # Now we can set the field
#                         audit.ReviewDate = current_time
#             except Exception as e:
#                 print(f"DEBUG: Error handling ReviewDate column: {str(e)}")
        
#         # Save the audit object with the updated review status
#         audit.save()
#         print(f"DEBUG: Audit review status updated in database to {new_status_str}")
        
#         # Get user ID from session or request
#         user_id = request.session.get('user_id')
#         if not user_id and hasattr(audit, 'reviewer_id'):
#             user_id = audit.reviewer_id
        
#         # Create a review version with all compliance review data 
#         if user_id:
#             # Extract compliance reviews from request if available
#             compliance_reviews = request.data.get('compliance_reviews', [])
            
#             # Get the next version number for review versions
#             review_version = "R1"  # Default
            
#             with connection.cursor() as cursor:
#                 cursor.execute("""
#                     SELECT Version FROM audit_version 
#                     WHERE AuditId = %s AND Version LIKE 'R%'
#                     ORDER BY Version DESC
#                     LIMIT 1
#                 """, [audit_id])
                
#                 existing_version = cursor.fetchone()
#                 if existing_version:
#                     current_version = existing_version[0]
#                     # Extract number part and increment it
#                     version_num = int(current_version[1:])
#                     review_version = f"R{version_num + 1}"
#                     print(f"DEBUG: Incrementing from existing review version {current_version} to {review_version}")
            
#             # Create review version with the overall review status and comments
#             # Always create a version, even if compliance_reviews is empty
#             review_version_result = create_review_version(
#                 audit_id, 
#                 user_id, 
#                 compliance_reviews, 
#                 review_comments,
#                 review_version
#             )
#             print(f"DEBUG: Review version result: {review_version_result}")
        
#         # If the review is rejected, update the audit status back to "Work In Progress"
#         if new_status_str == 'Reject' or has_rejected:
#             print(f"DEBUG: Audit {audit_id} rejected, changing status back to 'Work In Progress'")
#             audit.Status = 'Work In Progress'
#             audit.ReviewStatus = 0  # Set review status to 0
#             audit.save()
            
#             # Get the latest review version and JSON data
#             latest_version = None
#             should_reject = True
            
#             try:
#                     with connection.cursor() as cursor:
#                     # Get the latest R version
#                         cursor.execute("""
#                         SELECT Version, ExtractedInfo FROM audit_version 
#                             WHERE AuditId = %s AND Version LIKE 'R%'
#                             ORDER BY Version DESC
#                             LIMIT 1
#                         """, [audit_id])
                    
#                     version_row = cursor.fetchone()
#                     if version_row:
#                         latest_version = version_row[0]
#                         # Parse JSON data
#                         if isinstance(version_row[1], dict):
#                             version_data = version_row[1]
#                         else:
#                             version_data = json.loads(version_row[1])
                        
#                         print(f"DEBUG: Retrieved JSON data from latest review version {latest_version}")
                        
#                         # Update metadata to indicate rejection
#                         metadata = version_data.get('__metadata__', {})
#                         if not isinstance(metadata, dict):
#                             metadata = {}
                            
#                         metadata['overall_status'] = 'Reject'
#                         metadata['ApprovedRejected'] = 'Rejected'
#                         version_data['__metadata__'] = metadata
                        
#                         # Ensure we keep the overall comments
#                         if review_comments and 'overall_comments' not in version_data:
#                             version_data['overall_comments'] = review_comments
                        
#                         # Update the version in the database with the modified data
#                         cursor.execute("""
#                             UPDATE audit_version
#                             SET ExtractedInfo = %s,
#                                 ApprovedRejected = 'Rejected'
#                             WHERE AuditId = %s AND Version = %s
#                         """, [json.dumps(version_data), audit_id, latest_version])
                        
#                         print(f"DEBUG: Updated version {latest_version} with Rejected status and metadata")
#             except Exception as e:
#                 print(f"ERROR updating version with rejection status: {str(e)}")
#                 # Don't fail the whole request if this update fails
#         elif new_status_str == 'Accept' or (all_accepted and compliance_reviews):
#             print(f"DEBUG: Audit {audit_id} accepted, updating version data and changing status to Completed")
            
#             # Update audit status to Completed
#             audit.Status = 'Completed'
#             audit.save()
            
#             try:
#                 with connection.cursor() as cursor:
#                     # Get the latest R version
#                     cursor.execute("""
#                         SELECT Version, ExtractedInfo FROM audit_version 
#                         WHERE AuditId = %s AND Version LIKE 'R%'
#                         ORDER BY Version DESC
#                         LIMIT 1
#                     """, [audit_id])
                    
#                     version_row = cursor.fetchone()
#                     if version_row:
#                         latest_version = version_row[0]
#                         # Parse JSON data
#                         if isinstance(version_row[1], dict):
#                             version_data = version_row[1]
#                         else:
#                             version_data = json.loads(version_row[1])
                        
#                         print(f"DEBUG: Retrieved JSON data from latest review version {latest_version}")
                        
#                         # Update metadata to indicate acceptance
#                         metadata = version_data.get('__metadata__', {})
#                         if not isinstance(metadata, dict):
#                             metadata = {}
                            
#                         metadata['overall_status'] = 'Accept'
#                         metadata['ApprovedRejected'] = 'Approved'
#                         version_data['__metadata__'] = metadata
                        
#                         # Ensure we keep the overall comments
#                         if review_comments and 'overall_comments' not in version_data:
#                             version_data['overall_comments'] = review_comments
                        
#                         # Update the version in the database with the modified data
#                         cursor.execute("""
#                             UPDATE audit_version
#                             SET ExtractedInfo = %s,
#                                 ApprovedRejected = 'Approved'
#                             WHERE AuditId = %s AND Version = %s
#                         """, [json.dumps(version_data), audit_id, latest_version])
                        
#                         # Save the approved R version JSON data into the audit_findings table
#                         print(f"DEBUG: Updating audit_findings table with approved R version data")
                        
#                         # Process each compliance in the version data
#                         for compliance_id, compliance_data in version_data.items():
#                             # Skip metadata and overall_comments
#                             if compliance_id in ['__metadata__', 'overall_comments']:
#                                 continue
                                
#                             if not isinstance(compliance_data, dict):
#                                 continue
                                
#                             try:
#                                 # Map review status to ReviewRejected value
#                                 review_status = compliance_data.get('review_status', 'In Review')
#                                 review_comments = compliance_data.get('review_comments', '')
                                
#                                 # Get data from the version to update audit_findings
#                                 recommendation = compliance_data.get('recommendation', '')
#                                 details_of_finding = compliance_data.get('details_of_finding', '')
#                                 impact = compliance_data.get('impact', '')
#                                 how_to_verify = compliance_data.get('how_to_verify', '')
#                                 comments = compliance_data.get('comments', '')
                                
#                                 # Get the compliance_status and map it to Check value
#                                 compliance_status = compliance_data.get('compliance_status', '')
#                                 check_value = '0'  # Default: Not Started (Not Compliant)
                                
#                                 if compliance_status == 'Fully Compliant':
#                                     check_value = '2'  # Completed
#                                 elif compliance_status == 'Partially Compliant':
#                                     check_value = '1'  # In Progress
#                                 elif compliance_status == 'Not Applicable':
#                                     check_value = '3'  # Not Applicable
                                
#                                 # Get criticality for MajorMinor
#                                 criticality = compliance_data.get('criticality', '')
#                                 # Update the audit_findings table with the approved data
#                                 cursor.execute("""
#                                     UPDATE audit_findings
#                                     SET ReviewStatus = %s,
#                                         ReviewComments = %s,
#                                         ReviewRejected = %s,
#                                         ReviewDate = NOW(),
#                                         Recommendation = %s,
#                                         DetailsOfFinding = %s,
#                                         Impact = %s,
#                                         HowToVerify = %s,
#                                         Comments = %s,
#                                         `Check` = %s,
#                                         MajorMinor = %s,
#                                         CheckedDate = NOW()
#                                     WHERE AuditId = %s AND ComplianceId = %s
#                                 """, [
#                                     review_status,
#                                     review_comments,
#                                     0,  # Not rejected as this is approved
#                                     recommendation,
#                                     details_of_finding,
#                                     impact,
#                                     how_to_verify,
#                                     comments,
#                                     check_value,
#                                     criticality,
#                                     audit_id,
#                                     compliance_id
#                                 ])
                                
#                             except Exception as e:
#                                 print(f"ERROR updating audit_findings for compliance {compliance_id}: {str(e)}")
#                                 # Continue with next compliance
                                
#                         print(f"DEBUG: Successfully updated audit_findings table with approved R version data")
                        
#                         print(f"DEBUG: Updated version {latest_version} with Approved status and metadata")
#             except Exception as e:
#                 print(f"ERROR updating version with approval status: {str(e)}")
#                 # Don't fail the whole request if this update fails
        
#         # For case where we are only processing compliance_reviews without setting overall status
#         if compliance_reviews and new_status_str not in ['Accept', 'Reject']:
#             # Determine ApprovedRejected status based on individual compliance reviews
#             try:
#                 with connection.cursor() as cursor:
#                     # Get the latest R version
#                     cursor.execute("""
#                         SELECT Version, ExtractedInfo FROM audit_version 
#                         WHERE AuditId = %s AND Version LIKE 'R%'
#                         ORDER BY Version DESC
#                         LIMIT 1
#                     """, [audit_id])
                    
#                     version_row = cursor.fetchone()
#                     if version_row:
#                         latest_version = version_row[0]
#                         # Parse JSON data
#                         if isinstance(version_row[1], dict):
#                             version_data = version_row[1]
#                         else:
#                             version_data = json.loads(version_row[1])
                        
#                         # Update ApprovedRejected based on compliance reviews
#                         approvedRejected = 'Approved' if all_accepted else 'Rejected'
#                         print(f"DEBUG: Setting ApprovedRejected to {approvedRejected} based on compliance reviews")
                        
#                         # Update metadata
#                         metadata = version_data.get('__metadata__', {})
#                         if not isinstance(metadata, dict):
#                             metadata = {}
                            
#                         metadata['ApprovedRejected'] = approvedRejected
#                         version_data['__metadata__'] = metadata
                        
#                         # Update the version in the database
#                         cursor.execute("""
#                             UPDATE audit_version
#                             SET ExtractedInfo = %s,
#                                 ApprovedRejected = %s
#                             WHERE AuditId = %s AND Version = %s
#                         """, [json.dumps(version_data), approvedRejected, audit_id, latest_version])
                        
#                         print(f"DEBUG: Updated version {latest_version} ApprovedRejected to {approvedRejected}")
#             except Exception as e:
#                 print(f"ERROR updating ApprovedRejected status: {str(e)}")
#                 # Don't fail the whole request if this update fails
        
#         # Send notifications based on the review decision
#         try:
#             # Initialize notification service
#             notification_service = NotificationService()
            
#             # Get audit details using direct SQL to avoid column issues
#             with connection.cursor() as cursor:
#                 cursor.execute("""
#                     SELECT a.AuditId, a.Title, a.Auditor, a.Assignee, f.FrameworkName,
#                            p.PolicyName, sp.SubPolicyName
#                     FROM audit a
#                     LEFT JOIN frameworks f ON a.FrameworkId = f.FrameworkId
#                     LEFT JOIN policies p ON a.PolicyId = p.PolicyId
#                     LEFT JOIN subpolicies sp ON a.SubPolicyId = sp.SubPolicyId
#                     WHERE a.AuditId = %s
#                 """, [audit_id])
#                 audit_data = cursor.fetchone()
                
#                 if audit_data:
#                     audit_title = audit_data[1] or f"Audit #{audit_id}"
#                     auditor_id = audit_data[2]
#                     assignee_id = audit_data[3]
#                     framework_name = audit_data[4] or "Unknown Framework"
#                     policy_name = audit_data[5] or ""
#                     subpolicy_name = audit_data[6] or ""
                    
#                     # Get user emails
#                     cursor.execute("SELECT UserId, UserName, Email FROM users WHERE UserId IN (%s, %s)", 
#                                   [auditor_id, assignee_id])
#                     user_data = cursor.fetchall()
                    
#                     user_info = {}
#                     for user in user_data:
#                         user_info[user[0]] = {
#                             'name': user[1] or f"User {user[0]}",
#                             'email': user[2]
#                         }
                    
#                     # Get reviewer info (current user)
#                     reviewer_id = request.session.get('user_id')
#                     reviewer_name = "System"
#                     if reviewer_id:
#                         cursor.execute("SELECT UserName FROM users WHERE UserId = %s", [reviewer_id])
#                         reviewer_row = cursor.fetchone()
#                         if reviewer_row:
#                             reviewer_name = reviewer_row[0] or f"User {reviewer_id}"
                    
#                     # Send notification based on review status
#                     if new_status_str == 'Accept' or (all_accepted and compliance_reviews):
#                         # Send Approval notification
#                         if auditor_id in user_info and user_info[auditor_id]['email']:
#                             notification_data = {
#                                 'notification_type': 'auditReviewed',
#                                 'email': user_info[auditor_id]['email'],
#                                 'email_type': 'gmail',
#                                 'template_data': [
#                                     user_info[auditor_id]['name'],
#                                     audit_title,
#                                     'Approved',
#                                     reviewer_name,
#                                     review_comments or 'Audit has been approved.'
#                                 ]
#                             }
#                             notification_service.send_multi_channel_notification(notification_data)
#                             print(f"DEBUG: Sent 'audit approved' notification to auditor {user_info[auditor_id]['email']}")
                        
#                         # If assignee is different from auditor, notify them too
#                         if assignee_id != auditor_id and assignee_id in user_info and user_info[assignee_id]['email']:
#                             notification_data = {
#                                 'notification_type': 'auditReviewed',
#                                 'email': user_info[assignee_id]['email'],
#                                 'email_type': 'gmail',
#                                 'template_data': [
#                                     user_info[assignee_id]['name'],
#                                     audit_title,
#                                     'Approved',
#                                     reviewer_name,
#                                     review_comments or 'Audit has been approved.'
#                                 ]
#                             }
#                             notification_service.send_multi_channel_notification(notification_data)
#                             print(f"DEBUG: Sent 'audit approved' notification to assignee {user_info[assignee_id]['email']}")
                            
#                     elif new_status_str == 'Reject' or has_rejected:
#                         # Send Rejection notification
#                         if auditor_id in user_info and user_info[auditor_id]['email']:
#                             notification_data = {
#                                 'notification_type': 'auditReviewed',
#                                 'email': user_info[auditor_id]['email'],
#                                 'email_type': 'gmail',
#                                 'template_data': [
#                                     user_info[auditor_id]['name'],
#                                     audit_title,
#                                     'Rejected',
#                                     reviewer_name,
#                                     review_comments or 'Audit requires revisions. Please review and update.'
#                                 ]
#                             }
#                             notification_service.send_multi_channel_notification(notification_data)
#                             print(f"DEBUG: Sent 'audit rejected' notification to auditor {user_info[auditor_id]['email']}")
                        
#                         # If assignee is different from auditor, notify them too
#                         if assignee_id != auditor_id and assignee_id in user_info and user_info[assignee_id]['email']:
#                             notification_data = {
#                                 'notification_type': 'auditReviewed',
#                                 'email': user_info[assignee_id]['email'],
#                                 'email_type': 'gmail',
#                                 'template_data': [
#                                     user_info[assignee_id]['name'],
#                                     audit_title,
#                                     'Rejected',
#                                     reviewer_name,
#                                     review_comments or 'Audit requires revisions. Please review and update.'
#                                 ]
#                             }
#                             notification_service.send_multi_channel_notification(notification_data)
#                             print(f"DEBUG: Sent 'audit rejected' notification to assignee {user_info[assignee_id]['email']}")
                    
#                     elif new_status_str == 'In Review':
#                         # Send notification that review has started
#                         if auditor_id in user_info and user_info[auditor_id]['email']:
#                             notification_data = {
#                                 'notification_type': 'policyStatusChange',
#                                 'email': user_info[auditor_id]['email'],
#                                 'email_type': 'gmail',
#                                 'template_data': [
#                                     user_info[auditor_id]['name'],
#                                     audit_title,
#                                     'In Review',
#                                     reviewer_name,
#                                     timezone.now().strftime('%Y-%m-%d %H:%M:%S')
#                                 ]
#                             }
#                             notification_service.send_multi_channel_notification(notification_data)
#                             print(f"DEBUG: Sent 'review started' notification to auditor {user_info[auditor_id]['email']}")
        
#         except Exception as e:
#             print(f"ERROR: Failed to send notifications: {str(e)}")
#             # Don't fail the whole operation if notifications fail
        
#         # Return the updated status information
#         return Response({
#             'success': True,
#             'audit_id': audit_id,
#             'review_status': new_status_str,
#             'review_status_int': new_status_int,
#             'audit_status': audit.Status,
#             'review_comments': review_comments,
#             'has_rejected': has_rejected,
#             'all_accepted': all_accepted
#         }, status=status.HTTP_200_OK)
        
#     except Exception as e:
#         print(f"ERROR in update_review_status: {str(e)}")
#         import traceback
#         traceback.print_exc()
#         return Response({'error': str(e)}, status=status.HTTP_400_BAD_REQUEST)



@api_view(['POST'])
def save_review_json(request, audit_id):
    """
    Save review JSON directly to audit_version table
    - Takes a JSON structure with compliance IDs as keys
    - Updates only accept_reject and comments fields
    - Preserves all other data
    - Updates latest version instead of creating a new one
    """
    try:
        print(f"DEBUG: save_review_json called for audit_id: {audit_id}")
        
        # Get the JSON data from the request
        review_data = request.data
        print("DEBUG: Review data received:")
        print(json.dumps(review_data, indent=2, default=str))
        
        # Find the audit
        try:
            audit = Audit.objects.get(AuditId=audit_id)
        except Audit.DoesNotExist:
            return Response({'error': 'Audit not found'}, status=status.HTTP_404_NOT_FOUND)
        
        # Get user ID from session or request
        user_id = request.session.get('user_id')
        if not user_id and hasattr(audit, 'reviewer_id'):
            user_id = audit.reviewer_id
        if not user_id:
            user_id = 1020  # Default reviewer ID if not found
        
        # First check if there's an existing R-version to update
        existing_version = None
        with connection.cursor() as cursor:
            try:
                # First look for any existing R-version (reviewer)
                cursor.execute(
                    "SELECT Version FROM audit_version WHERE AuditId = %s AND Version LIKE %s ORDER BY Version DESC LIMIT 1", 
                    [audit_id, "R%"]
                )
                version_row = cursor.fetchone()
                if version_row:
                    existing_version = version_row[0]
                    print(f"DEBUG: Found existing reviewer version {existing_version} to update")
                else:
                    # If no R-version, check for A-version (auditor)
                    cursor.execute(
                        "SELECT Version FROM audit_version WHERE AuditId = %s AND Version LIKE %s ORDER BY Version DESC LIMIT 1", 
                        [audit_id, "A%"]
                    )
                    version_row = cursor.fetchone()
                    if version_row:
                        # Use A1 prefix but convert to R1 for the first reviewer version
                        existing_version = "R1"
                        print(f"DEBUG: Found auditor version {version_row[0]}, will create first reviewer version R1")
                    else:
                        # No versions at all, create new R1
                        existing_version = "R1"
                        print(f"DEBUG: No existing versions found, will create new R1")
            except Exception as e:
                print(f"DEBUG: Error checking for existing versions: {str(e)}")
                existing_version = "R1"  # Default if error occurs
        # Update the metadata in the JSON
        if "__metadata__" in review_data:
            review_data["__metadata__"]["review_date"] = timezone.now().strftime('%Y-%m-%d %H:%M:%S')
            review_data["__metadata__"]["reviewer_id"] = user_id
            review_data["__metadata__"]["version_type"] = "Reviewer"
        else:
            review_data["__metadata__"] = {
                "review_date": timezone.now().strftime('%Y-%m-%d %H:%M:%S'),
                "reviewer_id": user_id,
                "version_type": "Reviewer"
            }
        # Serialize the JSON data
        json_data = json.dumps(review_data)
        print("DEBUG: JSON data to be saved:")
        print(json.dumps(review_data, indent=2, default=str))
        
        # Check if version already exists, then update or create
        with connection.cursor() as cursor:
            try:
                cursor.execute(
                    "SELECT COUNT(*) FROM audit_version WHERE AuditId = %s AND Version = %s",
                    [audit_id, existing_version]
                )
                version_exists = cursor.fetchone()[0] > 0
                
                if version_exists:
                    # Update existing version
                    cursor.execute(
                        "UPDATE audit_version SET ExtractedInfo = %s, UserId = %s, Date = %s WHERE AuditId = %s AND Version = %s",
                        [json_data, user_id, timezone.now(), audit_id, existing_version]
                    )
                    print(f"DEBUG: Updated existing review version {existing_version} for audit {audit_id}")
                else:
                    # Insert new version
                    print(timezone.now(),"------------------------------------------------------------------------------")
                    cursor.execute(
                        "INSERT INTO audit_version (AuditId, Version, ExtractedInfo, UserId, Date) VALUES (%s, %s, %s, %s, %s)",
                        [audit_id, existing_version, json_data, user_id, timezone.localtime(current_time_utc)]
                    )
                    print(f"DEBUG: Created new review version {existing_version} for audit {audit_id}")
            except Exception as e:
                print(f"ERROR saving to audit_version table: {str(e)}")
                return Response({'error': f'Database error: {str(e)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        
        # Update audit review status if needed
        try:
            # Check if we need to update the ReviewStatus
            all_accepted = True
            has_rejected = False
            
            # Count review statuses
            for key, value in review_data.items():
                if key in ['__metadata__', 'overall_comments']:
                    continue
                if isinstance(value, dict) and 'accept_reject' in value:
                    accept_reject = value['accept_reject']
                    if accept_reject == '2':  # Rejected
                        has_rejected = True
                        all_accepted = False
                    elif accept_reject != '1':  # Not accepted
                        all_accepted = False
            # Set the appropriate ReviewStatus value
            review_status_int = 1  # Default to "In Review" (1)
            if has_rejected:
                review_status_int = 3  # "Reject" (3)
            elif all_accepted:
                review_status_int = 2  # "Accept" (2)
            # Update the audit
            if hasattr(audit, 'ReviewStatus'):
                # Only update if not already in a final state
                if audit.ReviewStatus not in [2, 3]:  # Not Accept or Reject
                    audit.ReviewStatus = review_status_int
                    if hasattr(audit, 'ReviewDate'):
                        audit.ReviewDate = timezone.now()
                    # If any reviews are rejected, change audit status to Work In Progress
                    if has_rejected and audit.Status == 'Under review':
                        print(f"DEBUG: Rejection detected in save_review_json, changing audit status from 'Under review' to 'Work In Progress'")
                        audit.Status = 'Work In Progress'
                    
                    audit.save()
            else:
                # Try to update with SQL
                with connection.cursor() as cursor:
                    cursor.execute(
                        "SELECT COUNT(*) FROM information_schema.COLUMNS WHERE TABLE_NAME = 'audit' AND COLUMN_NAME = 'ReviewStatus'",
                        []
                    )
                    field_exists = cursor.fetchone()[0] > 0
                    
                    if field_exists:
                        cursor.execute(
                            "UPDATE audit SET ReviewStatus = %s, ReviewDate = %s WHERE AuditId = %s",
                            [review_status_int, timezone.now(), audit_id]
                        )
                        
                        # If any reviews are rejected, change audit status to Work In Progress
                        if has_rejected:
                            cursor.execute(
                                "SELECT Status FROM audit WHERE AuditId = %s",
                                [audit_id]
                            )
                            current_audit_status = cursor.fetchone()[0]
                            
                            if current_audit_status == 'Under review':
                                cursor.execute(
                                    "UPDATE audit SET Status = %s WHERE AuditId = %s",
                                    ['Work In Progress', audit_id]
                                )
                                print(f"DEBUG: SQL - Changed audit status from 'Under review' to 'Work In Progress' due to rejection")
        except Exception as e:
            print(f"WARNING: Failed to update audit ReviewStatus: {str(e)}")
            # Continue even if this part fails
        
        return Response({
            'message': 'Review data saved successfully',
            'audit_id': audit_id,
            'version': existing_version,
            'updated': True,
            'saved_at': timezone.now().strftime('%Y-%m-%d %H:%M:%S')
        }, status=status.HTTP_200_OK)
        
    except Exception as e:
        print(f"ERROR in save_review_json: {str(e)}")
        import traceback
        traceback.print_exc()
        return Response({'error': f'Error saving review data: {str(e)}'}, status=status.HTTP_400_BAD_REQUEST)

@api_view(['GET'])
def update_audit_version_table(request):
    """
    Update audit_version table structure to match required schema
    Columns:
    - AuditId (int PK)
    - Version (varchar(45) PK)
    - ExtractedInfo (json)
    - UserId (int)
    - ApproverId (int)
    - ApprovedRejected (varchar(45))
    - Date (datetime)
    """
    try:
        print("DEBUG: update_audit_version_table called")
        
        # Check if audit_version table exists
        with connection.cursor() as cursor:
            try:
                # Check if table exists
                cursor.execute(
                    "SELECT COUNT(*) FROM information_schema.tables WHERE table_schema = DATABASE() AND table_name = 'audit_version'",
                    []
                )
                table_exists = cursor.fetchone()[0] > 0
                
                if not table_exists:
                    # Create table with proper schema
                    print("DEBUG: audit_version table doesn't exist, creating it")
                    cursor.execute("""
                        CREATE TABLE audit_version (
                            AuditId INT NOT NULL,
                            Version VARCHAR(45) NOT NULL,
                            ExtractedInfo JSON NULL,
                            UserId INT NULL,
                            ApproverId INT NULL,
                            ApprovedRejected VARCHAR(45) NULL,
                            Date DATETIME NULL,
                            PRIMARY KEY (AuditId, Version)
                        )
                    """)
                    return Response({
                        'message': 'audit_version table created successfully',
                        'action': 'created'
                    }, status=status.HTTP_200_OK)
                else:
                    # Table exists, check if it has the correct structure
                    print("DEBUG: audit_version table exists, checking structure")
                    
                    # Get current columns
                    cursor.execute("""
                        SELECT COLUMN_NAME, DATA_TYPE, IS_NULLABLE, COLUMN_KEY
                        FROM information_schema.COLUMNS 
                        WHERE TABLE_SCHEMA = DATABASE() AND TABLE_NAME = 'audit_version'
                        ORDER BY ORDINAL_POSITION
                    """)
                    
                    columns = cursor.fetchall()
                    column_names = [col[0] for col in columns]
                    
                    # Check each required column and add if missing
                    corrections = []
                    expected_columns = [
                        ('AuditId', 'int', 'NO', 'PRI'),
                        ('Version', 'varchar', 'NO', 'PRI'),
                        ('ExtractedInfo', 'json', 'YES', ''),
                        ('UserId', 'int', 'YES', ''),
                        ('ApproverId', 'int', 'YES', ''),
                        ('ApprovedRejected', 'varchar', 'YES', ''),
                        ('Date', 'datetime', 'YES', '')
                    ]
                    
                    for expected_col in expected_columns:
                        col_name, col_type, nullable, key = expected_col
                        
                        if col_name not in column_names:
                            # Column is missing, add it
                            null_str = "NULL" if nullable == "YES" else "NOT NULL"
                            type_str = f"VARCHAR(45)" if col_type == "varchar" else col_type.upper()
                            
                            try:
                                cursor.execute(f"""
                                    ALTER TABLE audit_version
                                    ADD COLUMN {col_name} {type_str} {null_str}
                                """)
                                corrections.append(f"Added missing column {col_name}")
                            except Exception as e:
                                corrections.append(f"Error adding column {col_name}: {str(e)}")
                    
                    # Check primary key
                    cursor.execute("""
                        SELECT COUNT(*) FROM information_schema.TABLE_CONSTRAINTS 
                        WHERE TABLE_SCHEMA = DATABASE() 
                        AND TABLE_NAME = 'audit_version'
                        AND CONSTRAINT_TYPE = 'PRIMARY KEY'
                    """)
                    
                    has_pk = cursor.fetchone()[0] > 0
                    
                    if not has_pk:
                        try:
                            cursor.execute("""
                                ALTER TABLE audit_version
                                ADD PRIMARY KEY (AuditId, Version)
                            """)
                            corrections.append("Added missing primary key on (AuditId, Version)")
                        except Exception as e:
                            corrections.append(f"Error adding primary key: {str(e)}")
                    
                    if corrections:
                        return Response({
                            'message': 'audit_version table updated successfully',
                            'changes': corrections
                        }, status=status.HTTP_200_OK)
                    else:
                        return Response({
                            'message': 'audit_version table already has the correct structure',
                            'action': 'none'
                        }, status=status.HTTP_200_OK)
                    
            except Exception as e:
                print(f"ERROR checking/creating audit_version table: {str(e)}")
                import traceback
                traceback.print_exc()
                return Response({
                    'error': f'Error updating audit_version table: {str(e)}'
                }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
    
    except Exception as e:
        print(f"ERROR in update_audit_version_table: {str(e)}")
        import traceback
        traceback.print_exc()
        return Response({'error': str(e)}, status=status.HTTP_400_BAD_REQUEST)

@api_view(['GET'])
def load_latest_review_version(request, audit_id):
    """
    Load the latest review version JSON for an audit when a reviewer continues a task
    - Prioritizes loading the latest auditor version (A-prefix) for reviewers to see
    - Falls back to R-prefix (reviewer) versions if no A versions exist
    - Returns the JSON data for the review UI to restore the previous state
    """
    try:
        print(f"DEBUG: load_latest_review_version called for audit_id: {audit_id}")
        
        # Check if the audit exists
        try:
            audit = Audit.objects.get(AuditId=audit_id)
        except Audit.DoesNotExist:
            return Response({'error': 'Audit not found'}, status=status.HTTP_404_NOT_FOUND)
        
        # Check if audit is in the correct state for review
        if audit.Status != 'Under review':
            return Response({
                'error': 'Cannot load review data when audit is not under review',
                'current_audit_status': audit.Status
            }, status=status.HTTP_400_BAD_REQUEST)
        
        # Get user ID from session
        user_id = request.session.get('user_id')
        if not user_id and hasattr(audit, 'reviewer_id'):
            user_id = audit.reviewer_id
        if not user_id:
            user_id = 1020  # Default reviewer ID if not found
            
        print(f"DEBUG: Using reviewer ID: {user_id}")
        
        # Get the latest version from audit_version table (prioritize auditor versions)
        latest_review_data = None
        review_version = None
        
        with connection.cursor() as cursor:
            try:
                # First try to get the latest auditor version (A-prefix)
                cursor.execute(
                    "SELECT Version, ExtractedInfo, Date FROM audit_version WHERE AuditId = %s AND Version LIKE 'A%' ORDER BY Date DESC LIMIT 1",
                    [audit_id]
                )
                
                version_row = cursor.fetchone()
                if version_row:
                    review_version = version_row[0]
                    print(f"DEBUG: Found latest auditor version: {review_version} from {version_row[2]}")
                    
                    # Parse the JSON data
                    if isinstance(version_row[1], dict):
                        latest_review_data = version_row[1]
                    else:
                        try:
                            latest_review_data = json.loads(version_row[1]) if version_row[1] else {}
                        except Exception as e:
                            print(f"ERROR parsing JSON from latest auditor version: {str(e)}")
                            latest_review_data = {}
                        
                    print(f"DEBUG: Loaded auditor version with {len(latest_review_data) if latest_review_data else 0} findings")
                else:
                    # If no auditor version found, try to get the latest reviewer version (R-prefix)
                    cursor.execute(
                        "SELECT Version, ExtractedInfo, Date FROM audit_version WHERE AuditId = %s AND Version LIKE 'R%' ORDER BY Date DESC LIMIT 1",
                        [audit_id]
                    )
                    
                    version_row = cursor.fetchone()
                    if version_row:
                        review_version = version_row[0]
                        print(f"DEBUG: Found latest reviewer version: {review_version} from {version_row[2]}")
                        
                        # Parse the JSON data
                        if isinstance(version_row[1], dict):
                            latest_review_data = version_row[1]
                        else:
                            try:
                                latest_review_data = json.loads(version_row[1]) if version_row[1] else {}
                            except Exception as e:
                                print(f"ERROR parsing JSON from latest reviewer version: {str(e)}")
                                latest_review_data = {}
                            
                        print(f"DEBUG: Loaded reviewer version with {len(latest_review_data) if latest_review_data else 0} findings")


                # For auditor versions (A-prefix), ensure it has the reviewer fields
                if latest_review_data and review_version and review_version.startswith('A'):
                    # Process each finding to add accept_reject and comments fields if needed
                    for key, value in list(latest_review_data.items()):
                        # Skip metadata and overall_comments
                        if key in ['__metadata__', 'overall_comments']:
                            continue
                            
                        if isinstance(value, dict):
                            # Add required reviewer fields if they don't exist
                            if 'accept_reject' not in value:
                                value['accept_reject'] = "0"  # Default to "Not reviewed"
                            if 'comments' not in value:
                                value['comments'] = ""  # Empty comments
                            
                    # Update metadata for reviewer
                    if '__metadata__' in latest_review_data:
                        latest_review_data['__metadata__']['reviewer_id'] = user_id
                        latest_review_data['__metadata__']['review_date'] = timezone.now().strftime('%Y-%m-%d %H:%M:%S')
                        latest_review_data['__metadata__']['version_type'] = 'Reviewer'
                    else:
                        latest_review_data['__metadata__'] = {
                            'reviewer_id': user_id,
                            'review_date': timezone.now().strftime('%Y-%m-%d %H:%M:%S'),
                            'version_type': 'Reviewer'
                        }
                    
                    # Ensure overall_comments exists
                    if 'overall_comments' not in latest_review_data:
                        latest_review_data['overall_comments'] = ""
            except Exception as e:
                print(f"ERROR querying audit_version: {str(e)}")
                import traceback
                traceback.print_exc()
                return Response({'error': f'Database error: {str(e)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        
        if not latest_review_data:
            # If we still don't have data, return an error
            return Response({'error': 'No review data found and unable to create baseline'}, status=status.HTTP_404_NOT_FOUND)
        
        # Return the data for the review UI
        return Response({
            'audit_id': audit_id,
            'version': review_version,
            'review_data': latest_review_data,
            'message': f'Loaded review data from version {review_version} (latest)',
            'timestamp': timezone.now().strftime('%Y-%m-%d %H:%M:%S')
        }, status=status.HTTP_200_OK)
        
    except Exception as e:
        print(f"ERROR in load_latest_review_version: {str(e)}")
        import traceback
        traceback.print_exc()
        return Response({'error': str(e)}, status=status.HTTP_400_BAD_REQUEST)

def get_empty_review_structure(audit_id, user_id):
    """
    Create an empty review structure by fetching compliance items for an audit
    Used when no previous version exists
    """
    try:
        # Get compliance items for this audit
        findings_data = {}
        
        with connection.cursor() as cursor:
            cursor.execute("""
                SELECT 
                    af.ComplianceId,
                    c.ComplianceItemDescription,
                    af.`Check`,
                    af.Comments,
                    af.Evidence,
                    af.HowToVerify,
                    af.Impact,
                    af.Recommendation,
                    af.DetailsOfFinding,
                    af.MajorMinor
                FROM 
                    audit_findings af
                JOIN
                    compliance c ON af.ComplianceId = c.ComplianceId
                WHERE 
                    af.AuditId = %s
            """, [audit_id])
            
            findings = cursor.fetchall()
            column_names = ['ComplianceId', 'ComplianceItemDescription', 'Check', 'Comments', 
                          'Evidence', 'HowToVerify', 'Impact', 'Recommendation', 
                          'DetailsOfFinding', 'MajorMinor']
            
            # Build the JSON structure
            for finding in findings:
                compliance_id = str(finding[0])
                finding_data = {}
                
                # Add data from all fields
                for i in range(2, len(column_names)):
                    field_name = column_names[i]
                    field_value = finding[i]
                    finding_data[field_name] = field_value
                
                # Add required reviewer fields
                finding_data['description'] = finding[1]  # ComplianceItemDescription
                
                # Ensure all necessary fields are present with consistent names
                finding_data['accept_reject'] = "0"  # Default to "Not reviewed"
                finding_data['review_status'] = "In Review"  # Same as accept_reject=0
                finding_data['review_comments'] = ""  # Empty review comments
                finding_data['reviewer_comments'] = ""  # Alternate field name
                finding_data['comments'] = ""  # Original comments field
                
                findings_data[compliance_id] = finding_data
        
        # Add metadata
        findings_data['__metadata__'] = {
            'reviewer_id': user_id,
            'review_date': timezone.now().strftime('%Y-%m-%d %H:%M:%S'),
            'version_type': 'Reviewer'
        }
        
        # Add overall comments
        findings_data['overall_comments'] = ""
        
        return findings_data
    except Exception as e:
        print(f"ERROR in get_empty_review_structure: {str(e)}")
        return {
            '__metadata__': {
                'reviewer_id': user_id,
                'review_date': timezone.now().strftime('%Y-%m-%d %H:%M:%S'),
                'version_type': 'Reviewer',
                'error': str(e)
            },
            'overall_comments': ""
        }

@api_view(['GET'])
def load_continuing_data(request, audit_id):
    """
    Load the appropriate version data for an auditor continuing after reviewer feedback
    - First checks if there's a reviewer version (R-prefix) that should be loaded
    - Falls back to latest auditor version (A-prefix) if no reviewer version exists
    """
    try:
        print(f"DEBUG: load_continuing_data called for audit_id: {audit_id}")
        
        # Check if the audit exists
        try:
            audit = Audit.objects.get(AuditId=audit_id)
        except Audit.DoesNotExist:
            return Response({'error': 'Audit not found'}, status=status.HTTP_404_NOT_FOUND)
        
        # Get user ID from session
        user_id = request.session.get('user_id')
        if not user_id and hasattr(audit, 'auditor_id'):
            user_id = audit.auditor_id
        if not user_id:
            user_id = 1050  # Default auditor ID if not found
            
        print(f"DEBUG: Using auditor ID: {user_id}")
        
        # Check if there's a reviewer version (R-prefix)
        latest_version_data = None
        version_id = None
        version_type = None
        
        with connection.cursor() as cursor:
            # First look for reviewer versions
            cursor.execute(
                "SELECT Version, ExtractedInfo, Date FROM audit_version WHERE AuditId = %s AND Version LIKE %s ORDER BY Version DESC LIMIT 1",
                [audit_id, "R%"]
            )
            
            r_version_row = cursor.fetchone()
            if r_version_row:
                version_id = r_version_row[0]
                version_type = "reviewer"
                
                # Parse the JSON data
                if isinstance(r_version_row[1], dict):
                    latest_version_data = r_version_row[1]
                else:
                    try:
                        latest_version_data = json.loads(r_version_row[1]) if r_version_row[1] else {}
                    except Exception as e:
                        print(f"ERROR parsing JSON from reviewer version: {str(e)}")
                        latest_version_data = None
                
                print(f"DEBUG: Found reviewer version {version_id} from {r_version_row[2]}")
            
            # If no reviewer version, fall back to latest auditor version
            if not latest_version_data:
                cursor.execute(
                    "SELECT Version, ExtractedInfo, Date FROM audit_version WHERE AuditId = %s AND Version LIKE %s ORDER BY Version DESC LIMIT 1",
                    [audit_id, "A%"]
                )
                
                a_version_row = cursor.fetchone()
                if a_version_row:
                    version_id = a_version_row[0]
                    version_type = "auditor"
                    
                    # Parse the JSON data
                    if isinstance(a_version_row[1], dict):
                        latest_version_data = a_version_row[1]
                    else:
                        try:
                            latest_version_data = json.loads(a_version_row[1]) if a_version_row[1] else {}
                        except Exception as e:
                            print(f"ERROR parsing JSON from auditor version: {str(e)}")
                            latest_version_data = None
                    
                    print(f"DEBUG: Found auditor version {version_id} from {a_version_row[2]}")
        
        if not latest_version_data:
            # If no version found, create empty structure
            print(f"DEBUG: No version data found for audit {audit_id}")
            return Response({
                'error': 'No version data found for this audit',
                'audit_id': audit_id
            }, status=status.HTTP_404_NOT_FOUND)
        
        # Check if this is a reviewer version (R-prefix)
        is_reviewer_version = version_id and version_id.startswith('R')
        
        # Get the current audit status for context
        audit_status = audit.Status
        
        # Determine the next step based on the loaded version and audit status
        next_step = None
        if is_reviewer_version:
            if audit_status == 'Work In Progress':
                next_step = "Update findings based on reviewer feedback"
            elif audit_status == 'Under review':
                next_step = "Wait for reviewer to complete the review"
            elif audit_status == 'Completed':
                next_step = "Audit is completed, no further action needed"
        else:
            if audit_status == 'Work In Progress':
                next_step = "Continue updating findings"
            elif audit_status == 'Under review':
                next_step = "Wait for reviewer feedback"
            elif audit_status == 'Completed':
                next_step = "Audit is completed, no further action needed"
        
        return Response({
            'audit_id': audit_id,
            'version_id': version_id,
            'version_type': version_type,
            'is_reviewer_version': is_reviewer_version,
            'audit_status': audit_status,
            'next_step': next_step,
            'data': latest_version_data,
            'message': f"Loaded {'reviewer' if is_reviewer_version else 'auditor'} version {version_id}"
        }, status=status.HTTP_200_OK)
    
    except Exception as e:
        print(f"ERROR in load_continuing_data: {str(e)}")
        import traceback
        traceback.print_exc()
        return Response({'error': str(e)}, status=status.HTTP_400_BAD_REQUEST)

@api_view(['GET'])
def load_audit_continuing_data(request, audit_id):
    """
    Load the appropriate version data for an auditor continuing after reviewer feedback
    - First checks if there's a reviewer version (R-prefix) that should be loaded
    - Falls back to latest auditor version (A-prefix) if no reviewer version exists
    """
    try:
        print(f"DEBUG: load_audit_continuing_data called for audit_id: {audit_id}")
        
        # Check if the audit exists
        try:
            audit = Audit.objects.get(AuditId=audit_id)
        except Audit.DoesNotExist:
            return Response({'error': 'Audit not found'}, status=status.HTTP_404_NOT_FOUND)
        
        # Get user ID from session
        user_id = request.session.get('user_id')
        if not user_id and hasattr(audit, 'auditor_id'):
            user_id = audit.auditor_id
        if not user_id:
            user_id = 1050  # Default auditor ID if not found
            
        print(f"DEBUG: Using auditor ID: {user_id}")
        
        # Check if there's a reviewer version (R-prefix)
        latest_version_data = None
        version_id = None
        version_type = None
        
        with connection.cursor() as cursor:
            # First look for reviewer versions
            cursor.execute(
                "SELECT Version, ExtractedInfo, Date FROM audit_version WHERE AuditId = %s AND Version LIKE %s ORDER BY Version DESC LIMIT 1",
                [audit_id, "R%"]
            )
            
            r_version_row = cursor.fetchone()
            if r_version_row:
                version_id = r_version_row[0]
                version_type = "reviewer"
                
                # Parse the JSON data
                if isinstance(r_version_row[1], dict):
                    latest_version_data = r_version_row[1]
                else:
                    try:
                        latest_version_data = json.loads(r_version_row[1]) if r_version_row[1] else {}
                    except Exception as e:
                        print(f"ERROR parsing JSON from reviewer version: {str(e)}")
                        latest_version_data = None
                
                print(f"DEBUG: Found reviewer version {version_id} from {r_version_row[2]}")
                
                # Process reviewer data to ensure review comments are properly formatted
                if latest_version_data:
                    # Process each compliance item to ensure consistent field names
                    for compliance_id, finding in list(latest_version_data.items()):
                        if compliance_id in ['__metadata__', 'overall_comments']:
                            continue
                            
                        if isinstance(finding, dict):
                            # Ensure both comments field formats exist
                            if 'review_comments' in finding and 'comments' not in finding:
                                finding['comments'] = finding['review_comments']
                            elif 'comments' in finding and 'review_comments' not in finding:
                                finding['review_comments'] = finding['comments']
                                
                            # Ensure review_status is consistent
                            if 'review_status' not in finding:
                                # Derive from accept_reject
                                accept_reject = finding.get('accept_reject', "0")
                                if accept_reject == "1" or accept_reject == 1:
                                    finding['review_status'] = 'Accept'
                                elif accept_reject == "2" or accept_reject == 2:
                                    finding['review_status'] = 'Reject'
                                else:
                                    finding['review_status'] = 'In Review'
                    
                    # Make sure overall_comments is at the root level for consistency
                    if 'overall_comments' not in latest_version_data and '__metadata__' in latest_version_data:
                        metadata = latest_version_data['__metadata__']
                        if isinstance(metadata, dict) and 'overall_comments' in metadata:
                            latest_version_data['overall_comments'] = metadata['overall_comments']
                    
                    print(f"DEBUG: Processed reviewer data with {len(latest_version_data) - (1 if '__metadata__' in latest_version_data else 0)} compliance items")
            
            # If no reviewer version, fall back to latest auditor version
            if not latest_version_data:
                cursor.execute(
                    "SELECT Version, ExtractedInfo, Date FROM audit_version WHERE AuditId = %s AND Version LIKE %s ORDER BY Version DESC LIMIT 1",
                    [audit_id, "A%"]
                )
                
                a_version_row = cursor.fetchone()
                if a_version_row:
                    version_id = a_version_row[0]
                    version_type = "auditor"
                    
                    # Parse the JSON data
                    if isinstance(a_version_row[1], dict):
                        latest_version_data = a_version_row[1]
                    else:
                        try:
                            latest_version_data = json.loads(a_version_row[1]) if a_version_row[1] else {}
                        except Exception as e:
                            print(f"ERROR parsing JSON from auditor version: {str(e)}")
                            latest_version_data = None
                    
                    print(f"DEBUG: Found auditor version {version_id} from {a_version_row[2]}")
            
            # If we still don't have version data, check if there's review data in the audit_findings table
            if not latest_version_data or version_type != "reviewer":
                cursor.execute("""
                    SELECT 
                        af.ComplianceId, 
                        af.ReviewStatus, 
                        af.ReviewComments,
                        af.ReviewRejected,
                        c.ComplianceItemDescription
                    FROM 
                        audit_findings af
                    JOIN
                        compliance c ON af.ComplianceId = c.ComplianceId
                    WHERE 
                        af.AuditId = %s AND 
                        (af.ReviewStatus IS NOT NULL OR af.ReviewComments IS NOT NULL)
                """, [audit_id])
                
                review_findings = cursor.fetchall()
                
                if review_findings:
                    print(f"DEBUG: Found {len(review_findings)} review findings in audit_findings table")
                    
                    # If we don't have version data yet, create a basic structure
                    if not latest_version_data:
                        latest_version_data = {'__metadata__': {}}
                        version_type = "derived"
                        version_id = "Derived"
                    
                    # Add the review data from audit_findings to the version data
                    for finding in review_findings:
                        compliance_id = str(finding[0])
                        review_status = finding[1]
                        review_comments = finding[2]
                        review_rejected = finding[3]
                        compliance_description = finding[4]
                        
                        # Skip if we already have this compliance from a proper version
                        if compliance_id in latest_version_data and version_type in ["reviewer", "auditor"]:
                            continue
                        
                        # Create or update the finding data
                        if compliance_id not in latest_version_data:
                            latest_version_data[compliance_id] = {
                                'description': compliance_description
                            }
                        
                        # Set the review status (convert from database format)
                        if review_status:
                            latest_version_data[compliance_id]['review_status'] = review_status
                        elif review_rejected == 1:
                            latest_version_data[compliance_id]['review_status'] = 'Reject'
                        elif review_rejected == 0:
                            latest_version_data[compliance_id]['review_status'] = 'Accept'
                        else:
                            latest_version_data[compliance_id]['review_status'] = 'In Review'
                        
                        # Set accept_reject based on review_status
                        if latest_version_data[compliance_id]['review_status'] == 'Accept':
                            latest_version_data[compliance_id]['accept_reject'] = "1"
                        elif latest_version_data[compliance_id]['review_status'] == 'Reject':
                            latest_version_data[compliance_id]['accept_reject'] = "2"
                        else:
                            latest_version_data[compliance_id]['accept_reject'] = "0"
                        
                        # Set review comments
                        if review_comments:
                            latest_version_data[compliance_id]['review_comments'] = review_comments
                            latest_version_data[compliance_id]['comments'] = review_comments
                    
                    # Check if there are overall comments in the audit table
                    cursor.execute("""
                        SELECT ReviewComments FROM audit WHERE AuditId = %s AND ReviewComments IS NOT NULL
                    """, [audit_id])
                    
                    overall_row = cursor.fetchone()
                    if overall_row and overall_row[0]:
                        latest_version_data['overall_comments'] = overall_row[0]
                        print(f"DEBUG: Added overall comments from audit table: {overall_row[0]}")
        
        if not latest_version_data:
            # If no version found, create empty structure
            print(f"DEBUG: No version data found for audit {audit_id}")
            return Response({
                'error': 'No version data found for this audit',
                'audit_id': audit_id
            }, status=status.HTTP_404_NOT_FOUND)
        
        # Check if this is a reviewer version (R-prefix)
        is_reviewer_version = version_id and version_id.startswith('R')
        
        # Get the current audit status for context
        audit_status = audit.Status
        
        # Determine the next step based on the loaded version and audit status
        next_step = None
        if is_reviewer_version:
            if audit_status == 'Work In Progress':
                next_step = "Update findings based on reviewer feedback"
            elif audit_status == 'Under review':
                next_step = "Wait for reviewer to complete the review"
            elif audit_status == 'Completed':
                next_step = "Audit is completed, no further action needed"
        else:
            if audit_status == 'Work In Progress':
                next_step = "Continue updating findings"
            elif audit_status == 'Under review':
                next_step = "Wait for reviewer feedback"
            elif audit_status == 'Completed':
                next_step = "Audit is completed, no further action needed"
        
        # Add metadata about approved/rejected status if missing
        if '__metadata__' in latest_version_data:
            metadata = latest_version_data['__metadata__']
            if not isinstance(metadata, dict):
                metadata = {}
                
            # If we have ApprovedRejected in the metadata, make sure it's also in the version_info
            if 'ApprovedRejected' not in metadata:
                # Try to determine from ReviewStatus in the audit table
                if hasattr(audit, 'ReviewStatus'):
                    if audit.ReviewStatus == 2:  # Accept
                        metadata['ApprovedRejected'] = 'Approved'
                    elif audit.ReviewStatus == 3:  # Reject
                        metadata['ApprovedRejected'] = 'Rejected'
            
            latest_version_data['__metadata__'] = metadata
        
        return Response({
            'audit_id': audit_id,
            'version_id': version_id,
            'version_type': version_type,
            'is_reviewer_version': is_reviewer_version,
            'audit_status': audit_status,
            'next_step': next_step,
            'data': latest_version_data,
            'message': f"Loaded {'reviewer' if is_reviewer_version else 'auditor'} version {version_id}",
            'version_info': latest_version_data.get('__metadata__', {})
        }, status=status.HTTP_200_OK)
    
    except Exception as e:
        print(f"ERROR in load_audit_continuing_data: {str(e)}")
        import traceback
        traceback.print_exc()
        return Response({'error': str(e)}, status=status.HTTP_400_BAD_REQUEST)
        
@api_view(['POST'])
def save_audit_version(request, audit_id):
    """
    Save a new version of audit findings when auditor makes changes after review
    """
    try:
        print(f"\n\n==== DEBUG: save_audit_version called for audit_id: {audit_id} ====")
        
        # Get the audit object
        audit = Audit.objects.get(pk=audit_id)
        user_id = request.data.get('user_id')
        
        print(f"DEBUG: User ID: {user_id}, Auditor ID: {audit.Auditor.UserId}")
        
        # Check if user is the assigned auditor
        if int(user_id) != audit.Auditor.UserId:
            return Response(
                {'error': 'Only the assigned auditor can submit changes'}, 
                status=status.HTTP_403_FORBIDDEN
            )
            
        # Get the audit's current status
        if audit.Status not in ['In Progress', 'Rejected']:
            return Response(
                {'error': f'Cannot update audit in its current state ({audit.Status})'}, 
                status=status.HTTP_400_BAD_REQUEST
            )
        
        # Get the compliance data from request
        compliance_data = request.data.get('compliance_data', {})
        overall_comments = request.data.get('overall_comments', '')
        
        print(f"DEBUG: Got compliance data for {len(compliance_data)} items")
        print(f"DEBUG: Overall comments: {overall_comments}")
        
        # Log a sample of the actual compliance data received
        if compliance_data:
            first_key = next(iter(compliance_data))
            print(f"DEBUG: Sample compliance data for ID {first_key}: {compliance_data[first_key]}")
        
        # First, let's save the individual compliance findings
        processed_compliance_count = 0
        for compliance_id, data in compliance_data.items():
            print(f"DEBUG: Processing compliance ID: {compliance_id}")
            
            # Update audit finding in database
            finding, created = AuditFinding.objects.get_or_create(
                AuditId=audit,
                ComplianceId_id=compliance_id,
                defaults={
                    'UserId_id': user_id,
                    'Evidence': data.get('evidence', ''),
                    'AssignedDate': timezone.now()
                }
            )
            
            print(f"DEBUG: Finding {'created' if created else 'already exists'}")
            
            # Convert status values
            status_value = data.get('compliance_status')
            if status_value == 'Not Compliant':
                check_value = '0'
            elif status_value == 'Partially Compliant':
                check_value = '1'
            elif status_value == 'Fully Compliant':
                check_value = '2'
            elif status_value == 'Not Applicable':
                check_value = '3'
            else:
                check_value = '0'  # Default to Not Compliant
                
            # Convert criticality values
            criticality = data.get('criticality', '').strip().lower() if data.get('criticality') else ''
            if criticality == 'minor':
                major_minor = '0'  # Minor: 0
            elif criticality == 'major':
                major_minor = '1'  # Major: 1
            elif criticality == 'not applicable':
                major_minor = '2'
            else:
                major_minor = '0'  # Default to Minor
                
            # Update finding fields
            finding.Check = check_value
            finding.MajorMinor = major_minor
            finding.HowToVerify = data.get('how_to_verify', '')
            finding.Impact = data.get('impact', '')
            finding.Recommendation = data.get('recommendation', '')
            finding.DetailsOfFinding = data.get('details_of_finding', '')
            finding.Comments = data.get('comments', '')
            finding.CheckedDate = timezone.now()
            finding.save()
            processed_compliance_count += 1
            
        print(f"DEBUG: Successfully processed {processed_compliance_count} compliance items")
        
        # Now get the complete set of findings for the version
        version_data = get_audit_findings_json(audit_id, overall_comments)
        
        # Check if we got valid JSON data
        if not version_data:
            print(f"ERROR: Failed to generate valid JSON data from audit findings")
            return Response(
                {'error': 'Failed to generate version data'}, 
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )
            
        # Print complete JSON for better debugging (up to 10KB)
        json_data = json.dumps(version_data, indent=2)
        json_size = len(json_data)
        max_print_size = 10000  # 10KB limit for logs
        print(f"DEBUG: Version data JSON size: {json_size} bytes")
        print(f"DEBUG: Complete JSON data (truncated if > 10KB):\n{json_data[:max_print_size]}")
        if json_size > max_print_size:
            print("... (JSON truncated due to size)")
        
        # CRITICAL: Check for existing versions and determine the next version number
        next_version = ''
        with connection.cursor() as cursor:
            # Check if there are any existing versions
            cursor.execute(
                """
                SELECT Version FROM audit_version 
                WHERE AuditId = %s AND Version LIKE 'A%' 
                ORDER BY CAST(SUBSTRING(Version, 2) AS UNSIGNED) DESC, Version DESC
                LIMIT 1
                """,
                [audit_id]
            )
            
            row = cursor.fetchone()
            
            if row:
                current_version = row[0]
                print(f"DEBUG: Found existing version: {current_version}")
                
                # Extract the number part and increment
                if current_version.startswith('A'):
                    # Handle case where version might have non-numeric parts after the A prefix
                    num_part = ''
                    for c in current_version[1:]:
                        if c.isdigit():
                            num_part += c
                        else:
                            break
                    
                    if num_part:
                        version_num = int(num_part)
                        next_version = f"A{version_num + 1}"
                        print(f"DEBUG: Incrementing from {current_version} to {next_version}")
                    else:
                        # Default to A2 if we couldn't parse the number
                        next_version = "A2"
                        print(f"DEBUG: Couldn't parse number from {current_version}, using {next_version}")
                else:
                    # If the version doesn't start with A, just start with A1
                    next_version = "A1"
                    print(f"DEBUG: Version {current_version} doesn't start with 'A', using {next_version}")
            else:
                # No existing versions, start with A1
                next_version = "A1"
                print(f"DEBUG: No existing versions found, starting with {next_version}")
        
        # Force a transaction to ensure the version is created correctly
        with connection.cursor() as cursor:
            cursor.execute("START TRANSACTION")
            try:
                # Double-check that this version doesn't already exist (race condition protection)
                cursor.execute(
                    "SELECT COUNT(*) FROM audit_version WHERE AuditId = %s AND Version = %s",
                    [audit_id, next_version]
                )
                if cursor.fetchone()[0] > 0:
                    # If duplicate, add timestamp to ensure uniqueness
                    timestamp = timezone.now().strftime('%H%M%S')
                    next_version = f"{next_version}_{timestamp}"
                    print(f"DEBUG: Version already exists! Adjusted to ensure uniqueness: {next_version}")
                
                # Insert the new version
                print(f"DEBUG: INSERTING NEW VERSION {next_version} for audit {audit_id} with JSON size {json_size}")
                
                # Execute the insert
                cursor.execute(
                    """
                    INSERT INTO audit_version (AuditId, Version, ExtractedInfo, UserId, Date)
                    VALUES (%s, %s, %s, %s, NOW())
                    """,
                    [audit_id, next_version, json_data, user_id,timezone.now()]
                )
                
                # Check rows affected
                rows_affected = cursor.rowcount
                print(f"DEBUG: Insert affected {rows_affected} rows")
                
                # Commit transaction
                cursor.execute("COMMIT")
                print(f"DEBUG: Successfully committed transaction for version {next_version}")
                
                # Verify the insert worked by querying it back
                cursor.execute(
                    "SELECT COUNT(*) FROM audit_version WHERE AuditId = %s AND Version = %s",
                    [audit_id, next_version]
                )
                verify_count = cursor.fetchone()[0]
                print(f"DEBUG: Verification query found {verify_count} records")
                
            except Exception as e:
                cursor.execute("ROLLBACK")
                print(f"ERROR creating version: {str(e)}")
                import traceback
                traceback.print_exc()
                raise e
        
        # Update audit status to indicate it needs review
        audit.Status = 'Pending Review'
        audit.save()
        print(f"DEBUG: Updated audit status to 'Pending Review'")
        
        print(f"==== DEBUG: save_audit_version completed successfully with version {next_version} ====\n")
        
        return Response({
            'success': True,
            'message': f'Audit changes saved as version {next_version}',
            'version': next_version,
            'audit_id': audit_id
        }, status=status.HTTP_200_OK)
        
    except Audit.DoesNotExist:
        print(f"ERROR: Audit not found for ID {audit_id}")
        return Response({'error': 'Audit not found'}, status=status.HTTP_404_NOT_FOUND)
    except Exception as e:
        print(f"ERROR in save_audit_version: {str(e)}")
        import traceback
        traceback.print_exc()
        return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        
# @api_view(['POST'])
# def save_review_version(request, audit_id):
#     """
#     Save reviewer decisions and create a new review version
#     """
#     try:
#         # Get the audit object
#         audit = Audit.objects.get(pk=audit_id)
#         user_id = request.data.get('user_id')
        
#         # Check if user is the assigned reviewer
#         if int(user_id) != audit.Reviewer.UserId:
#             return Response(
#                 {'error': 'Only the assigned reviewer can submit reviews'}, 
#                 status=status.HTTP_403_FORBIDDEN
#             )
            
#         # Get the audit's current status
#         if audit.Status != 'Pending Review':
#             return Response(
#                 {'error': 'Cannot review audit in its current state'}, 
#                 status=status.HTTP_400_BAD_REQUEST
#             )
        
#         # Get review data
#         compliance_reviews = request.data.get('compliance_reviews', {})
#         overall_status = request.data.get('overall_status', '')
#         overall_comments = request.data.get('overall_comments', '')
        
#         # Get the next version number (R1, R2, etc.)
#         next_version = get_next_version_number(audit_id, 'R')
        
#         # Create a new review version with the data
#         version_id = create_review_version(
#             audit_id, 
#             user_id, 
#             compliance_reviews=compliance_reviews,
#             overall_comments=overall_comments,
#             custom_version=next_version
#         )
        
#         # Update audit status based on overall review decision
#         if overall_status == 'Accept':
#             audit.Status = 'Completed'
#             audit.ReviewStatus = 'Approved'
#         else:
#             audit.Status = 'Rejected'
#             audit.ReviewStatus = 'Rejected'
            
#         audit.ReviewComments = overall_comments
#         audit.ReviewDate = timezone.now()
#         audit.save()
                
#         return Response({
#             'success': True,
#             'message': f'Review saved as version {next_version}',
#             'version': next_version,
#             'version_id': version_id,
#             'status': audit.Status
#         }, status=status.HTTP_200_OK)
        
#     except Audit.DoesNotExist:
#         return Response({'error': 'Audit not found'}, status=status.HTTP_404_NOT_FOUND)
#     except Exception as e:
#         return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        
@api_view(['GET'])
def get_latest_reviewer_data(request, audit_id):
    """
    Get the latest reviewer data for an audit to show to the auditor
    """
    try:
        print(f"DEBUG: get_latest_reviewer_data called for audit_id: {audit_id}")
        
        # Store the audit ID in the session for future requests
        # request.session['current_audit_id'] = audit_id
        
        # Get the audit to check status
        try:
            audit = Audit.objects.get(AuditId=audit_id)
        except Audit.DoesNotExist:
            return Response({'error': 'Audit not found'}, status=status.HTTP_404_NOT_FOUND)
        
        # Find the latest reviewer version (R-prefix)
        latest_reviewer_data = None
        latest_version_id = None
        
        with connection.cursor() as cursor:
            # Get the latest R-prefix version
            cursor.execute(
                """
                SELECT Version, ExtractedInfo, Date 
                FROM audit_version 
                WHERE AuditId = %s AND Version LIKE 'R%' 
                ORDER BY Date DESC LIMIT 1
                """,
                [audit_id]
            )
            
            row = cursor.fetchone()
            
            if row:
                latest_version_id = row[0]
                extracted_info = row[1]
                version_date = row[2]
                
                print(f"DEBUG: Found latest reviewer version: {latest_version_id} from {version_date}")
                
                try:
                    # Parse the JSON data
                    latest_reviewer_data = json.loads(extracted_info)
                except Exception as e:
                    print(f"ERROR: Could not parse JSON from reviewer version: {str(e)}")
                    return Response(
                        {'error': 'Error parsing reviewer data'}, 
                        status=status.HTTP_500_INTERNAL_SERVER_ERROR
                    )
            else:
                print(f"DEBUG: No reviewer versions found for audit {audit_id}")
                return Response(
                    {'error': 'No reviewer data available for this audit'}, 
                    status=status.HTTP_404_NOT_FOUND
                )
        
        # Add some helpful metadata 
        response_data = {
            'audit_id': audit_id,
            'audit_status': audit.Status,
            'version_id': latest_version_id,
            'version_type': 'Reviewer',
            'review_data': latest_reviewer_data,
            'review_comments': audit.ReviewComments or '',
            'review_status': audit.ReviewStatus or '',
            'message': f"Loaded reviewer version {latest_version_id} for auditor to view"
        }
        
        # Return the reviewer data
        return Response(response_data, status=status.HTTP_200_OK)
    
    except Exception as e:
        print(f"ERROR in get_latest_reviewer_data: {str(e)}")
        import traceback
        traceback.print_exc()
        return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        
@api_view(['GET'])
def load_audit_with_reviewer_feedback(request, audit_id):
    """
    Load latest reviewer feedback for an auditor to continue working
    """
    try:
        print(f"DEBUG: load_audit_with_reviewer_feedback called for audit_id: {audit_id}")
        
        # Store the audit ID in the session for future requests
        audit_id=request.session.get('current_audit_id')
        
        # Get the audit to check status
        try:
            audit = Audit.objects.get(AuditId=audit_id)
        except Audit.DoesNotExist:
            return Response({'error': 'Audit not found'}, status=status.HTTP_404_NOT_FOUND)
        
        # Check audit status to ensure it's been reviewed
        if audit.Status != 'Rejected':
            return Response({
                'error': 'This audit does not have reviewer feedback to load',
                'current_status': audit.Status
            }, status=status.HTTP_400_BAD_REQUEST)
        
        # Get the latest reviewer (R-prefix) version
        latest_reviewer_data = None
        latest_reviewer_version = None
        latest_reviewer_date = None
        
        with connection.cursor() as cursor:
            # Get the most recent R-prefix version
            cursor.execute(
                """
                SELECT Version, ExtractedInfo, Date 
                FROM audit_version 
                WHERE AuditId = %s AND Version LIKE 'R%' 
                ORDER BY Date DESC LIMIT 1
                """,
                [audit_id]
            )
            
            row = cursor.fetchone()
            
            if row:
                latest_reviewer_version = row[0]
                reviewer_json = row[1]
                latest_reviewer_date = row[2]
                
                print(f"DEBUG: Found latest reviewer version: {latest_reviewer_version} from {latest_reviewer_date}")
                
                try:
                    # Parse the JSON data
                    latest_reviewer_data = json.loads(reviewer_json)
                except Exception as e:
                    print(f"ERROR: Could not parse JSON from reviewer version: {str(e)}")
                    
        # Get the latest auditor (A-prefix) version as a base
        latest_auditor_data = None
        latest_auditor_version = None
        
        with connection.cursor() as cursor:
            # Get the most recent A-prefix version
            cursor.execute(
                """
                SELECT Version, ExtractedInfo, Date 
                FROM audit_version 
                WHERE AuditId = %s AND Version LIKE 'A%' 
                ORDER BY Date DESC LIMIT 1
                """,
                [audit_id]
            )
            
            row = cursor.fetchone()
            
            if row:
                latest_auditor_version = row[0]
                auditor_json = row[1]
                
                print(f"DEBUG: Found latest auditor version: {latest_auditor_version}")
                
                try:
                    # Parse the JSON data
                    latest_auditor_data = json.loads(auditor_json)
                except Exception as e:
                    print(f"ERROR: Could not parse JSON from auditor version: {str(e)}")
        
        # If we didn't find any versions, use the current audit findings
        if not latest_auditor_data and not latest_reviewer_data:
            print(f"DEBUG: No versions found, getting current audit findings")
            data = get_audit_findings_json(audit_id)
            if data:
                latest_auditor_data = data
        
        # Use reviewer data if available, otherwise fall back to auditor data
        response_data = latest_reviewer_data if latest_reviewer_data else latest_auditor_data
        
        # Add reviewer feedback to the response
        if latest_reviewer_data and '__review_data__' in latest_reviewer_data:
            response_data['__reviewer_feedback__'] = latest_reviewer_data['__review_data__']
        
        # Add reviewer comments
        if audit.ReviewComments:
            response_data['reviewer_comments'] = audit.ReviewComments
        
        # Add some helpful metadata
        return Response({
            'audit_id': audit_id,
            'audit_status': audit.Status,
            'reviewer_version': latest_reviewer_version,
            'reviewer_date': latest_reviewer_date,
            'auditor_version': latest_auditor_version,
            'review_status': audit.ReviewStatus or '',
            'data': response_data,
            'message': f"Loaded reviewer feedback from {latest_reviewer_version} for auditor to continue"
        }, status=status.HTTP_200_OK)
    
    except Exception as e:
        print(f"ERROR in load_audit_with_reviewer_feedback: {str(e)}")
        import traceback
        traceback.print_exc()
        return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        
@api_view(['POST'])
def debug_audit_status_transition(request, audit_id):
    """Debug endpoint for tracking the audit status transition process"""
    print(f"=== DEBUG: Status transition for audit_id={audit_id} ===")
    print(f"DEBUG: Request data: {request.data}")
    
    # Get the audit record
    try:
        audit = Audit.objects.get(AuditId=audit_id)
        print(f"DEBUG: Found audit: {audit.AuditId}, current status: {audit.Status}")
        
        # Get all audit findings for this audit
        findings = AuditFinding.objects.filter(AuditId=audit.AuditId)
        print(f"DEBUG: Found {findings.count()} findings for audit {audit.AuditId}")
        
        for finding in findings:
            print(f"Finding ID: {finding.AuditFindingsId}, Compliance ID: {finding.ComplianceId.ComplianceId}, Status: {finding.Check}")
        
        # Continue with normal status update logic
        new_status = request.data.get('status')
        print(f"Changing status from {audit.Status} to {new_status}")
        
        # Update the audit status
        audit.Status = new_status
        audit.save()
        
        return Response({"status": "success", "message": f"Audit status updated to {new_status}"})
    
    except Audit.DoesNotExist:
        print(f"ERROR: Audit with ID {audit_id} not found")
        return Response({"status": "error", "message": "Audit not found"}, status=404)
    except Exception as e:
        print(f"ERROR: {str(e)}")
        return Response({"status": "error", "message": str(e)}, status=500)
        
@api_view(['GET'])
def debug_audit_version_schema(request):
    """Debug endpoint to check the audit_version table schema"""
    try:
        with connection.cursor() as cursor:
            cursor.execute("DESCRIBE audit_version")
            columns = cursor.fetchall()
            column_names = [column[0] for column in columns]
            
            return Response({
                "status": "success",
                "columns": column_names
            })
    except Exception as e:
        print(f"ERROR: {str(e)}")
        return Response({"status": "error", "message": str(e)}, status=500)
        
@api_view(['GET'])
def debug_audit_versions(request, audit_id):
    """Debug endpoint to check existing versions for an audit"""
    try:
        with connection.cursor() as cursor:
            cursor.execute(
                "SELECT * FROM audit_version WHERE AuditId = %s",
                [audit_id]
            )
            versions = cursor.fetchall()
            
            return Response({
                "status": "success",
                "count": len(versions),
                "versions": [{"id": v[0], "version": v[2]} for v in versions]
            })
    except Exception as e:
        print(f"ERROR: {str(e)}")
        return Response({"status": "error", "message": str(e)}, status=500)
        
@api_view(['GET'])
def debug_review_process(request, audit_id):
    """Debug endpoint to track the entire review process and status transitions"""
    try:
        print(f"=== DEBUG: Review process debug for audit_id={audit_id} ===")
        
        # Get the audit record
        audit = Audit.objects.get(AuditId=audit_id)
        
        # Get all version data
        versions = []
        with connection.cursor() as cursor:
            cursor.execute(
                """
                SELECT Version, ExtractedInfo, Date, UserId, ApproverId, ApprovedRejected 
                FROM audit_version 
                WHERE AuditId = %s 
                ORDER BY Date ASC
                """,
                [audit_id]
            )
            
            for row in cursor.fetchall():
                version_id = row[0]
                json_data = row[1]
                date = row[2]
                user_id = row[3]
                approver_id = row[4]
                approved_rejected = row[5]
                
                # Parse the JSON data
                if isinstance(json_data, str):
                    try:
                        extracted_data = json.loads(json_data)
                    except:
                        extracted_data = {"error": "Could not parse JSON"}
                else:
                    extracted_data = json_data
                
                # Add to versions list
                versions.append({
                    "version_id": version_id,
                    "date": date.strftime('%Y-%m-%d %H:%M:%S') if date else None,
                    "user_id": user_id,
                    "approver_id": approver_id,
                    "approved_rejected": approved_rejected,
                    "data": extracted_data
                })
        
        # Get audit findings
        findings = []
        with connection.cursor() as cursor:
            cursor.execute(
                """
                SELECT 
                    af.ComplianceId,
                    af.`Check`,
                    af.Comments,
                    af.Evidence,
                    af.HowToVerify,
                    af.Impact,
                    af.Recommendation,
                    af.DetailsOfFinding,
                    af.MajorMinor,
                    af.CheckedDate,
                    af.ReviewRejected,
                    af.ReviewComments,
                    af.ReviewStatus,
                    af.ReviewDate
                FROM 
                    audit_findings af
                WHERE 
                    af.AuditId = %s
                """,
                [audit_id]
            )
            
            columns = [col[0] for col in cursor.description]
            
            for row in cursor.fetchall():
                finding_data = dict(zip(columns, row))
                
                # Format dates
                if finding_data.get('CheckedDate'):
                    finding_data['CheckedDate'] = finding_data['CheckedDate'].strftime('%Y-%m-%d %H:%M:%S')
                if finding_data.get('ReviewDate'):
                    finding_data['ReviewDate'] = finding_data['ReviewDate'].strftime('%Y-%m-%d %H:%M:%S')
                
                findings.append(finding_data)
        
        # Audit details
        audit_data = {
            "AuditId": audit.AuditId,
            "Status": audit.Status,
            "ReviewStatus": audit.ReviewStatus,
            "ReviewComments": audit.ReviewComments,
            "DueDate": audit.DueDate.strftime('%Y-%m-%d') if audit.DueDate else None,
            "CompletionDate": audit.CompletionDate.strftime('%Y-%m-%d %H:%M:%S') if audit.CompletionDate else None,
            "ReviewStartDate": audit.ReviewStartDate.strftime('%Y-%m-%d %H:%M:%S') if audit.ReviewStartDate else None,
            "ReviewDate": audit.ReviewDate.strftime('%Y-%m-%d %H:%M:%S') if audit.ReviewDate else None
        }
        
        # Get the latest reviewer version (R-prefix)
        latest_reviewer_version = next((v for v in versions if v["version_id"].startswith("R")), None)
        
        # Get the latest auditor version (A-prefix)
        latest_auditor_version = next((v for v in versions if v["version_id"].startswith("A")), None)
        
        # Response data
        response_data = {
            "audit": audit_data,
            "versions_count": len(versions),
            "latest_auditor_version": latest_auditor_version["version_id"] if latest_auditor_version else None,
            "latest_reviewer_version": latest_reviewer_version["version_id"] if latest_reviewer_version else None,
            "findings_count": len(findings),
            "versions": versions,
            "findings": findings
        }
        
        return Response(response_data, status=status.HTTP_200_OK)
    
    except Audit.DoesNotExist:
        return Response({"error": f"Audit with ID {audit_id} not found"}, status=status.HTTP_404_NOT_FOUND)
    except Exception as e:
        print(f"ERROR in debug_review_process: {str(e)}")
        import traceback
        traceback.print_exc()
        return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        
@api_view(['GET'])
def debug_rejection_process(request, audit_id):
    """Debug endpoint to specifically track the rejection process and data flow"""
    try:
        print(f"=== DEBUG: Rejection process debug for audit_id={audit_id} ===")
        
        # Get the audit record
        audit = Audit.objects.get(AuditId=audit_id)
        
        # Check if this audit has been rejected
        is_rejected = audit.Status == 'Rejected' or (audit.ReviewStatus and audit.ReviewStatus in ['3', 'Rejected'])
        
        if not is_rejected:
            return Response({
                "message": f"Audit {audit_id} has not been rejected",
                "current_status": audit.Status,
                "review_status": audit.ReviewStatus
            }, status=status.HTTP_200_OK)
        
        # Get the latest reviewer data
        latest_reviewer_data = None
        
        with connection.cursor() as cursor:
            cursor.execute(
                """
                SELECT Version, ExtractedInfo, Date 
                FROM audit_version 
                WHERE AuditId = %s AND Version LIKE 'R%'
                ORDER BY Date DESC LIMIT 1
                """,
                [audit_id]
            )
            
            row = cursor.fetchone()
            
            if row:
                rejection_version = row[0]
                rejection_json = row[1]
                rejection_date = row[2]
                
                # Parse JSON data
                if isinstance(rejection_json, str):
                    try:
                        latest_reviewer_data = json.loads(rejection_json)
                    except Exception as e:
                        print(f"ERROR parsing rejection JSON: {str(e)}")
                        latest_reviewer_data = {"error": "Could not parse JSON"}
                else:
                    latest_reviewer_data = rejection_json
                
                print(f"DEBUG: Found reviewer version: {rejection_version} from {rejection_date}")
            else:
                print(f"DEBUG: No reviewer version found")
        
        # Extract rejection comments from all sources
        rejection_data = {
            "audit_id": audit_id,
            "status": audit.Status,
            "review_status": audit.ReviewStatus,
            "review_comments": audit.ReviewComments,
            "rejection_version": rejection_version if 'rejection_version' in locals() else None,
            "rejection_date": rejection_date.strftime('%Y-%m-%d %H:%M:%S') if 'rejection_date' in locals() and rejection_date else None,
            "reviewer_data": latest_reviewer_data,
        }
        
        # Also get individual finding rejections
        rejected_findings = []
        
        with connection.cursor() as cursor:
            cursor.execute(
                """
                SELECT 
                    af.ComplianceId,
                    af.ReviewRejected,
                    af.ReviewComments,
                    af.ReviewStatus,
                    af.ReviewDate,
                    c.ComplianceItemDescription
                FROM 
                    audit_findings af
                JOIN
                    compliance c ON af.ComplianceId = c.ComplianceId
                WHERE 
                    af.AuditId = %s AND 
                    (af.ReviewRejected = 1 OR af.ReviewStatus = 'Reject')
                """,
                [audit_id]
            )
            
            columns = [col[0] for col in cursor.description]
            
            for row in cursor.fetchall():
                finding_data = dict(zip(columns, row))
                
                # Format dates
                if finding_data.get('ReviewDate'):
                    finding_data['ReviewDate'] = finding_data['ReviewDate'].strftime('%Y-%m-%d %H:%M:%S')
                
                rejected_findings.append(finding_data)
        
        rejection_data["rejected_findings"] = rejected_findings
        rejection_data["rejected_findings_count"] = len(rejected_findings)
        
        # Get any work that happened after rejection
        post_rejection_versions = []
        
        if 'rejection_date' in locals() and rejection_date:
            with connection.cursor() as cursor:
                cursor.execute(
                    """
                    SELECT Version, ExtractedInfo, Date 
                    FROM audit_version 
                    WHERE AuditId = %s AND Date > %s
                    ORDER BY Date ASC
                    """,
                    [audit_id, rejection_date]
                )
                
                for row in cursor.fetchall():
                    post_version = row[0]
                    post_date = row[2]
                    
                    post_rejection_versions.append({
                        "version": post_version,
                        "date": post_date.strftime('%Y-%m-%d %H:%M:%S') if post_date else None
                    })
        
        rejection_data["post_rejection_versions"] = post_rejection_versions
        rejection_data["post_rejection_versions_count"] = len(post_rejection_versions)
        
        return Response(rejection_data, status=status.HTTP_200_OK)
    
    except Audit.DoesNotExist:
        return Response({"error": f"Audit with ID {audit_id} not found"}, status=status.HTTP_404_NOT_FOUND)
    except Exception as e:
        print(f"ERROR in debug_rejection_process: {str(e)}")
        import traceback
        traceback.print_exc()
        return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        
@api_view(['POST'])
def debug_save_reviewer_comments(request, audit_id):
    """Debug endpoint to save reviewer comments directly to the JSON data"""
    try:
        print(f"=== DEBUG: Saving reviewer comments for audit_id={audit_id} ===")
        
        # Get the audit record
        audit = Audit.objects.get(AuditId=audit_id)
        
        # Check if the audit is in a reviewable state
        if audit.Status not in ['Under review', 'Pending Review']:
            return Response({
                'error': f'Cannot save reviewer comments when audit is in {audit.Status} state'
            }, status=status.HTTP_400_BAD_REQUEST)
        
        # Get comments data from request
        comments_data = request.data.get('comments', {})
        overall_comments = request.data.get('overall_comments', '')
        user_id = request.data.get('user_id', audit.Reviewer.UserId if audit.Reviewer else None)
        
        print(f"DEBUG: Received {len(comments_data)} comments and overall comment: '{overall_comments}'")
        
        # Get the latest version to update
        latest_version = None
        version_data = None
        
        with connection.cursor() as cursor:
            # First check for an existing R version
            cursor.execute(
                """
                SELECT Version, ExtractedInfo, Date 
                FROM audit_version 
                WHERE AuditId = %s AND Version LIKE 'R%'
                ORDER BY Date DESC LIMIT 1
                """,
                [audit_id]
            )
            
            row = cursor.fetchone()
            if row:
                latest_version = row[0]
                json_data = row[1]
                version_date = row[2]
                
                # Parse the JSON data
                if isinstance(json_data, str):
                    try:
                        version_data = json.loads(json_data)
                    except Exception as e:
                        print(f"ERROR parsing JSON: {str(e)}")
                        return Response({"error": f"Could not parse JSON: {str(e)}"}, 
                                       status=status.HTTP_500_INTERNAL_SERVER_ERROR)
                else:
                    version_data = json_data
                
                print(f"DEBUG: Found existing reviewer version {latest_version} from {version_date}")
            else:
                # If no R version exists, try to get latest A version
                cursor.execute(
                    """
                    SELECT Version, ExtractedInfo, Date 
                    FROM audit_version 
                    WHERE AuditId = %s AND Version LIKE 'A%'
                    ORDER BY Date DESC LIMIT 1
                    """,
                    [audit_id]
                )
                
                row = cursor.fetchone()
                if row:
                    auditor_version = row[0]
                    json_data = row[1]
                    version_date = row[2]
                    
                    # Create new R version based on this A version
                    latest_version = "R1"  # Start with R1
                    
                    # Parse the JSON data
                    if isinstance(json_data, str):
                        try:
                            version_data = json.loads(json_data)
                        except Exception as e:
                            print(f"ERROR parsing JSON: {str(e)}")
                            return Response({"error": f"Could not parse JSON: {str(e)}"}, 
                                           status=status.HTTP_500_INTERNAL_SERVER_ERROR)
                    else:
                        version_data = json_data
                    
                    print(f"DEBUG: Using auditor version {auditor_version} as base for new reviewer version")
                else:
                    return Response({"error": "No existing version found to update"}, 
                                   status=status.HTTP_404_NOT_FOUND)
        
        # Update version data with reviewer comments
        updated = False
        
        # Update overall comments if provided
        if overall_comments:
            version_data['overall_comments'] = overall_comments
            updated = True
            print(f"DEBUG: Updated overall comments: '{overall_comments}'")
        
        # Update metadata
        if '__metadata__' in version_data:
            metadata = version_data['__metadata__']
            metadata['reviewer_id'] = user_id
            metadata['review_date'] = timezone.now().strftime('%Y-%m-%d %H:%M:%S')
            metadata['version_type'] = 'Reviewer'
            updated = True
        else:
            version_data['__metadata__'] = {
                'reviewer_id': user_id,
                'review_date': timezone.now().strftime('%Y-%m-%d %H:%M:%S'),
                'version_type': 'Reviewer'
            }
            updated = True
        
        # Update individual compliance comments
        for compliance_id, comment_data in comments_data.items():
            # Skip if compliance_id is not in the data
            if compliance_id not in version_data:
                print(f"DEBUG: Compliance ID {compliance_id} not found in version data, skipping")
                continue
            
            # Get the finding data
            finding_data = version_data[compliance_id]
            
            # Update review comments
            if 'comments' in comment_data:
                finding_data['comments'] = comment_data['comments']
                finding_data['review_comments'] = comment_data['comments']
                updated = True
                print(f"DEBUG: Updated comments for compliance {compliance_id}: '{comment_data['comments']}'")
            
            # Update review status
            if 'status' in comment_data:
                review_status = comment_data['status']
                finding_data['review_status'] = review_status
                
                # Update accept_reject field based on review_status
                if review_status == 'Accept':
                    finding_data['accept_reject'] = "1"
                elif review_status == 'Reject':
                    finding_data['accept_reject'] = "2"
                else:
                    finding_data['accept_reject'] = "0"
                
                updated = True
                print(f"DEBUG: Updated status for compliance {compliance_id}: '{review_status}'")
            
            # Save back to version_data
            version_data[compliance_id] = finding_data
        
        if not updated:
            return Response({"message": "No changes were made"}, status=status.HTTP_200_OK)
        
        # Check if this is a new version or update to existing
        is_new_version = latest_version.startswith("R") and not row
        
        # Save the updated version data
        json_data = json.dumps(version_data, indent=2)
        
        # Print the full updated JSON data
        print(f"DEBUG: Updated JSON data for version {latest_version}:\n{json_data}")
        
        with connection.cursor() as cursor:
            if is_new_version:
                # Insert new version
                cursor.execute(
                    """
                    INSERT INTO audit_version (AuditId, Version, ExtractedInfo, UserId, Date)
                    VALUES (%s, %s, %s, %s, NOW())
                    """,
                    [audit_id, latest_version, json_data, user_id,timezone.now()]
                )
                print(f"DEBUG: Created new reviewer version {latest_version}")
            else:
                # Update existing version
                cursor.execute(
                    """
                    UPDATE audit_version 
                    SET ExtractedInfo = %s, UserId = %s, Date = NOW()
                    WHERE AuditId = %s AND Version = %s
                    """,
                    [json_data, user_id, audit_id, latest_version]
                )
                print(f"DEBUG: Updated existing reviewer version {latest_version}")
        
        # Also update the audit_findings table with the same comments
        for compliance_id, comment_data in comments_data.items():
            if 'comments' in comment_data or 'status' in comment_data:
                review_rejected = 1 if comment_data.get('status') == 'Reject' else 0
                review_status = comment_data.get('status', 'In Review')
                review_comments = comment_data.get('comments', '')
                
                cursor.execute(
                    """
                        UPDATE audit_findings
                    SET ReviewRejected = %s,
                        ReviewStatus = %s,
                            ReviewComments = %s,
                        ReviewDate = NOW()
                    WHERE AuditId = %s AND ComplianceId = %s
                    """,
                    [review_rejected, review_status, review_comments, audit_id, compliance_id]
                )
                print(f"DEBUG: Updated audit_findings table for compliance {compliance_id}")
        
        # Update the audit's overall review comments if provided
        if overall_comments:
            cursor.execute(
                """
                UPDATE audit
                SET ReviewComments = %s,
                    ReviewDate = NOW()
                        WHERE AuditId = %s
                """,
                [overall_comments, audit_id]
            )
            print(f"DEBUG: Updated audit.ReviewComments with '{overall_comments}'")
        
        return Response({
            'success': True,
            'message': f"{'Created' if is_new_version else 'Updated'} reviewer version {latest_version}",
            'audit_id': audit_id,
            'version': latest_version,
            'comments_updated': len(comments_data),
            'overall_comments_updated': bool(overall_comments)
        }, status=status.HTTP_200_OK)
    
    except Audit.DoesNotExist:
        return Response({"error": f"Audit with ID {audit_id} not found"}, status=status.HTTP_404_NOT_FOUND)
    except Exception as e:
        print(f"ERROR in debug_save_reviewer_comments: {str(e)}")
        import traceback
        traceback.print_exc()
        return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        
@api_view(['GET'])
def debug_auditor_view_of_reviewer_comments(request, audit_id):
    """Debug endpoint to show what the auditor sees after reviewer rejection"""
    try:
        print(f"=== DEBUG: Auditor view of reviewer comments for audit_id={audit_id} ===")
        
        # Get the audit record
        audit = Audit.objects.get(AuditId=audit_id)
        
        # Check if this audit has been rejected
        is_rejected = audit.Status == 'Rejected' or (audit.ReviewStatus and audit.ReviewStatus in ['3', 'Rejected'])
        
        if not is_rejected:
            return Response({
                'warning': f"Audit {audit_id} has not been rejected",
                'current_status': audit.Status,
                'review_status': audit.ReviewStatus
            }, status=status.HTTP_200_OK)
        
        # Get the latest reviewer version with rejection data
        reviewer_version = None
        reviewer_data = None
        
        with connection.cursor() as cursor:
            cursor.execute(
                """
                SELECT Version, ExtractedInfo, Date 
                FROM audit_version 
                WHERE AuditId = %s AND Version LIKE 'R%'
                ORDER BY Date DESC LIMIT 1
                """,
                [audit_id]
            )
            
            row = cursor.fetchone()
            if row:
                reviewer_version = row[0]
                json_data = row[1]
                version_date = row[2]
                
                # Parse the JSON data
                if isinstance(json_data, str):
                    try:
                        reviewer_data = json.loads(json_data)
                    except Exception as e:
                        print(f"ERROR parsing JSON: {str(e)}")
                        return Response({"error": f"Could not parse JSON: {str(e)}"}, 
                                       status=status.HTTP_500_INTERNAL_SERVER_ERROR)
                else:
                    reviewer_data = json_data
                
                print(f"DEBUG: Found reviewer version {reviewer_version} from {version_date}")
            else:
                print(f"DEBUG: No reviewer version found for audit {audit_id}")
        
        # Format the data for auditor display
        auditor_view = {
            'audit_id': audit_id,
            'audit_status': audit.Status,
            'review_status': audit.ReviewStatus,
            'overall_comments': audit.ReviewComments or reviewer_data.get('overall_comments', '') if reviewer_data else '',
            'rejected_items': []
        }
        
        # Extract rejected items from version data
        if reviewer_data:
            for compliance_id, finding in reviewer_data.items():
                # Skip metadata and overall comments
                if compliance_id in ['__metadata__', 'overall_comments']:
                    continue
                
                if not isinstance(finding, dict):
                    continue
                
                # Check if this item was rejected
                was_rejected = False
                
                if 'review_status' in finding and finding['review_status'] == 'Reject':
                    was_rejected = True
                elif 'accept_reject' in finding and finding['accept_reject'] == "2":
                    was_rejected = True
                
                if was_rejected:
                    # Get the compliance description
                    description = finding.get('description', 'Unknown')
                    
                    # Get reviewer comments
                    comments = finding.get('review_comments', finding.get('comments', ''))
                    
                    auditor_view['rejected_items'].append({
                        'compliance_id': compliance_id,
                        'description': description,
                        'reviewer_comments': comments
                    })
        
        # Also get rejected items from audit_findings table
        with connection.cursor() as cursor:
            cursor.execute(
                """
                SELECT 
                    af.ComplianceId,
                    af.ReviewComments,
                    af.ReviewStatus,
                    c.ComplianceItemDescription
                FROM 
                    audit_findings af
                JOIN
                    compliance c ON af.ComplianceId = c.ComplianceId
                WHERE 
                    af.AuditId = %s AND 
                    (af.ReviewRejected = 1 OR af.ReviewStatus = 'Reject')
                """,
                [audit_id]
            )
            
            for row in cursor.fetchall():
                compliance_id = str(row[0])
                review_comments = row[1]
                review_status = row[2]
                description = row[3]
                
                # Skip if this compliance is already in the list
                if any(item['compliance_id'] == compliance_id for item in auditor_view['rejected_items']):
                    continue
                
                auditor_view['rejected_items'].append({
                    'compliance_id': compliance_id,
                    'description': description,
                    'reviewer_comments': review_comments,
                    'source': 'audit_findings'
                })
        
        # Get all compliances for this audit that need to be addressed
        compliances_to_update = []
        
        with connection.cursor() as cursor:
            cursor.execute(
                """
                SELECT 
                    af.ComplianceId,
                    c.ComplianceItemDescription,
                    af.`Check`,
                    af.Comments,
                    af.Evidence,
                    af.HowToVerify,
                    af.Impact,
                    af.Recommendation,
                    af.DetailsOfFinding,
                    af.MajorMinor,
                    af.ReviewRejected,
                    af.ReviewComments,
                    af.ReviewStatus
                FROM 
                    audit_findings af
                JOIN
                    compliance c ON af.ComplianceId = c.ComplianceId
                WHERE 
                    af.AuditId = %s
                """,
                [audit_id]
            )
            
            columns = [col[0] for col in cursor.description]
            
            for row in cursor.fetchall():
                finding_data = dict(zip(columns, row))
                
                # Only include items that were rejected or are incomplete
                review_rejected = finding_data.get('ReviewRejected') == 1
                review_status = finding_data.get('ReviewStatus') == 'Reject'
                check_incomplete = finding_data.get('Check') in ['0', '1']  # Not started or in progress
                
                if review_rejected or review_status or check_incomplete:
                    compliances_to_update.append(finding_data)
        
        # Get current state and what needs to be done
        auditor_view['compliances_to_update'] = compliances_to_update
        auditor_view['rejected_count'] = len(auditor_view['rejected_items'])
        auditor_view['total_to_update'] = len(compliances_to_update)
        
        return Response(auditor_view, status=status.HTTP_200_OK)
    
    except Audit.DoesNotExist:
        return Response({"error": f"Audit with ID {audit_id} not found"}, status=status.HTTP_404_NOT_FOUND)
    except Exception as e:
        print(f"ERROR in debug_auditor_view_of_reviewer_comments: {str(e)}")
        import traceback
        traceback.print_exc()
        return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        
@api_view(['POST'])
def save_audit_json_version(request, audit_id):
    """
    Save the complete audit data to the audit_version table
    - This function just saves the JSON to the version table
    - It doesn't update the audit status or the audit_findings table
    - Each compliance in the JSON will have review fields initialized with consistent format
    """
    try:
        print(f"\n==== DEBUG: save_audit_json_version called for audit_id: {audit_id} ====")
        
        # Get the audit object
        try:
            audit = Audit.objects.get(pk=audit_id)
        except Audit.DoesNotExist:
            return Response({'error': 'Audit not found'}, status=status.HTTP_404_NOT_FOUND)
        
        # Get user ID from session or request
        user_id = request.session.get('user_id')
        if not user_id and hasattr(audit, 'auditor_id'):
            user_id = audit.auditor_id
        if not user_id:
            # Try to get auditor from audit object
            user_id = audit.Auditor.UserId if hasattr(audit, 'Auditor') else None
        if not user_id:
            user_id = 1050  # Default auditor ID if not found
            
        print(f"DEBUG: Using user_id: {user_id} for saving audit version")
        
        # Get the audit data from request
        audit_data = request.data.get('audit_data', {})
        if not audit_data:
            return Response({'error': 'No audit data provided'}, status=status.HTTP_400_BAD_REQUEST)
        
        print(f"DEBUG: Got audit data with {len(audit_data)} entries")
        
        # Make sure every compliance has consistent review fields initialization
        for key, value in audit_data.items():
            if key != '__metadata__' and isinstance(value, dict):
                # Initialize all review fields consistently for every compliance item
                # This ensures both A and R versions have the same field structure
                
                # Make sure review_status field exists and is initialized
                if 'review_status' not in value:
                    audit_data[key]['review_status'] = 'In Review'
                
                # Make sure review_comments field exists and is initialized
                if 'review_comments' not in value:
                    audit_data[key]['review_comments'] = ''
                
                # Make sure reviewer_comments field exists (alternate field name)
                if 'reviewer_comments' not in value:
                    audit_data[key]['reviewer_comments'] = audit_data[key]['review_comments']
                
                # Make sure accept_reject field exists and is initialized
                if 'accept_reject' not in value:
                    audit_data[key]['accept_reject'] = '0'  # 0=In Review, 1=Accept, 2=Reject
        
        # Add created timestamp to metadata
        if '__metadata__' not in audit_data:
            audit_data['__metadata__'] = {}
        
        audit_data['__metadata__']['created_date'] = timezone.now().strftime('%Y-%m-%d %H:%M:%S')
        audit_data['__metadata__']['version_type'] = 'Auditor'
        audit_data['__metadata__']['user_id'] = user_id
        
        # Convert the audit_data to JSON
        json_data = json.dumps(audit_data)
        json_size = len(json_data)
        print(f"DEBUG: JSON data size: {json_size} bytes")
        
        # Get the next version number (A1, A2, etc.)
        next_version = get_next_version_number(audit_id, "A")
        print(f"DEBUG: Next version number: {next_version}")
        
        # Insert the version data into the audit_version table
        with connection.cursor() as cursor:
            cursor.execute("START TRANSACTION")
            try:
                # Double-check that this version doesn't already exist
                cursor.execute(
                    "SELECT COUNT(*) FROM audit_version WHERE AuditId = %s AND Version = %s",
                    [audit_id, next_version]
                )
                if cursor.fetchone()[0] > 0:
                    # If duplicate, add timestamp to ensure uniqueness
                    timestamp = timezone.now().strftime('%H%M%S')
                    next_version = f"{next_version}_{timestamp}"
                    print(f"DEBUG: Version already exists! Adjusted to ensure uniqueness: {next_version}")
                
                # Insert the new version
                print(f"DEBUG: INSERTING NEW VERSION {next_version} for audit {audit_id} with JSON size {json_size}")
                
                cursor.execute(
                    """
                    INSERT INTO audit_version (AuditId, Version, ExtractedInfo, UserId, Date)
                    VALUES (%s, %s, %s, %s, NOW())
                    """,
                    [audit_id, next_version, json_data, user_id,timezone.now()]
                )
                
                # Commit transaction
                cursor.execute("COMMIT")
                print(f"DEBUG: Successfully committed transaction for version {next_version}")
                
                # Verify the insert worked
                cursor.execute(
                    "SELECT COUNT(*) FROM audit_version WHERE AuditId = %s AND Version = %s",
                    [audit_id, next_version]
                )
                verify_count = cursor.fetchone()[0]
                if verify_count != 1:
                    print(f"WARNING: Verification found {verify_count} records instead of expected 1")
                
            except Exception as e:
                cursor.execute("ROLLBACK")
                print(f"ERROR creating version: {str(e)}")
                import traceback
                traceback.print_exc()
                return Response({'error': f'Database error: {str(e)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
        
        print(f"==== DEBUG: save_audit_json_version completed successfully with version {next_version} ====\n")
        
        return Response({
            'success': True,
            'message': f'Audit data saved as version {next_version}',
            'version': next_version,
            'audit_id': audit_id
        }, status=status.HTTP_200_OK)
        
    except Exception as e:
        print(f"ERROR in save_audit_json_version: {str(e)}")
        import traceback
        traceback.print_exc()
        return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

@api_view(['GET'])
def generate_audit_report(request, audit_id):
    """
    Generate and download an audit report in DOCX format
    """
    try:
        # Get the audit details
        audit = Audit.objects.get(pk=audit_id)
        
        # Check if the audit is in ACCEPTED status (ACCEPTED or APPROVED)
        if audit.status not in ['ACCEPTED', 'APPROVED']:
            return Response({"error": "Report can only be downloaded for accepted or approved audits"}, 
                           status=status.HTTP_400_BAD_REQUEST)
        
        # Get the template path
        template_path = os.path.join(settings.BASE_DIR, 'frontend', 'src', 'assets', 'Template', 'AuditReportTemplate.docx')
        
        if not os.path.exists(template_path):
            return Response({"error": "Report template not found"}, status=status.HTTP_404_NOT_FOUND)
        
        # Count and fetch all audit findings for this audit
        with connection.cursor() as cursor:
            cursor.execute("""
                SELECT 
                    af.AuditFindingId,
                    af.ComplianceId,
                    af.MajorMinor,
                    af.`Check`,
                    af.HowToVerify,
                    af.Evidence,
                    af.DetailsOfFinding,
                    af.Impact,
                    af.Recommendation,
                    af.Comments,
                    af.ReviewDate,
                    c.ComplianceItemCode,
                    c.ComplianceItemDescription,
                    sp.SubPolicyName,
                    p.PolicyName
                FROM 
                    audit_findings af
                JOIN 
                    compliance c ON af.ComplianceId = c.ComplianceId
                JOIN 
                    sub_policy sp ON c.SubPolicyId = sp.SubPolicyId
                JOIN 
                    policy p ON sp.PolicyId = p.PolicyId
                WHERE 
                    af.AuditId = %s
            """, [audit_id])
            
            columns = [col[0] for col in cursor.description]
            findings = [dict(zip(columns, row)) for row in cursor.fetchall()]
        
        # Count the number of findings
        findings_count = len(findings)
        if findings_count == 0:
            return Response({"error": "No findings available for this audit"}, status=status.HTTP_404_NOT_FOUND)
        
        # Create a temporary directory for the process
        import tempfile
        temp_dir = tempfile.mkdtemp()
        
        try:
            # Using docxtpl for template rendering
            from docxtpl import DocxTemplate
            
            # Using docxcompose for combining documents
            from docx import Document
            from docxcompose.composer import Composer
            
            # Load the master document
            master = Document(template_path)
            composer = Composer(master)
            
            # Add audit information to the first page
            # This could be done with a separate template for the cover page
            
            # Process each finding
            for i, finding in enumerate(findings):
                # Create a context dictionary for this finding
                context = {
                    "AuditName": audit.name,
                    "Framework": audit.framework.name if hasattr(audit, 'framework') and audit.framework else "N/A",
                    "Auditor": audit.auditor.name if hasattr(audit, 'auditor') and audit.auditor else "N/A",
                    "Reviewer": audit.reviewer.name if hasattr(audit, 'reviewer') and audit.reviewer else "N/A",
                    "AuditDate": audit.created_at.strftime("%Y-%m-%d") if audit.created_at else "N/A",
                    "MajorMinor": finding.get("MajorMinor", "N/A"),
                    "ComplianceId": finding.get("ComplianceId", "N/A"),
                    "ComplianceCode": finding.get("ComplianceItemCode", "N/A"),
                    "ComplianceDescription": finding.get("ComplianceItemDescription", "N/A"),
                    "PolicyName": finding.get("PolicyName", "N/A"),
                    "SubPolicyName": finding.get("SubPolicyName", "N/A"),
                    "ReviewDate": finding.get("ReviewDate").strftime("%Y-%m-%d %H:%M") if finding.get("ReviewDate") else "N/A",
                    "SubPolicy": finding.get("Check", "N/A"),
                    "HowToVerify": finding.get("HowToVerify", "N/A"),
                    "Proof": finding.get("Evidence", "N/A"),
                    "DetailsOfFinding": finding.get("DetailsOfFinding", "N/A"),
                    "Impact": finding.get("Impact", "N/A"),
                    "Recommendation": finding.get("Recommendation", "N/A"),
                    "Comments": finding.get("Comments", "N/A")
                }
                
                # Create a temp file path for this finding
                temp_file = os.path.join(temp_dir, f"finding_{i}.docx")
                
                # Load the template and render it with the context
                doc = DocxTemplate(template_path)
                doc.render(context)
                doc.save(temp_file)
                
                # Add the rendered document to the master
                finding_doc = Document(temp_file)
                
                # Skip the first append as it's the master document itself
                if i > 0:
                    composer.append(finding_doc)
            
            # Save the final composed document to a bytes buffer
            final_doc_path = os.path.join(temp_dir, "final_report.docx")
            composer.save(final_doc_path)
            
            # Read the final document into a buffer
            with open(final_doc_path, 'rb') as f:
                buffer = io.BytesIO(f.read())
            
            # Create response with the document
            response = HttpResponse(
                buffer.getvalue(),
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            response['Content-Disposition'] = f'attachment; filename=Audit_Report_{audit_id}.docx'
            
            return response
            
        finally:
            # Clean up temporary files
            import shutil
            shutil.rmtree(temp_dir, ignore_errors=True)
        
    except ImportError as e:
        return Response({
            "error": f"Required package not installed: {str(e)}. Please install docxtpl and docxcompose packages."
        }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
    except Audit.DoesNotExist:
        return Response({"error": "Audit not found"}, status=status.HTTP_404_NOT_FOUND)
    except Exception as e:
        import traceback
        traceback.print_exc()
        return Response({"error": str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
